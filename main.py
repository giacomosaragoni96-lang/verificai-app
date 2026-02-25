import streamlit as st
import os
import subprocess
import tempfile
import base64
import re
import io
import time
import google.generativeai as genai
from dotenv import load_dotenv
from supabase import create_client, Client

# ── CONFIGURAZIONE ──────────────────────────────────────────────────────────────
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")
if not API_KEY:
    st.error("⚠️ Chiave API mancante! Crea un file .env con: GOOGLE_API_KEY=la_tua_chiave")
    st.stop()
genai.configure(api_key=API_KEY)
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
SUPABASE_SERVICE_KEY = st.secrets["SUPABASE_SERVICE_KEY"]
supabase_admin: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)

LIMITE_MENSILE = 5  # ← numero massimo di verifiche al mese per utente free

# ── PERSISTENT LOGIN ─────────────────────────────────────────────────────────────
def _ripristina_sessione():
    if st.session_state.get('utente') is not None:
        return
    params = st.query_params
    access_token  = params.get("_at", None)
    refresh_token = params.get("_rt", None)
    if not access_token:
        access_token  = st.session_state.get("_sb_access_token")
        refresh_token = st.session_state.get("_sb_refresh_token")
    if access_token and refresh_token:
        try:
            sess = supabase.auth.set_session(access_token, refresh_token)
            if sess and sess.user:
                st.session_state.utente = sess.user
                new_at = sess.session.access_token
                new_rt = sess.session.refresh_token
                st.session_state["_sb_access_token"] = new_at
                st.session_state["_sb_refresh_token"] = new_rt
                st.query_params["_at"] = new_at
                st.query_params["_rt"] = new_rt
                return
        except Exception:
            st.query_params.pop("_at", None)
            st.query_params.pop("_rt", None)
            st.session_state.pop("_sb_access_token", None)
            st.session_state.pop("_sb_refresh_token", None)

_ripristina_sessione()

# ── AUTENTICAZIONE ──────────────────────────────────────────────────────────────
# ── AUTENTICAZIONE ──────────────────────────────────────────────────────────────
def mostra_auth():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,300;0,400;0,600;0,700;0,900;1,400&display=swap');

    html, body, [data-testid="stAppViewContainer"], .stApp {
        background: #0C0C0B !important;
        font-family: 'DM Sans', sans-serif !important;
    }
    [data-testid="stHeader"],
    [data-testid="stDecoration"],
    [data-testid="stToolbar"],
    #MainMenu, footer { display: none !important; }

    .block-container {
        padding: 0 !important;
        max-width: 480px !important;
        margin: 0 auto !important;
    }
    [data-testid="stMainBlockContainer"] { padding: 0 !important; }

    [data-testid="stTextInput"] input {
        background: #1A1916 !important;
        border: 1.5px solid #2A2926 !important;
        border-radius: 10px !important;
        color: #F5F4EF !important;
        font-family: 'DM Sans', sans-serif !important;
        font-size: 1rem !important;
        padding: 14px 16px !important;
        min-height: 50px !important;
    }
    [data-testid="stTextInput"] input:focus {
        border-color: #D97706 !important;
        box-shadow: 0 0 0 3px rgba(217,119,6,0.18) !important;
        outline: none !important;
    }
    [data-testid="stTextInput"] input::placeholder {
        color: #4A4840 !important;
        opacity: 1 !important;
    }
    [data-testid="stTextInput"] label p {
        color: #C8C6BC !important;
        font-size: 0.85rem !important;
        font-weight: 600 !important;
        font-family: 'DM Sans', sans-serif !important;
    }
    [data-testid="stTabs"] [data-baseweb="tab-list"] {
        background: #1A1916 !important;
        border-radius: 10px !important;
        padding: 4px !important;
        gap: 2px !important;
        border: 1px solid #2A2926 !important;
        margin-bottom: 1.5rem !important;
    }
    [data-testid="stTabs"] [data-baseweb="tab"] {
        border-radius: 7px !important;
        font-size: 0.85rem !important;
        font-weight: 600 !important;
        color: #6B6960 !important;
        padding: 0.45rem 1.2rem !important;
        background: transparent !important;
    }
    [data-testid="stTabs"] [aria-selected="true"] {
        background: #2A2926 !important;
        color: #F5F4EF !important;
    }
    [data-testid="stTabs"] [data-baseweb="tab-highlight"] { display: none !important; }
    div.stButton > button[kind="primary"] {
        background: #D97706 !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 700 !important;
        font-size: 1rem !important;
        min-height: 52px !important;
        box-shadow: 0 2px 20px rgba(217,119,6,0.35) !important;
        width: 100% !important;
        transition: filter 0.15s, transform 0.15s !important;
    }
    div.stButton > button[kind="primary"]:hover {
        filter: brightness(1.1) !important;
        transform: translateY(-1px) !important;
    }
    .stAlert { border-radius: 8px !important; }
    .stAlert p, .stAlert div { font-size: 0.88rem !important; }
    </style>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style="padding:3rem 2rem 0 2rem;text-align:center;">
      <div style="display:inline-flex;align-items:center;gap:7px;
                  background:rgba(217,119,6,0.12);border:1px solid rgba(217,119,6,0.3);
                  border-radius:100px;padding:5px 14px;margin-bottom:1.8rem;">
        <span style="width:6px;height:6px;border-radius:50%;background:#F59E0B;display:inline-block;"></span>
        <span style="font-size:0.72rem;font-weight:700;color:#F59E0B;letter-spacing:0.07em;text-transform:uppercase;">
          Generazione AI · Beta
        </span>
      </div>
      <div style="font-size:3.2rem;font-weight:900;letter-spacing:-0.04em;
                  color:#F5F4EF;line-height:1;margin-bottom:0.5rem;">
        📝 Verific<span style="background:linear-gradient(135deg,#D97706,#FF8C00);
                               -webkit-background-clip:text;-webkit-text-fill-color:transparent;
                               background-clip:text;">AI</span>
      </div>
      <p style="font-size:1rem;color:#8C8A82;font-weight:400;
                margin:0 auto 2rem auto;line-height:1.5;max-width:360px;">
        Crea <strong style="color:#C8C6BC;">verifiche scolastiche professionali</strong>
        in 30 secondi. Tu dici l'argomento, l'AI fa il resto.
      </p>
      <div style="display:flex;flex-wrap:wrap;gap:0.5rem;justify-content:center;margin-bottom:2.5rem;">
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">🧠 Generazione AI</span>
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">📄 PDF & Word</span>
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">🔀 Fila A/B</span>
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">🎯 BES/DSA</span>
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">✅ Soluzioni</span>
        <span style="background:#161614;border:1px solid #2A2926;border-radius:20px;padding:5px 12px;font-size:0.75rem;color:#C8C6BC;">🏫 Tutti i livelli</span>
      </div>
      <div style="background:#111110;border:1px solid #1E1D1A;border-radius:20px;
                  padding:1.8rem 2rem 0.5rem 2rem;text-align:left;margin-bottom:0.5rem;">
        <div style="font-size:1.3rem;font-weight:800;color:#F5F4EF;
                    margin-bottom:0.3rem;letter-spacing:-0.02em;">Inizia subito</div>
        <div style="font-size:0.85rem;color:#6B6960;margin-bottom:0;line-height:1.4;">
          Gratuito durante il periodo Beta · Nessuna carta richiesta
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── FORM ──────────────────────────────────────────────────────────────────
    tab_login, tab_reg, tab_reset = st.tabs(["  Accedi  ", "  Registrati  ", "  Password  "])

    with tab_login:
        st.write("")
        email = st.text_input("Email", key="login_email", placeholder="docente@scuola.it")
        password = st.text_input("Password", type="password", key="login_pass", placeholder="••••••••")
        st.write("")
        if st.button("Accedi →", type="primary", use_container_width=True, key="btn_login"):
            if not email or not password:
                st.warning("Inserisci email e password.")
            else:
                try:
                    res = supabase.auth.sign_in_with_password({"email": email, "password": password})
                    st.session_state.utente = res.user
                    at = res.session.access_token
                    rt = res.session.refresh_token
                    st.session_state["_sb_access_token"] = at
                    st.session_state["_sb_refresh_token"] = rt
                    st.query_params["_at"] = at
                    st.query_params["_rt"] = rt
                    st.rerun()
                except Exception as e:
                    err_str = str(e).lower()
                    if "invalid login" in err_str or "invalid credentials" in err_str:
                        st.warning("Password errata. Usa il tab 'Password' per reimpostarla.")
                    elif "email not confirmed" in err_str:
                        st.warning("Email non confermata. Controlla la tua casella di posta.")
                    elif "user not found" in err_str or "no user" in err_str:
                        st.warning("Nessun account trovato. Registrati prima.")
                    else:
                        st.warning("Accesso non riuscito. Controlla email e password.")
                    time.sleep(2)

    with tab_reg:
        st.write("")
        email    = st.text_input("Email", key="reg_email", placeholder="docente@scuola.it")
        password = st.text_input("Password (min. 6 caratteri)", type="password", key="reg_pass", placeholder="••••••••")
        st.write("")
        if st.button("Crea account gratuito →", type="primary", use_container_width=True, key="btn_reg"):
            if not email or not password:
                st.warning("Inserisci email e password.")
            elif len(password) < 6:
                st.warning("La password deve essere di almeno 6 caratteri.")
            else:
                try:
                    res = supabase.auth.sign_up({"email": email, "password": password})
                    st.session_state.utente = res.user
                    if res.session:
                        at = res.session.access_token
                        rt = res.session.refresh_token
                        st.session_state["_sb_access_token"] = at
                        st.session_state["_sb_refresh_token"] = rt
                        st.query_params["_at"] = at
                        st.query_params["_rt"] = rt
                    st.success("Benvenuto su VerificAI! Account creato.")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"Errore durante la registrazione: {e}")

    with tab_reset:
        st.write("")
        st.markdown("""
        <div style="font-size:0.82rem;color:#8C8A82;line-height:1.5;
                    padding:0.8rem 1rem;background:#161614;border-radius:8px;
                    border-left:2px solid #2A2926;margin-bottom:0.8rem;">
          Inserisci la tua email. Riceverai un link per reimpostare la password.
        </div>
        """, unsafe_allow_html=True)
        email_reset = st.text_input("Email", key="reset_email", placeholder="docente@scuola.it")
        st.write("")
        if st.button("Invia link di reset →", type="primary", use_container_width=True, key="btn_reset"):
            if not email_reset:
                st.warning("Inserisci la tua email.")
            else:
                try:
                    supabase.auth.reset_password_email(email_reset)
                    st.success("Email inviata! Controlla la casella di posta.")
                except Exception as e:
                    st.error(f"Errore nell'invio: {e}")

    # ── FOOTER ────────────────────────────────────────────────────────────────
    st.markdown("""
    <div style="display:flex;align-items:center;gap:0.8rem;
                justify-content:center;padding:2rem 1rem 3rem 1rem;">
      <div style="display:flex;">
        <div style="width:28px;height:28px;border-radius:50%;border:2px solid #0C0C0B;
                    background:linear-gradient(135deg,#D97706,#92400E);
                    display:flex;align-items:center;justify-content:center;
                    font-size:0.6rem;font-weight:700;color:white;">GM</div>
        <div style="width:28px;height:28px;border-radius:50%;border:2px solid #0C0C0B;
                    background:linear-gradient(135deg,#D97706,#92400E);
                    display:flex;align-items:center;justify-content:center;
                    font-size:0.6rem;font-weight:700;color:white;margin-left:-7px;">AR</div>
        <div style="width:28px;height:28px;border-radius:50%;border:2px solid #0C0C0B;
                    background:linear-gradient(135deg,#D97706,#92400E);
                    display:flex;align-items:center;justify-content:center;
                    font-size:0.6rem;font-weight:700;color:white;margin-left:-7px;">FL</div>
        <div style="width:28px;height:28px;border-radius:50%;border:2px solid #0C0C0B;
                    background:linear-gradient(135deg,#444,#222);
                    display:flex;align-items:center;justify-content:center;
                    font-size:0.6rem;font-weight:700;color:white;margin-left:-7px;">+</div>
      </div>
      <div style="font-size:0.78rem;color:#6B6960;line-height:1.4;">
        Già usato da docenti di tutta Italia.<br>
        <strong style="color:#C8C6BC;">Gratis durante il periodo Beta.</strong>
      </div>
    </div>
    """, unsafe_allow_html=True)


if 'utente' not in st.session_state:
    st.session_state.utente = None

if st.session_state.utente is None:
    mostra_auth()
    st.stop()

APP_NAME    = "VerificAI"
APP_ICON    = "📝"
APP_TAGLINE = "Crea verifiche su misura in pochi secondi"
SHARE_URL   = "https://verificai.streamlit.app"
FEEDBACK_FORM_URL = "https://forms.gle/KNu8v8iDVUiGkQUL8"

MODELLI_DISPONIBILI = {
    "⚡ Flash 2.5 Lite (velocissimo)": "gemini-2.5-flash-lite",
    "⚡ Flash 2.5 (bilanciato)":        "gemini-2.5-flash",
    "🧠 Pro 2.5 (massima qualità)":     "gemini-2.5-pro",
}

THEMES = {
    "light": {
        "bg":           "#F7F7F5",
        "bg2":          "#FFFFFF",
        "card":         "#FFFFFF",
        "card2":        "#F0EFEB",
        "border":       "#E5E4DF",
        "border2":      "#D1D0CA",
        "text":         "#1A1915",
        "text2":        "#4A4840",
        "muted":        "#8C8A82",
        "accent":       "#D97706",
        "accent_light": "#FEF3C7",
        "accent2":      "#059669",
        "success":      "#059669",
        "warn":         "#D97706",
        "err":          "#DC2626",
        "shadow":       "0 1px 3px rgba(0,0,0,.08), 0 1px 2px rgba(0,0,0,.05)",
        "shadow_md":    "0 4px 12px rgba(0,0,0,.08)",
        "input_bg":     "#FFFFFF",
        "hover":        "#F0EFEB",
    },
    "dark": {
        "bg":           "#0F0F0E",
        "bg2":          "#1A1916",
        "card":         "#1A1916",
        "card2":        "#232320",
        "border":       "#2E2D28",
        "border2":      "#3D3C36",
        "text":         "#F5F4EF",
        "text2":        "#C8C6BC",
        "muted":        "#6B6960",
        "accent":       "#F59E0B",
        "accent_light": "#292215",
        "accent2":      "#10B981",
        "success":      "#10B981",
        "warn":         "#F59E0B",
        "err":          "#EF4444",
        "shadow":       "0 1px 3px rgba(0,0,0,.3)",
        "shadow_md":    "0 4px 16px rgba(0,0,0,.4)",
        "input_bg":     "#232320",
        "hover":        "#232320",
    },
}

if "theme" not in st.session_state:
    st.session_state.theme = "light"
T = THEMES[st.session_state.theme]

SCUOLE = [
    "Generico — adatta alle istruzioni",
    "Scuola Primaria (Elementari)",
    "Scuola Secondaria I grado (Medie)",
    "Liceo Scientifico",
    "Liceo Classico",
    "Liceo Linguistico",
    "Liceo delle Scienze Umane",
    "Liceo Artistico",
    "Istituto Tecnico Tecnologico/Industriale",
    "Istituto Tecnico Economico",
    "Istituto Tecnico Agrario/Ambientale",
    "Istituto Professionale",
]

CALIBRAZIONE_SCUOLA = {
    "Generico — adatta alle istruzioni": (
        "Livello NON specificato: adatta autonomamente difficoltà, registro linguistico e complessità "
        "in base all'argomento e alle istruzioni del docente. "
        "Se l'argomento suggerisce un livello (es. 'derivate' → superiori, 'addizioni' → primaria), calibra di conseguenza. "
        "Usa un linguaggio chiaro, diretto e professionale. Nessun vincolo di scuola."
    ),
    "Scuola Primaria (Elementari)": (
        "Target: 6-11 anni. Linguaggio ludico-concreto. "
        "Contesto: vita quotidiana familiare, gioco, spesa. "
        "Usa frasi brevi e numeri entro il 1000. Evita simboli astratti, preferisci il testo narrativo."
    ),
    "Scuola Secondaria I grado (Medie)": (
        "Target: 11-14 anni. Linguaggio in transizione verso il tecnico. "
        "Contesto: scuola, sport, socialità, prime esplorazioni scientifiche. "
        "Difficoltà bilanciata: calcolo procedurale e primi problemi logici con frazioni e variabili."
    ),
    "Liceo Scientifico": (
        "Target: 14-19 anni. Linguaggio rigoroso e accademico. "
        "Contesto: ricerca scientifica, astrazione pura, modellizzazione complessa. "
        "Livello elevato: stimola il ragionamento deduttivo e la giustificazione dei passaggi."
    ),
    "Liceo Classico": (
        "Target: 14-19 anni. Linguaggio formale ma DIRETTO e CHIARO. "
        "Contesto: storia, letteratura, filosofia, lingue classiche. "
        "IMPORTANTE: registro colto ma non aulico. Frasi dirette, domande nette. "
        "La profondità sta nel contenuto, NON nel lessico. Evita perifrasi, paroloni rari e fronzoli retorici. "
        "Una domanda ben posta vale più di tre righe di introduzione filosofica."
    ),
    "Liceo Linguistico": (
        "Target: 14-19 anni. Linguaggio chiaro e internazionale. "
        "Contesto: lingue straniere, cultura, comunicazione, letteratura comparata. "
        "Stile neutro e diretto. Privilegia la chiarezza espositiva. "
        "Evita tecnicismi inutili e riferimenti troppo specialistici."
    ),
    "Liceo delle Scienze Umane": (
        "Target: 14-19 anni. Linguaggio accessibile con riferimenti alle scienze sociali. "
        "Contesto: psicologia, sociologia, pedagogia, antropologia. "
        "Stile semplice e concreto, con esempi pratici tratti dalla realtà quotidiana. "
        "Evita gergo accademico pesante."
    ),
    "Liceo Artistico": (
        "Target: 14-19 anni. Linguaggio descrittivo e visivo. "
        "Contesto: storia dell'arte, teoria, tecniche artistiche, progettazione. "
        "Stile diretto e pratico. Privilegia domande operative e descrittive. "
        "Evita astrattismi: chiedi di osservare, descrivere, confrontare opere concrete."
    ),
    "Istituto Tecnico Tecnologico/Industriale": (
        "Target: 14-19 anni. Linguaggio tecnico-professionale. "
        "Contesto: laboratorio, tecnologia, elettronica, meccanica, informatica applicata. "
        "Enfasi su applicazione pratica, dati reali, scenari lavorativi concreti. "
        "Esercizi con misure, tolleranze, schemi, calcoli ingegneristici di base."
    ),
    "Istituto Tecnico Economico": (
        "Target: 14-19 anni. Linguaggio economico-aziendale. "
        "Contesto: azienda, contabilità, economia, diritto commerciale, marketing. "
        "Privilegia casi aziendali reali, calcoli su bilanci, problemi di gestione. "
        "Stile professionale ma accessibile."
    ),
    "Istituto Tecnico Agrario/Ambientale": (
        "Target: 14-19 anni. Linguaggio tecnico-naturalistico. "
        "Contesto: agricoltura, ambiente, biologia applicata, chimica agraria, territorio. "
        "Privilegia esempi concreti legati a colture, ecosistemi, analisi del suolo, sostenibilità. "
        "Stile pratico e operativo."
    ),
    "Istituto Professionale": (
        "Target: 14-19 anni. Linguaggio pratico e operativo. "
        "Contesto: situazioni lavorative simulate, compiti di realtà, problem solving guidato. "
        "Suddividi i problemi complessi in step chiari ed espliciti. "
        "Domande brevi, dirette, con un unico obiettivo per volta."
    ),
}

MATERIE = [
    "Matematica", "Fisica", "Chimica", "Biologia", "Scienze della Terra",
    "Italiano", "Storia", "Geografia", "Latino", "Greco",
    "Inglese", "Francese", "Spagnolo", "Tedesco",
    "Filosofia", "Storia dell'Arte", "Musica",
    "Informatica", "Economia", "Diritto",
    "Educazione Civica", "Scienze Motorie",
]

NOTE_PLACEHOLDER = {
    "Matematica":         "es. Includi un esercizio sul teorema di Pitagora e due problemi algebrici.",
    "Fisica":             "es. Un esercizio sulla seconda legge di Newton, uno sul moto uniformemente accelerato.",
    "Chimica":            "es. Bilanciamento di reazioni ed esercizi sulla mole.",
    "Biologia":           "es. Struttura della cellula e ciclo cellulare. Privilegia la comprensione.",
    "Scienze della Terra":"es. Tettonica a placche e ciclo delle rocce. Schema da completare.",
    "Italiano":           "es. Analisi del testo narrativo, figure retoriche. Testo di circa 15 righe.",
    "Storia":             "es. Prima Guerra Mondiale: cause, fasi, conseguenze. Includi una fonte.",
    "Geografia":          "es. Climi e biomi. Una domanda su cartina muta e una su dati demografici.",
    "Latino":             "es. Versione dal De Bello Gallico (40-50 parole). Declinazioni I e II.",
    "Greco":              "es. Versione da Lisia (30-40 parole). Presente e imperfetto attivo.",
    "Inglese":            "es. Reading comprehension (150 parole), present perfect, short writing task.",
    "Francese":           "es. Compréhension écrite, passé composé, production écrite guidée.",
    "Spagnolo":           "es. Comprensión lectora, pretérito indefinido vs imperfecto, escritura breve.",
    "Tedesco":            "es. Leseverstehen, Perfekt und Präteritum, kurze Schreibaufgabe.",
    "Filosofia":          "es. Problema mente-corpo in Cartesio vs Spinoza. Una domanda aperta e una di definizione.",
    "Storia dell'Arte":   "es. Impressionismo: tecnica e confronto Monet/Renoir. Analisi di un'opera.",
    "Musica":             "es. Sonata classica e confronto Mozart/Beethoven. Domande di teoria.",
    "Informatica":        "es. Bubble sort e selection sort. Pseudocodice e complessità computazionale.",
    "Economia":           "es. Domanda e offerta, elasticità dei prezzi. Un esercizio numerico.",
    "Diritto":            "es. Fonti del diritto e gerarchia normativa. Un caso pratico.",
    "Educazione Civica":  "es. Struttura della Costituzione italiana. Diritti e doveri dei cittadini.",
    "Scienze Motorie":    "es. Apparato muscolare e scheletrico. Norme di sicurezza in palestra.",
}

TIPI_ESERCIZIO = ["Aperto", "Scelta multipla", "Vero/Falso", "Completamento"]

# ── FUNZIONI ───────────────────────────────────────────────────────────────────

# ── NUOVO: contatore verifiche mensili ────────────────────────────────────────
def _get_verifiche_mese(user_id):
    """Restituisce il numero di verifiche generate dall'utente nel mese corrente.
    Conta ANCHE le verifiche eliminate dallo storico (soft-delete) — il contatore
    non deve scendere se l'utente cancella una verifica."""
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc)
    primo_mese = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0).isoformat()
    try:
        # Non filtriamo su deleted_at: contiamo TUTTE le generazioni del mese
        res = supabase_admin.table("verifiche_storico") \
            .select("id", count="exact") \
            .eq("user_id", user_id) \
            .gte("created_at", primo_mese) \
            .execute()
        return res.count or 0
    except Exception:
        return 0


def _giorni_al_reset():
    """Restituisce (giorni, ore) al primo del mese prossimo."""
    from datetime import datetime, timezone, timedelta
    import calendar
    now = datetime.now(timezone.utc)
    # Primo giorno del mese prossimo
    if now.month == 12:
        reset = now.replace(year=now.year+1, month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
    else:
        reset = now.replace(month=now.month+1, day=1, hour=0, minute=0, second=0, microsecond=0)
    delta = reset - now
    giorni = delta.days
    ore    = delta.seconds // 3600
    return giorni, ore


def parse_esercizi(latex):
    esercizi = []
    ex_blocks = re.split(r'\\subsection\*\{', latex)[1:]
    for i, block in enumerate(ex_blocks):
        header_match = re.match(r'([^}]*)', block)
        header = header_match.group(1) if header_match else f"Es {i+1}"
        if any(x in header for x in ["Sezione", "Chiusa", "Risposta", "Quesiti"]):
            num_label = "A"
        else:
            num_match = re.search(r'(?:Esercizio\s+)?(\d+)', header)
            num_label = num_match.group(1) if num_match else str(i + 1)

        items_found = []
        lines = block.split('\n')
        lettere = 'abcdefghijklmnopqrstuvwxyz'
        lettera_idx = 0

        for li, line in enumerate(lines):
            item_label_match = re.search(r'\\item\[([^\]]+)\]', line)
            item_plain_match = re.search(r'\\item(?!\[)', line)

            if not item_label_match and not item_plain_match:
                continue

            if item_label_match:
                raw_label = item_label_match.group(1).replace('*', '').strip()
            else:
                raw_label = lettere[lettera_idx % 26] + ")"
                lettera_idx += 1

            window_lines = []
            for lj in range(li, min(li + 15, len(lines))):
                if lj > li and (re.search(r'\\item(?:\[|(?!\w))', lines[lj])):
                    break
                window_lines.append(lines[lj])
            search_window = '\n'.join(window_lines)
            search_window = re.sub(
                r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', '',
                search_window, flags=re.DOTALL
            )

            pt_match = re.search(
                r'[\(\[]?\s*(\d+(?:[.,]\d+)?)\s*(?:pt|punt[io]|p\.?)\s*[\)\]]?',
                search_window, re.IGNORECASE
            )
            if not pt_match:
                continue

            punti = pt_match.group(1)
            items_found.append((raw_label, punti))

        if items_found:
            esercizi.append({'num': num_label, 'items': items_found})
        else:
            # Esercizio senza \item ma con punteggio nel testo del blocco (es. "Determina X. (25 pt)")
            pt_global = re.search(
                r'[\(\[]?\s*(\d+(?:[.,]\d+)?)\s*(?:pt|punt[io]|p\.?)\s*[\)\]]?',
                block, re.IGNORECASE
            )
            if pt_global:
                esercizi.append({'num': num_label, 'items': [('', pt_global.group(1))]})
    return esercizi

def build_griglia_latex(esercizi, punti_totali):
    if not esercizi:
        return ""
    col_spec = "|l|" + "".join("c|" * len(ex['items']) for ex in esercizi) + "c|"
    row_es = "\\textbf{Es.}" + "".join(
        f" & \\multicolumn{{{len(ex['items'])}}}{{c||}}{{\\textbf{{{ex['num']}}}}}"
        for ex in esercizi
    ) + " & \\textbf{Tot} \\\\ \\hline"
    row_sotto = "\\textbf{Sotto.}" + "".join(
        f" & {label if label else '—'}" for ex in esercizi for label, _ in ex['items']
    ) + " & \\\\ \\hline"
    row_max = "\\textbf{Max}" + "".join(
        f" & {pts}" for ex in esercizi for _, pts in ex['items']
    ) + f" & {punti_totali} \\\\ \\hline"
    total_cols = sum(len(ex['items']) for ex in esercizi) + 1
    row_punti = "\\textbf{Punti}" + " &" * total_cols + " \\\\ \\hline"
    return (
        "% GRIGLIA\n\\begin{center}\n\\textbf{Griglia di Valutazione}\\\\[0.3cm]\n"
        "{\\renewcommand{\\arraystretch}{1.8}\n"
        f"\\adjustbox{{max width=\\textwidth}}{{\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
        f"{row_es}\n{row_sotto}\n{row_max}\n{row_punti}\n\\end{{tabular}}\n}}}}\n\\end{{center}}"
    )


def fix_items_environment(latex):
    import re as _r
    lines = latex.split('\n')
    result = []
    list_depth = 0
    in_bare_block = False
    for line in lines:
        stripped = line.strip()
        list_opens  = len(_r.findall(r'\\begin\{(?:enumerate|itemize)', line))
        list_closes = len(_r.findall(r'\\end\{(?:enumerate|itemize)', line))
        is_bare_item = bool(_r.match(r'\\item\[', stripped)) and list_depth == 0
        if is_bare_item and not in_bare_block:
            result.append(r'\begin{enumerate}[a)]')
            in_bare_block = True
        elif not is_bare_item and in_bare_block:
            result.append(r'\end{enumerate}')
            in_bare_block = False
        result.append(line)
        list_depth = max(0, list_depth + list_opens - list_closes)
    if in_bare_block:
        result.append(r'\end{enumerate}')
    return '\n'.join(result)

def rimuovi_vspace_corpo(latex):
    idx = latex.find('\\subsection*')
    if idx == -1:
        return latex
    preambolo = latex[:idx]
    corpo = latex[idx:]
    corpo = re.sub(r'\\vspace\*?\{[^}]*\}', '', corpo)
    corpo = re.sub(r'\\hspace\*?\{[^}]*\}', '', corpo)
    corpo = re.sub(r'\\(?:big|med|small)skip\b', '', corpo)
    corpo = re.sub(r'\n{3,}', '\n\n', corpo)
    return preambolo + corpo

def pulisci_corpo_latex(testo):
    idx = testo.find('\\subsection*')
    if idx == -1:
        testo = re.sub(r'^.*?\\begin\{document\}[^\n]*\n?', '', testo, flags=re.DOTALL)
        while re.match(r'^\s*\\begin\{center\}', testo):
            testo = re.sub(r'^\s*\\begin\{center\}.*?\\end\{center\}\s*', '', testo, flags=re.DOTALL)
    else:
        testo = testo[idx:]
    testo = re.sub(r'\\end\{document\}.*$', '', testo, flags=re.DOTALL).rstrip()
    testo += "\n\\end{document}"
    return testo

def rimuovi_punti_subsection(latex):
    latex = re.sub(
        r'(\\subsection\*\{[^}]*\}[^\n]*)\s*\((\d+(?:[.,]\d+)?)\s*pt\)',
        r'\1',
        latex
    )
    latex = re.sub(
        r'(\\subsection\*\{[^}]*\})\s*\n\s*\(\d+(?:[.,]\d+)?\s*pt\)\s*\n',
        r'\1\n',
        latex
    )
    return latex

def riscala_punti(latex, punti_totali_target):
    pattern = re.compile(r'\((\d+(?:[.,]\d+)?)\s*pt\)')
    matches = list(pattern.finditer(latex))
    if not matches:
        return latex

    valori = [float(m.group(1).replace(',', '.')) for m in matches]
    somma_attuale = sum(valori)
    if somma_attuale == 0:
        return latex

    fattore = punti_totali_target / somma_attuale
    nuovi_valori = [v * fattore for v in valori]

    nuovi_interi = [int(v) for v in nuovi_valori]
    resti = [(nuovi_valori[i] - nuovi_interi[i], i) for i in range(len(nuovi_valori))]
    differenza = punti_totali_target - sum(nuovi_interi)
    resti.sort(reverse=True)
    for i in range(int(round(differenza))):
        nuovi_interi[resti[i][1]] += 1

    risultato = latex
    offset = 0
    for i, m in enumerate(matches):
        vecchio = m.group(0)
        nuovo = f"({nuovi_interi[i]} pt)"
        start = m.start() + offset
        end = m.end() + offset
        risultato = risultato[:start] + nuovo + risultato[end:]
        offset += len(nuovo) - len(vecchio)

    return risultato

def inietta_griglia(latex, punti_totali):
    latex = re.sub(
        r'(\\vspace\{[^}]+\}\s*)?% GRIGLIA.*?\\end\{center\}',
        '', latex, flags=re.DOTALL
    )
    latex = re.sub(
        r'(\\vspace\{[^}]+\}\s*)?\\begin\{center\}\s*\\textbf\{Griglia[^}]*\}.*?\\end\{center\}',
        '', latex, flags=re.DOTALL
    )

    esercizi = parse_esercizi(latex)
    if not esercizi:
        return latex

    griglia = build_griglia_latex(esercizi, punti_totali)

    try:
        tot_reale = sum(
            float(pts.replace(',', '.'))
            for ex in esercizi for _, pts in ex['items']
        )
        if abs(tot_reale - punti_totali) > 0.5:
            griglia = build_griglia_latex(esercizi, int(tot_reale) if tot_reale == int(tot_reale) else round(tot_reale, 1))
    except Exception:
        pass

    if "\\end{document}" in latex:
        return latex.replace("\\end{document}", f"\n\\vfill\n{griglia}\n\\end{{document}}")
    else:
        return latex + f"\n\\vfill\n{griglia}\n\\end{{document}}"


def compila_pdf(codice_latex):
    with tempfile.TemporaryDirectory() as tmpdir:
        tex = os.path.join(tmpdir, "v.tex")
        pdf = os.path.join(tmpdir, "v.pdf")
        with open(tex, "w", encoding="utf-8") as f:
            f.write(codice_latex)
        r = subprocess.run(
            ["pdflatex", "-interaction=nonstopmode", "-output-directory", tmpdir, tex],
            capture_output=True
        )
        if os.path.exists(pdf):
            return open(pdf, "rb").read(), None
        return None, r.stdout.decode()


def pdf_to_images_bytes(pdf_bytes):
    try:
        from pdf2image import convert_from_bytes as cfb
        pages = cfb(pdf_bytes, dpi=150)
        out = []
        for p in pages:
            buf = io.BytesIO(); p.save(buf, "PNG"); out.append(buf.getvalue())
        return out, None
    except Exception:
        pass
    try:
        with tempfile.TemporaryDirectory() as d:
            src = os.path.join(d, "in.pdf")
            with open(src, "wb") as f:
                f.write(pdf_bytes)
            subprocess.run(
                ["pdftoppm", "-png", "-r", "150", src, os.path.join(d, "p")],
                capture_output=True, check=True
            )
            imgs = sorted(f for f in os.listdir(d) if f.startswith("p") and f.endswith(".png"))
            if imgs:
                return [open(os.path.join(d, f), "rb").read() for f in imgs], None
    except Exception as e:
        return None, str(e)
    return None, "pdf2image non installato e pdftoppm non trovato."


def _make_tc_xml(text, width_dxa, bold=False, gridSpan=1):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    tc = _OE('w:tc')
    tcPr = _OE('w:tcPr')
    tc.append(tcPr)
    if gridSpan > 1:
        gs = _OE('w:gridSpan')
        gs.set(_qn('w:val'), str(gridSpan))
        tcPr.append(gs)
    tcW = _OE('w:tcW')
    tcW.set(_qn('w:w'), str(width_dxa))
    tcW.set(_qn('w:type'), 'dxa')
    tcPr.append(tcW)
    p = _OE('w:p'); tc.append(p)
    pPr = _OE('w:pPr'); p.append(pPr)
    jc = _OE('w:jc'); jc.set(_qn('w:val'), 'center'); pPr.append(jc)
    if text:
        r_el = _OE('w:r'); p.append(r_el)
        if bold:
            rPr = _OE('w:rPr'); b_el = _OE('w:b'); rPr.append(b_el); r_el.append(rPr)
        t = _OE('w:t'); t.text = str(text); r_el.append(t)
    return tc


def _fix_tbl_grid(tbl_el, col_widths, _qn, _OE):
    old_grid = tbl_el.find(_qn('w:tblGrid'))
    if old_grid is not None:
        tbl_el.remove(old_grid)
    new_grid = _OE('w:tblGrid')
    for w in col_widths:
        gc = _OE('w:gridCol')
        gc.set(_qn('w:w'), str(max(1, w)))
        new_grid.append(gc)
    tbl_pr = tbl_el.find(_qn('w:tblPr'))
    if tbl_pr is not None:
        tbl_pr.addnext(new_grid)
    else:
        tbl_el.insert(0, new_grid)


def _build_griglia_xml(doc, row_es, row_sotto, row_max, PAGE_W_DXA=9638):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    n_cols = len(row_es)
    row_punti = ["Punti"] + [""] * (n_cols - 1)

    first_col = 1400
    last_col = 900
    n_mid = max(1, n_cols - 2)
    available = PAGE_W_DXA - first_col - last_col
    mid_col = max(400, available // n_mid)
    col_widths = [first_col] + [mid_col] * n_mid + [last_col]
    diff = PAGE_W_DXA - sum(col_widths)
    col_widths[-1] += diff

    def _setup_tbl(tbl, total_w, widths):
        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(_qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = _OE('w:tblPr'); tbl_el.insert(0, tbl_pr)
        tbl_w = _OE('w:tblW')
        tbl_w.set(_qn('w:w'), str(total_w)); tbl_w.set(_qn('w:type'), 'dxa')
        ex = tbl_pr.find(_qn('w:tblW'))
        if ex is not None: tbl_pr.remove(ex)
        tbl_pr.append(tbl_w)
        tbl_lay = _OE('w:tblLayout'); tbl_lay.set(_qn('w:type'), 'fixed')
        ex2 = tbl_pr.find(_qn('w:tblLayout'))
        if ex2 is not None: tbl_pr.remove(ex2)
        tbl_pr.append(tbl_lay)
        tbl_cm = _OE('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            cm_el = _OE(f'w:{side}')
            cm_el.set(_qn('w:w'), '50'); cm_el.set(_qn('w:type'), 'dxa')
            tbl_cm.append(cm_el)
        ex3 = tbl_pr.find(_qn('w:tblCellMar'))
        if ex3 is not None: tbl_pr.remove(ex3)
        tbl_pr.append(tbl_cm)
        _fix_tbl_grid(tbl_el, widths, _qn, _OE)

    def _fill_cell(cell, text, bold=False, w=None, font_pt=9):
        cell.text = str(text)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(font_pt)
            if bold: run.bold = True
        if w is not None:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = _OE('w:tcW'); tcW.set(_qn('w:w'), str(w)); tcW.set(_qn('w:type'), 'dxa')
            ex = tcPr.find(_qn('w:tcW'))
            if ex is not None: tcPr.remove(ex)
            tcPr.append(tcW)

    tbl = doc.add_table(rows=4, cols=n_cols)
    tbl.style = 'Table Grid'
    _setup_tbl(tbl, PAGE_W_DXA, col_widths)

    for r_idx, riga in enumerate([row_sotto, row_max, row_punti], start=1):
        for c_idx in range(n_cols):
            val = riga[c_idx] if c_idx < len(riga) else ''
            _fill_cell(tbl.cell(r_idx, c_idx), val,
                       bold=(c_idx == 0),
                       w=col_widths[c_idx] if c_idx < len(col_widths) else mid_col)

    groups = []
    c = 0
    while c < n_cols:
        val = row_es[c] if c < len(row_es) else ''
        if c == 0 or c == n_cols - 1:
            groups.append((val, 1, col_widths[c]))
            c += 1
        else:
            span = 1
            while (c + span < n_cols - 1 and
                   c + span < len(row_es) and
                   row_es[c + span] == val):
                span += 1
            total_w = sum(col_widths[c:c + span])
            groups.append((val, span, total_w))
            c += span

    tr_el = tbl.rows[0]._tr
    for tc_el in list(tr_el.findall(_qn('w:tc'))):
        tr_el.remove(tc_el)
    for label, span, width in groups:
        tr_el.append(_make_tc_xml(label, width, bold=True, gridSpan=span))

    return tbl


def _strip_latex_math(text):
    import re as _r
    if not text: return text
    text = _r.sub(r'\$\$(.+?)\$\$', lambda m: m.group(1).strip(), text, flags=_r.DOTALL)
    text = _r.sub(r'\$(.+?)\$',    lambda m: m.group(1).strip(), text)
    text = _r.sub(r'\\mathcal\{([^}]+)\}', r'\1', text)
    text = _r.sub(r'\\mathbf\{([^}]+)\}',  r'\1', text)
    text = _r.sub(r'\\mathrm\{([^}]+)\}',  r'\1', text)
    text = _r.sub(r'\\text\{([^}]+)\}',    r'\1', text)
    text = _r.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = _r.sub(r'\\sqrt\{([^}]+)\}',    r'sqrt(\1)', text)
    text = _r.sub(r'\^\{([^}]+)\}',         r'^\1', text)
    text = _r.sub(r'_\{([^}]+)\}',          r'_\1', text)
    for fr, to in [
        ("\\\\leq", "\u2264"), ("\\\\geq", "\u2265"), ("\\\\neq", "\u2260"),
        ("\\\\approx","\u2248"), ("\\\\cdot","\u00b7"), ("\\\\times","\u00d7"),
        ("\\\\pm","\u00b1"),  ("\\\\infty","\u221e"), ("\\\\alpha","\u03b1"),
        ("\\\\beta","\u03b2"), ("\\\\gamma","\u03b3"), ("\\\\delta","\u03b4"),
        ("\\\\theta","\u03b8"), ("\\\\pi","\u03c0"), ("\\\\sin","sin"),
        ("\\\\cos","cos"), ("\\\\tan","tan"), ("\\\\log","log"),
        ("\\\\ln","ln"), ("\\\\lim","lim"), ("\\\\forall","\u2200"),
        ("\\\\exists","\u2203"), ("\\\\in","\u2208"), ("\\\\notin","\u2209"),
        ("\\\\subset","\u2282"), ("\\\\cup","\u222a"), ("\\\\cap","\u2229"),
        ("\\\\emptyset","\u2205"), ("^2","\u00b2"), ("^3","\u00b3"),
        ("\\\\left(","("), ("\\\\right)",")"),
        ("\\\\left[","["), ("\\\\right]","]"),
    ]:
        text = text.replace(fr, to)
    text = _r.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = _r.sub(r'\\[a-zA-Z]+', '', text)
    text = text.replace('{', '').replace('}', '')
    return text.strip()


def _clean_latex_line(text):
    import re as _r
    if not text:
        return ''
    text = _r.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}', '[Grafico]', text, flags=_r.DOTALL)
    text = _r.sub(r'\\begin\{axis\}.*?\\end\{axis\}', '[Grafico]', text, flags=_r.DOTALL)
    text = _r.sub(r'\\vspace\*?\{[^}]*\}', '', text)
    text = _r.sub(r'\\hspace\*?\{[^}]*\}', '', text)
    text = _r.sub(r'\\noindent\b', '', text)
    text = _r.sub(r'\\newline\b', '', text)
    text = _r.sub(r'\\\\(\s)', r'\1', text)
    text = _r.sub(r'\\\\\s*$', '', text, flags=_r.MULTILINE)
    for cmd in ('textbf', 'textit', 'emph', 'underline', 'textrm', 'texttt'):
        text = _r.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', text)
    text = _r.sub(r'\\(?:small|large|Large|LARGE|huge|Huge|normalsize|footnotesize)\b', '', text)
    return _strip_latex_math(text).strip()


def _parse_latex_to_data(codice_latex):
    """
    Parsa il codice LaTeX e restituisce un dict con titolo, intestazione e lista esercizi.
    Gestisce sia \item[label] con etichetta esplicita sia \item senza etichetta
    all'interno di \begin{enumerate}[a)] o \begin{enumerate}[i)].
    """
    import re as _r
    LETTERE = 'abcdefghijklmnopqrstuvwxyz'
    ROMANI  = ['i','ii','iii','iv','v','vi','vii','viii','ix','x',
               'xi','xii','xiii','xiv','xv','xvi','xvii','xviii','xix','xx']

    data = {'titolo': '', 'intestazione_nota': '', 'esercizi': []}

    # ── titolo e nota ────────────────────────────────────────────────────────
    m = _r.search(r'\\textbf\{\\large ([^}]+)\}', codice_latex)
    if m:
        data['titolo'] = _clean_latex_line(m.group(1))

    m2 = _r.search(r'\\textit\{\\small ([^}]+)\}', codice_latex)
    if m2:
        data['intestazione_nota'] = m2.group(1).strip()

    # ── corpo dopo \end{center} ──────────────────────────────────────────────
    body_start = codice_latex.find('\\end{center}')
    body_start = (body_start + len('\\end{center}')) if body_start != -1 else 0
    corpus = codice_latex[body_start:].replace('\\end{document}', '')

    # ── split per esercizio ──────────────────────────────────────────────────
    blocks = _r.split(r'\\subsection\*\s*\{', corpus)

    for block in blocks[1:]:
        # estrai il titolo dell'esercizio (contenuto fino alla prima } non appaiate)
        brace_depth = 0
        header_end  = 0
        for ci, ch in enumerate(block):
            if   ch == '{': brace_depth += 1
            elif ch == '}':
                if brace_depth == 0:
                    header_end = ci
                    break
                brace_depth -= 1
        titolo_ex = _clean_latex_line(block[:header_end])
        body = block[header_end + 1:]

        # rimuovi grafica tikz
        body = _r.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
                      '\n[Grafico]\n', body, flags=_r.DOTALL)
        body = _r.sub(r'\\begin\{axis\}.*?\\end\{axis\}',
                      '\n[Grafico]\n', body, flags=_r.DOTALL)

        # ── testo introduttivo (prima del primo \begin{enumerate/itemize} o \item) ──
        first_env = len(body)
        for marker in [r'\begin{enumerate}', r'\begin{itemize}', r'\item']:
            idx = body.find(marker)
            if idx != -1 and idx < first_env:
                first_env = idx
        raw_intro = body[:first_env]
        raw_intro = _r.sub(r'\\begin\{tabular\}.*?\\end\{tabular\}', '', raw_intro, flags=_r.DOTALL)
        raw_intro = _r.sub(r'\\begin\{center\}.*?\\end\{center\}',   '', raw_intro, flags=_r.DOTALL)
        testo_intro = _clean_latex_line(raw_intro)

        sottopunti = []

        # ── helper: estrai items da un blocco enumerate/itemize ──────────────
        def _parse_items(items_block, label_style='alpha'):
            """
            Restituisce lista di (label, testo_pulito, opzioni, punti).
            label_style: 'alpha' → a) b) c)  |  'roman' → i) ii) iii)
            """
            risultati = []
            # match sia \item[X] con etichetta esplicita sia \item senza
            pat = _r.compile(
                r'\\item(?:\[([^\]]*)\])?\s*(.*?)(?=\\item(?:\[|\s)|$)',
                _r.DOTALL
            )
            auto_idx = 0
            for im in pat.finditer(items_block):
                explicit = im.group(1)
                if explicit is not None:
                    label = explicit.strip()
                    auto_idx += 1
                else:
                    # genera etichetta automatica
                    if label_style == 'roman':
                        label = ROMANI[min(auto_idx, len(ROMANI)-1)] + ')'
                    else:
                        label = LETTERE[auto_idx % 26] + ')'
                    auto_idx += 1

                raw_text = im.group(2).strip()
                if not label and not raw_text:
                    continue

                opzioni = []

                # sotto-enumerate (es. scelta multipla o domande i) ii) iii))
                inner = _r.search(
                    r'\\begin\{enumerate\}\s*\[([^\]]*)\]\s*(.*?)\\end\{enumerate\}',
                    raw_text, _r.DOTALL
                )
                if not inner:
                    inner = _r.search(
                        r'\\begin\{enumerate\}\s*(.*?)\\end\{enumerate\}',
                        raw_text, _r.DOTALL
                    )
                if inner:
                    inner_content = inner.group(inner.lastindex)
                    inner_label   = inner.group(1) if inner.lastindex > 1 else 'a)'
                    inner_style   = 'roman' if 'i' in inner_label else 'alpha'
                    for sub_label, sub_testo, _, _ in _parse_items(inner_content, inner_style):
                        opt_c = (f"{sub_label} {sub_testo}").strip()
                        if opt_c:
                            opzioni.append(opt_c)
                    raw_text = raw_text[:inner.start()].strip()

                # Vero/Falso
                vf_pairs = _r.findall(r'\$\\square\$\s*\\textbf\{([VF])\}', raw_text)
                if vf_pairs:
                    opzioni = [f"☐ {v}" for v in vf_pairs]
                    raw_text = _r.sub(
                        r'\$\\square\$\s*\\textbf\{[VF]\}\s*(?:\\quad)?', '', raw_text
                    ).strip()

                testo_clean = _clean_latex_line(raw_text)
                testo_clean = _r.sub(r'\n{2,}', '\n', testo_clean).strip()

                punti = _estrai_punti(label + ' ' + testo_clean)
                risultati.append((label, testo_clean, opzioni, punti))
            return risultati

        # ── scansiona tutti gli ambienti enumerate/itemize nel body ──────────
        env_pat = _r.compile(
            r'\\begin\{(enumerate|itemize)\}(\s*\[[^\]]*\])?\s*(.*?)\\end\{(?:enumerate|itemize)\}',
            _r.DOTALL
        )
        used_ranges = []
        for em in env_pat.finditer(body):
            used_ranges.append((em.start(), em.end()))
            opt_arg      = (em.group(2) or '').strip()  # es. [a)] o [i)] o vuoto
            items_block  = em.group(3)
            label_style  = 'roman' if opt_arg.startswith('[i') else 'alpha'

            # testo eventuale prima di questo ambiente (intestazione secondaria)
            preceding_start = used_ranges[-2][1] if len(used_ranges) >= 2 else first_env
            # (non necessario in docx, usiamo solo il testo_intro già estratto)

            for label, testo_clean, opzioni, punti in _parse_items(items_block, label_style):
                sottopunti.append({
                    'label':  label,
                    'testo':  testo_clean,
                    'opzioni': opzioni,
                    'punti':  punti,
                })

        # ── fallback: \item liberi fuori da qualsiasi ambiente ───────────────
        if not sottopunti:
            free_items = _r.compile(
                r'\\item(?:\[([^\]]*)\])?\s*(.*?)(?=\\item(?:\[|\s)|\\end\{|$)',
                _r.DOTALL
            )
            auto_idx = 0
            for im in free_items.finditer(body):
                skip = any(s <= im.start() < e for s, e in used_ranges)
                if skip:
                    continue
                explicit = im.group(1)
                label    = explicit.strip() if explicit is not None else (LETTERE[auto_idx % 26] + ')')
                auto_idx += 1
                raw_text = im.group(2).strip()
                testo_clean = _clean_latex_line(raw_text)
                testo_clean = _r.sub(r'\n{2,}', '\n', testo_clean).strip()
                sottopunti.append({
                    'label':  label,
                    'testo':  testo_clean,
                    'opzioni': [],
                    'punti':  _estrai_punti(label + ' ' + testo_clean),
                })

        if titolo_ex or sottopunti:
            data['esercizi'].append({
                'titolo':      titolo_ex,
                'testo_intro': testo_intro,
                'sottopunti':  sottopunti,
            })
    return data


def _estrai_punti(text):
    import re as _r
    m = _r.search(r'\((\d+(?:[.,]\d+)?)\s*pt\)', text)
    return m.group(1) if m else ''


def latex_to_docx_via_ai(codice_latex, con_griglia=True):
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
    except ImportError:
        return None, "python-docx non installato. Esegui: pip install python-docx"

    try:
        data = _parse_latex_to_data(codice_latex)
    except Exception as e:
        import traceback
        return None, f"Errore parsing LaTeX: {e}\n{traceback.format_exc()}"

    try:
        doc = DocxDocument()
        MARGIN_CM = 1.5
        PAGE_W_DXA = int((21.0 - 2 * MARGIN_CM) / 2.54 * 1440)
        GRIGLIA_W_DXA = PAGE_W_DXA - 80

        for section in doc.sections:
            section.page_width    = Cm(21.0)
            section.page_height   = Cm(29.7)
            section.left_margin   = Cm(MARGIN_CM); section.right_margin  = Cm(MARGIN_CM)
            section.top_margin    = Cm(MARGIN_CM); section.bottom_margin = Cm(MARGIN_CM)

        style = doc.styles['Normal']
        style.font.name = 'Arial'; style.font.size = Pt(11)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = p.add_run(data.get('titolo', 'Verifica'))
        rt.bold = True; rt.font.size = Pt(14)

        _cw = [int(PAGE_W_DXA * 0.62), PAGE_W_DXA - int(PAGE_W_DXA * 0.62)]
        hdr_tbl = doc.add_table(rows=1, cols=2)
        hdr_el = hdr_tbl._tbl
        tblPr_h = hdr_el.find(_qn('w:tblPr'))
        if tblPr_h is None:
            tblPr_h = _OE('w:tblPr'); hdr_el.insert(0, tblPr_h)
        for tag_h, attrs_h in [
            ('w:tblW',      {'w:w': str(PAGE_W_DXA), 'w:type': 'dxa'}),
            ('w:tblLayout', {'w:type': 'fixed'}),
        ]:
            el_h = _OE(tag_h)
            for k_h, v_h in attrs_h.items(): el_h.set(_qn(k_h), v_h)
            ex_h = tblPr_h.find(_qn(tag_h))
            if ex_h is not None: tblPr_h.remove(ex_h)
            tblPr_h.append(el_h)
        tblB_h = _OE('w:tblBorders')
        for side_h in ('top','left','bottom','right','insideH','insideV'):
            b_h = _OE(f'w:{side_h}'); b_h.set(_qn('w:val'), 'nil'); tblB_h.append(b_h)
        ex_b_h = tblPr_h.find(_qn('w:tblBorders'))
        if ex_b_h is not None: tblPr_h.remove(ex_b_h)
        tblPr_h.append(tblB_h)
        _fix_tbl_grid(hdr_el, _cw, _qn, _OE)
        hdr_cells_data = [
            [("Nome: ", "_______________________________")],
            [("Classe e Data: ", "______________________")],
        ]
        for ci_h, runs in enumerate(hdr_cells_data):
            cell_h = hdr_tbl.cell(0, ci_h)
            tc_h = cell_h._tc; tcPr_h = tc_h.get_or_add_tcPr()
            tcW_h = _OE('w:tcW'); tcW_h.set(_qn('w:w'), str(_cw[ci_h])); tcW_h.set(_qn('w:type'),'dxa')
            ex_w_h = tcPr_h.find(_qn('w:tcW'))
            if ex_w_h is not None: tcPr_h.remove(ex_w_h)
            tcPr_h.append(tcW_h)
            tcBdr_h = _OE('w:tcBorders')
            for s_h in ('top','left','bottom','right'):
                b2_h = _OE(f'w:{s_h}'); b2_h.set(_qn('w:val'),'nil'); tcBdr_h.append(b2_h)
            ex_bdr_h = tcPr_h.find(_qn('w:tcBorders'))
            if ex_bdr_h is not None: tcPr_h.remove(ex_bdr_h)
            tcPr_h.append(tcBdr_h)
            p_h = cell_h.paragraphs[0]
            for lbl_h, line_h in runs:
                r1_h = p_h.add_run(lbl_h); r1_h.bold = True; r1_h.font.size = Pt(10)
                r2_h = p_h.add_run(line_h); r2_h.font.size = Pt(10)

        nota = data.get('intestazione_nota', '')
        if nota:
            p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r3 = p3.add_run(nota); r3.italic = True; r3.font.size = Pt(10)
        doc.add_paragraph()

        for ex in data.get('esercizi', []):
            pe = doc.add_paragraph()
            pe.paragraph_format.space_before = Pt(10)
            pe.paragraph_format.space_after = Pt(3)
            rt = pe.add_run(ex.get('titolo', '')); rt.bold = True; rt.font.size = Pt(12)
            intro = ex.get('testo_intro', '').strip()
            if intro:
                pi = doc.add_paragraph(intro)
                pi.paragraph_format.space_before = Pt(0)
                pi.paragraph_format.space_after = Pt(3)
            for sp in ex.get('sottopunti', []):
                label = sp.get('label', '').strip()
                testo = sp.get('testo', '').strip()
                opzioni = sp.get('opzioni', [])
                ps = doc.add_paragraph()
                ps.paragraph_format.left_indent = Cm(0.4)
                ps.paragraph_format.space_before = Pt(2)
                ps.paragraph_format.space_after = Pt(2)
                rl = ps.add_run(label + "  "); rl.bold = True; rl.font.size = Pt(11)
                if testo:
                    ps.add_run(testo).font.size = Pt(11)
                if opzioni:
                    for opt in opzioni:
                        po = doc.add_paragraph()
                        po.paragraph_format.left_indent = Cm(1.2)
                        po.paragraph_format.space_before = Pt(0)
                        po.paragraph_format.space_after = Pt(1)
                        po.add_run(str(opt)).font.size = Pt(11)
                if not opzioni:
                    pr = doc.add_paragraph()
                    pr.paragraph_format.left_indent = Cm(0.4)
                    pr.paragraph_format.space_before = Pt(1)
                    pr.paragraph_format.space_after = Pt(8)

        esercizi_parsed = parse_esercizi(codice_latex) if con_griglia else []
        if con_griglia and esercizi_parsed:
            pg = doc.add_paragraph()
            rg = pg.add_run("Griglia di Valutazione"); rg.bold = True; rg.font.size = Pt(12)
            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER; pg.paragraph_format.space_before = Pt(12)

            n_sotto_totali = sum(len(ex['items']) for ex in esercizi_parsed)
            if n_sotto_totali > 12:
                row_es = ["Es."]
                row_max = ["Max"]
                for ex in esercizi_parsed:
                    row_es.append(f"Es. {ex['num']}")
                    tot_ex = 0
                    for _, pts in ex['items']:
                        try: tot_ex += float(pts.replace(',','.'))
                        except: pass
                    row_max.append(str(int(tot_ex)) if tot_ex == int(tot_ex) else str(round(tot_ex,1)))
                row_es.append("Tot")
                tot_pts = sum(float(x.replace(',','.')) for x in row_max[1:] if x.replace(',','.').replace('.','').isdigit())
                row_max.append(str(int(tot_pts)) if tot_pts == int(tot_pts) else str(round(tot_pts,1)))
                row_sotto_c = [""] * len(row_es)
                _build_griglia_xml(doc, row_es, row_sotto_c, row_max, GRIGLIA_W_DXA)
            else:
                row_es = ["Es."]; row_sotto = ["Sotto."]; row_max = ["Max"]
                for ex in esercizi_parsed:
                    for label_clean, pts in ex['items']:
                        lbl = label_clean.replace('*','').replace(')','').strip()[:3]
                        row_es.append(str(ex['num']))
                        row_sotto.append(lbl)
                        row_max.append(str(pts))
                row_es.append("Tot"); row_sotto.append("")
                tot_pts = 0
                for x in row_max[1:]:
                    try: tot_pts += float(x.replace(',','.'))
                    except: pass
                row_max.append(str(int(tot_pts)) if tot_pts == int(tot_pts) else str(round(tot_pts,1)))
                if len(row_es) >= 3:
                    _build_griglia_xml(doc, row_es, row_sotto, row_max, GRIGLIA_W_DXA)

        buf = io.BytesIO(); doc.save(buf)
        return buf.getvalue(), None

    except Exception as e:
        import traceback
        return None, f"Errore DOCX: {e}\n{traceback.format_exc()}"


def modifica_verifica_con_ai(latex_originale, richiesta_modifica, model):
    prompt = f"""Sei un docente esperto di LaTeX. Ti viene fornita una verifica già generata e una richiesta di modifica.

VERIFICA ORIGINALE:
{latex_originale}

RICHIESTA DI MODIFICA:
{richiesta_modifica}

ISTRUZIONI:
- Applica SOLO le modifiche richieste
- Mantieni la struttura LaTeX esistente
- NON rigenerare da zero, modifica l'esistente
- Mantieni lo stesso preambolo e intestazione
- Se la modifica riguarda punteggi, ricalcola la somma totale
- Restituisci il codice LaTeX completo modificato
- TERMINA con \\end{{document}}

OUTPUT: SOLO codice LaTeX completo, senza ```latex né spiegazioni."""

    response = model.generate_content(prompt)
    latex_modificato = response.text.replace("```latex", "").replace("```", "").strip()
    if "\\end{document}" not in latex_modificato:
        latex_modificato += "\n\\end{document}"
    return latex_modificato

def costruisci_prompt_esercizi(esercizi_custom, num_totale, punti_totali, mostra_punteggi):
    n_liberi = max(0, num_totale - len(esercizi_custom))
    righe = [
        f"\nSTRUTTURA ESERCIZI — REGOLA ASSOLUTA:",
        f"Devi generare ESATTAMENTE {num_totale} esercizi, né uno di più né uno di meno.",
        f"Ogni esercizio è un blocco \\subsection*{{Esercizio N: Titolo}}.",
        f"Conta i tuoi \\subsection* prima di terminare: devono essere ESATTAMENTE {num_totale}.",
    ]
    immagini = []

    if mostra_punteggi and num_totale > 0:
        pts_base = punti_totali // num_totale
        resto = punti_totali - pts_base * num_totale
        pts_per_ex = [pts_base] * num_totale
        if resto > 0:
            pts_per_ex[0] += resto
        righe.append(f"\nDISTRIBUZIONE PUNTI SUGGERITA (totale ESATTO: {punti_totali} pt):")
        for i_ex in range(num_totale):
            righe.append(f"  - Esercizio {i_ex+1}: circa {pts_per_ex[i_ex]} pt (distribuisci tra i sottopunti)")
        righe.append(f"REGOLA CRITICA: la somma di TUTTI i (X pt) nei sottopunti DEVE essere ESATTAMENTE {punti_totali} pt.")

    ha_primo_custom = len(esercizi_custom) > 0
    if not ha_primo_custom:
        righe.append(
            f"\nREGOLA PRIMO ESERCIZIO (Esercizio 1 — SEMPRE presente, NON modificabile):\n"
            f"Il primo esercizio DEVE chiamarsi 'Saperi Essenziali' e coprire i concetti fondamentali\n"
            f"dell'argomento che TUTTI gli studenti devono conoscere (definizioni, concetti base, formule\n"
            f"chiave, fatti imprescindibili). NON inserire mai il simbolo (*) in questo esercizio:\n"
            f"è obbligatorio per tutti, nessuna esclusione. Calibra il livello di difficoltà in modo\n"
            f"accessibile. Gli esercizi {2}–{num_totale} possono approfondire e variare."
        )

    righe.append(f"\nDETTAGLIO ESERCIZI ({num_totale} totali):")
    for i, ex in enumerate(esercizi_custom, 1):
        tipo, desc = ex.get('tipo', 'Aperto'), ex.get('descrizione', '').strip()
        riga = f"- Esercizio {i} [{tipo}]" + (f": {desc}" if desc else "")
        if ex.get('immagine'):
            riga += f" — vedi immagine allegata per l'esercizio {i}"
            immagini.append({'idx': i, 'data': ex['immagine'].getvalue(),
                             'mime_type': ex['immagine'].type})
        if tipo == "Scelta multipla":
            riga += " — opzioni con \\begin{enumerate}[a)] \\item prima \\item seconda \\end{enumerate}"
        elif tipo == "Vero/Falso":
            riga += " — $\\square$ \\textbf{V} $\\quad\\square$ \\textbf{F} \\quad testo (min 3)"
        elif tipo == "Completamento":
            riga += " — \\underline{\\hspace{3cm}} per gli spazi"
        righe.append(riga)
    if n_liberi > 0:
        start_idx = len(esercizi_custom) + 1
        end_idx   = num_totale
        righe.append(f"- Esercizi {start_idx}–{end_idx}: genera tu {n_liberi} esercizi coerenti con l'argomento.")
    return "\n".join(righe), immagini


# ── HELPER ───────────────────────────────────────────────────────────────────────
def _tempo_relativo(ts):
    if ts is None:
        return ""
    diff = time.time() - ts
    if diff < 60:
        return "pochi secondi fa"
    elif diff < 3600:
        m = int(diff // 60)
        return f"{m} min fa"
    elif diff < 86400:
        h = int(diff // 3600)
        return f"{h} ore fa"
    else:
        return time.strftime("%d/%m", time.localtime(ts))


def _stima_dimensione(data_bytes):
    if data_bytes is None:
        return "—"
    kb = len(data_bytes) / 1024
    if kb < 1024:
        return f"{kb:.0f} KB"
    return f"{kb/1024:.1f} MB"


# ── SESSION STATE ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title=APP_NAME,
    page_icon=APP_ICON,
    layout="wide",
    initial_sidebar_state="collapsed"
)


def _vf():
    return {'latex': '', 'pdf': None, 'preview': False,
            'soluzioni_latex': '', 'soluzioni_pdf': None, 'docx': None,
            'pdf_ts': None, 'docx_ts': None, 'latex_originale': ''}

if 'utente' not in st.session_state: st.session_state.utente = None
if 'verifiche' not in st.session_state: st.session_state.verifiche = {'A': _vf(), 'B': _vf(), 'R': _vf(), 'RB': _vf(), 'S': {'latex': None, 'testo': None}}
if 'esercizi_custom' not in st.session_state: st.session_state.esercizi_custom = []
if 'last_materia'    not in st.session_state: st.session_state.last_materia = None
if 'last_argomento'  not in st.session_state: st.session_state.last_argomento = None
if 'last_gen_ts'     not in st.session_state: st.session_state.last_gen_ts = None
if '_storico_refresh' not in st.session_state: st.session_state._storico_refresh = 0
if '_preferiti' not in st.session_state: st.session_state._preferiti = set()
if '_storico_page' not in st.session_state: st.session_state._storico_page = 1
if '_onboarding_done' not in st.session_state: st.session_state._onboarding_done = False

# ── CALCOLA VERIFICHE DEL MESE (una volta per rerun) ────────────────────────────
ADMIN_EMAILS = {"giacomosaragoni96@gmail.com"}  # ← email admin con verifiche illimitate

_verifiche_mese_count = _get_verifiche_mese(st.session_state.utente.id) if st.session_state.utente else 0
_is_admin = (st.session_state.utente.email in ADMIN_EMAILS) if st.session_state.utente else False
_limite_raggiunto = (not _is_admin) and (_verifiche_mese_count >= LIMITE_MENSILE)

# ── CSS GLOBALE ──────────────────────────────────────────────────────────────────
is_dark = (st.session_state.theme == "dark")

_SB_LABEL   = "#c8c6bc"
_SB_MUTED   = "#8a8880"
_SB_BORDER  = "#2a2926"
_SB_TEXT    = "#e8e6e0"

st.markdown(f"""
<style>
  @import url('https://fonts.googleapis.com/css2?family=DM+Sans:ital,wght@0,300;0,400;0,500;0,600;0,700;0,800;0,900;1,400&display=swap');

  *, *::before, *::after {{ box-sizing: border-box; }}

  .stApp {{
    background-color: {T['bg']} !important;
    font-family: 'DM Sans', sans-serif;
    color: {T['text']};
    transition: background-color 0.25s ease, color 0.25s ease;
  }}

  .block-container {{
    padding: 5rem 1.5rem 4rem !important;
    max-width: 780px !important;
    margin: 0 auto !important;
  }}

  #MainMenu, footer {{ visibility: hidden; }}
  .stDecoration {{ display: none; }}

  header[data-testid="stHeader"] {{
    background-color: {T['bg']} !important;
    border-bottom: 1px solid {T['border']} !important;
  }}

  header button svg {{
    fill: {T['text']} !important;
    color: {T['text']} !important;
    stroke: {T['text']} !important;
  }}
  header button {{
    background: {T['card']} !important;
    border: 1.5px solid {T['border2']} !important;
    border-radius: 8px !important;
    color: {T['text']} !important;
  }}
  header button:hover {{
    background: {T['hover']} !important;
    border-color: {T['accent']} !important;
  }}
  header button:hover svg {{
    fill: {T['accent']} !important;
  }}
  .stMarkdown h1 a, .stMarkdown h2 a, .stMarkdown h3 a {{ display: none !important; }}

  /* ════ SIDEBAR ════ */
  [data-testid="stSidebar"] {{
    background: #141412 !important;
    border-right: 1px solid {_SB_BORDER} !important;
  }}
  .sidebar-title {{
    font-family: 'DM Sans', sans-serif;
    font-size: 1.1rem !important;
    font-weight: 800 !important;
    letter-spacing: -0.01em;
    color: #f5f3ed !important;
    margin: 0.5rem 0 1.2rem 0;
    padding-bottom: 0.6rem;
    border-bottom: 1px solid {_SB_BORDER};
  }}
  [data-testid="stSidebar"] .block-container {{
    padding: 1.5rem 1.2rem !important;
    max-width: 100% !important;
  }}
  [data-testid="stSidebar"] p,
  [data-testid="stSidebar"] span,
  [data-testid="stSidebar"] label,
  [data-testid="stSidebar"] div {{
    color: {_SB_TEXT} !important;
  }}
  [data-testid="stSidebar"] .stTextInput label p,
  [data-testid="stSidebar"] .stSelectbox label p,
  [data-testid="stSidebar"] .stNumberInput label p {{
    color: {_SB_MUTED} !important;
    font-size: 0.7rem !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    font-weight: 700 !important;
  }}
  [data-testid="stSidebar"] .stCheckbox label {{
    color: {_SB_TEXT} !important;
    font-size: 0.9rem !important;
  }}
  [data-testid="stSidebar"] .stCheckbox [data-testid="stCheckbox"] span:first-child {{
    background-color: {T['input_bg']} !important;
    border: 1.5px solid {T['border2']} !important;
    border-radius: 5px !important;
  }}
  [data-testid="stSidebar"] .stTextInput input,
  [data-testid="stSidebar"] .stNumberInput input {{
    background: #232320 !important;
    border: 1.5px solid #3d3c36 !important;
    border-radius: 8px !important;
    color: #f5f3ed !important;
  }}
  [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] > div:first-child {{
    background: #232320 !important;
    border: 1.5px solid #3d3c36 !important;
    border-radius: 8px !important;
  }}
  [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] span {{
    color: #f5f3ed !important;
  }}
  [data-testid="stSidebar"] .stRadio label {{
    color: {_SB_TEXT} !important;
  }}
  [data-testid="stSidebar"] .stRadio [data-testid="stMarkdownContainer"] p {{
    color: {_SB_TEXT} !important;
  }}
  [data-testid="stSidebar"] .stButton button {{
    background: #232320 !important;
    color: #f5f3ed !important;
    border: 1.5px solid #3d3c36 !important;
    border-radius: 8px !important;
  }}
  [data-testid="stSidebar"] .stButton button:hover {{
    background: #2e2d28 !important;
    border-color: #5a5950 !important;
  }}
  [data-testid="stSidebar"] .stSelectSlider [data-testid="stMarkdownContainer"] p {{
    color: {_SB_TEXT} !important;
  }}
  [data-testid="stSidebar"] .section-label {{
    color: #5a5950 !important;
    border-bottom-color: {_SB_BORDER} !important;
  }}
  [data-testid="collapsedControl"] {{
    top: 0.75rem !important;
    left: 0.75rem !important;
    z-index: 999 !important;
  }}
  [data-testid="collapsedControl"] button {{
    background: {T['card']} !important;
    border: 2px solid {T['accent']} !important;
    border-radius: 10px !important;
    color: {T['accent']} !important;
    width: 40px !important;
    height: 40px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    box-shadow: 0 2px 12px {T['accent']}33 !important;
    transition: transform 0.15s ease, box-shadow 0.15s ease, background 0.15s ease !important;
    padding: 0 !important;
  }}
  [data-testid="collapsedControl"] button:hover {{
    background: {T['accent']} !important;
    box-shadow: 0 4px 18px {T['accent']}55 !important;
    transform: scale(1.08) !important;
  }}
  [data-testid="collapsedControl"] button:hover svg {{
    fill: #ffffff !important;
    color: #ffffff !important;
    stroke: #ffffff !important;
  }}
  [data-testid="collapsedControl"] button svg {{
    fill: {T['accent']} !important;
    color: {T['accent']} !important;
    stroke: {T['accent']} !important;
    width: 18px !important;
    height: 18px !important;
  }}

  [data-testid="stSidebar"] .sidebar-label,
  .sidebar-label {{
    font-size: 0.72rem !important;
    font-weight: 800 !important;
    letter-spacing: 0.1em !important;
    text-transform: uppercase !important;
    color: #D97706 !important;
    margin: 1rem 0 0.5rem 0 !important;
    padding-bottom: 0.35rem !important;
    border-bottom: 2px solid #3a3320 !important;
    display: block !important;
  }}

  [data-testid="stSidebar"] .logout-btn-wrap div.stButton > button,
  [data-testid="stSidebar"] .logout-btn-wrap .stButton button,
  [data-testid="stSidebar"] .logout-btn-wrap button {{
    background: transparent !important;
    color: #f87171 !important;
    border: 1px solid #5c2222 !important;
    border-radius: 8px !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    padding: 6px 14px !important;
    width: auto !important;
    min-height: unset !important;
    box-shadow: none !important;
    transition: background 0.15s ease, border-color 0.15s ease, color 0.15s ease !important;
    letter-spacing: 0.02em !important;
  }}
  [data-testid="stSidebar"] .logout-btn-wrap div.stButton > button:hover,
  [data-testid="stSidebar"] .logout-btn-wrap .stButton button:hover,
  [data-testid="stSidebar"] .logout-btn-wrap button:hover {{
    background: #2a0f0f !important;
    border-color: #f87171 !important;
    color: #fca5a5 !important;
    box-shadow: 0 0 0 1px #5c2222 !important;
  }}

  h1, h2, h3 {{
    font-family: 'DM Sans', sans-serif !important;
    color: {T['text']} !important;
    letter-spacing: -0.02em;
  }}

  .hero-wrap {{
    margin-bottom: 2.5rem;
    padding-bottom: 1.8rem;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-direction: column;
    flex-wrap: wrap;
    gap: 0;
    position: relative;
  }}
  .hero-left {{ flex: 1; min-width: 200px; text-align: center; }}
  @keyframes iconBounce {{
    0%   {{ transform: rotate(0deg) scale(1); }}
    15%  {{ transform: rotate(-12deg) scale(1.15); }}
    30%  {{ transform: rotate(8deg) scale(1.1); }}
    45%  {{ transform: rotate(-6deg) scale(1.05); }}
    60%  {{ transform: rotate(3deg) scale(1.02); }}
    75%  {{ transform: rotate(-1deg) scale(1.01); }}
    100% {{ transform: rotate(0deg) scale(1); }}
  }}
  @keyframes badgePop {{
    0%   {{ opacity: 0; transform: translateY(6px); }}
    100% {{ opacity: 1; transform: translateY(0); }}
  }}
  @keyframes subFadeIn {{
    0%   {{ opacity: 0; transform: translateY(4px); }}
    100% {{ opacity: 1; transform: translateY(0); }}
  }}
  .hero-title {{
    font-family: 'DM Sans', sans-serif;
    font-size: 96px !important;
    font-weight: 900 !important;
    color: {T['text']};
    line-height: 1.0;
    margin: 0 0 0.15rem 0;
    letter-spacing: -0.04em;
    display: inline-flex;
    align-items: center;
    gap: 0;
    justify-content: center;
  }}
  .hero-icon {{
    display: inline-block;
    margin-right: 0.3em;
    animation: iconBounce 1.1s ease 0.2s both;
    transform-origin: center bottom;
  }}
  .hero-ai {{
    background: linear-gradient(135deg, {T['accent']} 0%, #FF8C00 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    animation: badgePop 0.6s ease 0.5s both;
  }}
  .hero-sub {{
    font-size: 1.05rem;
    color: {T['text2']};
    margin: 0 0 0.55rem 0;
    font-weight: 500;
    letter-spacing: -0.01em;
    animation: subFadeIn 0.5s ease 0.35s both;
    opacity: 0;
  }}
  .hero-beta {{
    display: inline-block;
    font-size: 0.65rem;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: {T['muted']};
    background: {T['card2']};
    border: 1px solid {T['border2']};
    border-radius: 100px;
    padding: 2px 10px;
    font-family: 'DM Sans', sans-serif;
    animation: badgePop 0.5s ease 0.75s both;
    opacity: 0;
  }}

  .stTextInput label p,
  .stSelectbox label p,
  .stNumberInput label p,
  .stTextArea label p,
  .stFileUploader label p {{
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    color: {T['text2']} !important;
    letter-spacing: 0.01em;
    text-transform: uppercase;
    margin-bottom: 4px !important;
  }}

  .stTextInput input,
  .stNumberInput input {{
    background: {T['input_bg']} !important;
    border: 1.5px solid {T['border']} !important;
    border-radius: 10px !important;
    color: {T['text']} !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.95rem !important;
    padding: 14px 16px !important;
    min-height: 52px !important;
    height: 52px !important;
    box-sizing: border-box !important;
    transition: border-color 0.15s ease, box-shadow 0.15s ease;
  }}
  .stTextInput input::placeholder,
  .stNumberInput input::placeholder {{
    color: {T['muted']} !important;
    opacity: 1 !important;
  }}
  .stTextInput input:focus,
  .stNumberInput input:focus {{
    border-color: {T['accent']} !important;
    box-shadow: 0 0 0 3px {T['accent_light']} !important;
    outline: none !important;
  }}
  .stTextArea textarea {{
    background: {T['input_bg']} !important;
    border: 1.5px solid {T['border']} !important;
    border-radius: 10px !important;
    color: {T['text']} !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.9rem !important;
    padding: 10px 14px !important;
  }}
  .stTextArea textarea::placeholder {{
    color: {T['muted']} !important;
    opacity: 1 !important;
  }}
  .stTextArea textarea:focus {{
    border-color: {T['accent']} !important;
    box-shadow: 0 0 0 3px {T['accent_light']} !important;
  }}

  .stSelectbox [data-baseweb="select"] > div:first-child {{
    background: {T['input_bg']} !important;
    border: 1.5px solid {T['border']} !important;
    border-radius: 10px !important;
    color: {T['text']} !important;
  }}
  .stSelectbox [data-baseweb="select"] span {{
    color: {T['text']} !important;
  }}

  .stCheckbox label {{
    color: {T['text']} !important;
    font-size: 0.9rem !important;
    font-family: 'DM Sans', sans-serif !important;
  }}
  .stCheckbox [data-testid="stCheckbox"] span:first-child {{
    background-color: {T['input_bg']} !important;
    border: 1.5px solid {T['border2']} !important;
    border-radius: 5px !important;
  }}

  div.stButton > button[kind="primary"] {{
    background: #D97706 !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    transition: transform 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275), box-shadow 0.2s ease, filter 0.2s ease !important;
    box-shadow: 0 2px 12px rgba(217,119,6,0.35) !important;
    display: block !important;
  }}
  div.stButton > button[kind="primary"]:hover {{
    transform: scale(1.05) !important;
    box-shadow: 0 10px 25px rgba(217,119,6,0.5) !important;
    filter: brightness(1.1) !important;
    border: none !important;
  }}
  div.stButton > button[kind="primary"]:active {{
    transform: scale(0.98) !important;
  }}
  div.stButton > button[kind="primary"]:disabled {{
    background: #9CA3AF !important;
    box-shadow: none !important;
    transform: none !important;
    filter: none !important;
    cursor: not-allowed !important;
    opacity: 0.7 !important;
  }}

  .dl-card {{
    background: #FFFFFF !important;
    padding: 1.2rem;
    border-radius: 15px;
    border: 1px solid #E0E0E0;
    text-align: center;
    margin-bottom: 1rem;
  }}
  .dl-label {{
    font-size: 0.85rem;
    color: #666;
    margin-bottom: 0.8rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.5px;
  }}
  .hint-docx {{
    font-size: 0.78rem;
    color: #888;
    line-height: 1.3;
    margin-top: 12px;
    font-style: italic;
    text-align: left;
    border-top: 1px solid #EEE;
    padding-top: 8px;
  }}

  .output-card {{
    background: {T['card']} !important;
    border: 1.5px solid {T['border2']} !important;
    border-radius: 16px !important;
    padding: 0 !important;
    margin-bottom: 1.5rem !important;
    overflow: hidden !important;
    box-shadow: {T['shadow_md']} !important;
  }}
  .output-card > div {{
    padding: 1.2rem !important;
  }}
  .stDownloadButton button,
  [data-testid="stDownloadButton"] button,
  .stButton [data-testid="baseButton-secondary"],
  .stButton button[kind="secondary"],
  button[data-testid="baseButton-secondary"] {{
    background: {T['card']} !important;
    color: {T['text']} !important;
    border: 1.5px solid {T['border2']} !important;
    border-radius: 12px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 1.05rem !important;
    font-weight: 700 !important;
    padding: 1rem 1.4rem !important;
    transition: border-color 0.18s ease, box-shadow 0.18s ease, transform 0.18s ease !important;
    letter-spacing: 0.01em !important;
    box-shadow: 0 2px 6px rgba(0,0,0,0.04) !important;
    width: 100% !important;
  }}
  .stDownloadButton button:hover,
  [data-testid="stDownloadButton"] button:hover,
  .stButton [data-testid="baseButton-secondary"]:hover,
  button[data-testid="baseButton-secondary"]:hover {{
    background: {T['hover']} !important;
    border-color: {T['accent']} !important;
    transform: translateY(-3px) !important;
    box-shadow: 0 6px 20px rgba(217,119,6,0.18) !important;
    color: {T['text']} !important;
  }}

  .action-card {{
    background: {T['card']} !important;
    border: 2px solid {T['border']} !important;
    border-radius: 14px !important;
    padding: 1.2rem !important;
    margin-bottom: 1rem !important;
    transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1) !important;
    position: relative !important;
    overflow: hidden !important;
  }}
  .action-card::before {{
    content: '';
    position: absolute;
    top: 0;
    left: 0;
    width: 4px;
    height: 100%;
    background: linear-gradient(180deg, {T['accent']} 0%, {T['accent']}00 100%);
    opacity: 0;
    transition: opacity 0.25s ease;
  }}
  .action-card:hover {{
    border-color: {T['accent']}88 !important;
    box-shadow: 0 8px 24px {T['accent']}15 !important;
    transform: translateY(-2px) !important;
  }}
  .action-card:hover::before {{
    opacity: 1;
  }}
  .action-card .stDownloadButton button,
  .action-card .stButton button {{
    background: linear-gradient(135deg, {T['accent']} 0%, {T['accent']}ee 100%) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 12px !important;
    font-size: 1rem !important;
    font-weight: 700 !important;
    padding: 1rem 1.5rem !important;
    box-shadow: 0 4px 14px {T['accent']}35 !important;
    transition: all 0.2s cubic-bezier(0.175, 0.885, 0.32, 1.275) !important;
  }}
  .action-card .stDownloadButton button:hover,
  .action-card .stButton button:hover {{
    transform: scale(1.03) !important;
    box-shadow: 0 6px 20px {T['accent']}50 !important;
    filter: brightness(1.08) !important;
  }}
  .action-card .stDownloadButton button:active,
  .action-card .stButton button:active {{
    transform: scale(0.98) !important;
  }}
  .action-card [data-testid="stExpander"] {{
    background: transparent !important;
    border: 2px dashed {T['border2']} !important;
    border-radius: 12px !important;
    transition: all 0.2s ease !important;
  }}
  .action-card [data-testid="stExpander"]:hover {{
    border-color: {T['accent']} !important;
    background: {T['accent_light']} !important;
  }}
  .action-card [data-testid="stExpander"] summary {{
    background: transparent !important;
    color: {T['text']} !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
    padding: 1rem 1.2rem !important;
  }}
  .action-card [data-testid="stExpander"] summary:hover {{
    color: {T['accent']} !important;
  }}
  .action-card [data-testid="stExpander"] summary svg {{
    color: {T['accent']} !important;
  }}
  .tex-download {{
    opacity: 0.7 !important;
    transition: opacity 0.2s ease, transform 0.2s ease !important;
  }}
  .tex-download:hover {{
    opacity: 1 !important;
    transform: translateX(3px) !important;
  }}
  .hint-docx {{
    background: {T['accent_light']} !important;
    border-left: 3px solid {T['accent']} !important;
    border-radius: 8px !important;
    padding: 10px 14px !important;
    margin-top: 8px !important;
    font-size: 0.8rem !important;
    color: {T['text2']} !important;
    line-height: 1.5 !important;
  }}
  .hint-docx strong {{
    color: {T['accent']} !important;
  }}
  .section-action-label {{
    display: inline-flex;
    align-items: center;
    gap: 8px;
    font-size: 0.7rem;
    font-weight: 800;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: {T['muted']};
    margin-bottom: 0.8rem;
    padding-bottom: 0.5rem;
    border-bottom: 2px solid {T['border']};
  }}
  .section-action-label::before {{
    content: '';
    width: 8px;
    height: 8px;
    border-radius: 50%;
    background: {T['accent']};
    display: inline-block;
  }}

  .compact-uploader [data-testid="stFileUploader"] section {{
    padding: 0 !important;
    border: none !important;
    background: transparent !important;
    min-height: unset !important;
  }}
  .compact-uploader [data-testid="stFileUploadDropzone"] {{
    display: none !important;
  }}
  .compact-uploader [data-testid="stFileUploader"] button {{
    background: {T['card2']} !important;
    color: {T['text2']} !important;
    border: 1px solid {T['border2']} !important;
    border-radius: 8px !important;
    font-size: 0.8rem !important;
    padding: 6px 14px !important;
  }}
  .compact-uploader [data-testid="stFileUploader"] button:hover {{
    border-color: {T['accent']} !important;
    color: {T['accent']} !important;
  }}

  .tex-btn-wrap .stDownloadButton button,
  .tex-btn-wrap [data-testid="stDownloadButton"] button {{
    background: transparent !important;
    color: {T['muted']} !important;
    border: 1px solid {T['border']} !important;
    border-radius: 6px !important;
    font-size: 0.72rem !important;
    font-weight: 500 !important;
    padding: 0.3rem 0.7rem !important;
    width: 100% !important;
    box-shadow: none !important;
    transform: none !important;
  }}
  .tex-btn-wrap .stDownloadButton button:hover,
  .tex-btn-wrap [data-testid="stDownloadButton"] button:hover {{
    color: {T['text2']} !important;
    border-color: #5a5950 !important;
    background: {T['card2']} !important;
    box-shadow: none !important;
    transform: translateY(-1px) !important;
  }}

  [data-testid="stExpander"] {{
    background: {T['card']} !important;
    border: 1px solid {T['border']} !important;
    border-radius: 12px !important;
    margin-bottom: 0.75rem !important;
    overflow: hidden;
  }}
  [data-testid="stExpander"] summary {{
    padding: 0.85rem 1.1rem !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    color: {T['text']} !important;
    background: {T['card']} !important;
  }}
  [data-testid="stExpander"] summary:hover {{
    background: {T['hover']} !important;
  }}
  [data-testid="stExpander"] > div > div {{
    padding: 0.5rem 1.1rem 1rem !important;
  }}

  hr {{
    border: none;
    border-top: 1px solid {T['border']} !important;
    margin: 2rem 0 !important;
  }}

  [data-testid="stFileUploader"] section {{
    background: {T['card2']} !important;
    border: 1.5px dashed {T['border2']} !important;
    border-radius: 10px !important;
    padding: 0.75rem 1rem !important;
    min-height: unset !important;
  }}
  [data-testid="stFileUploader"] section > div {{ gap: 0.3rem !important; }}
  [data-testid="stFileUploadDropzone"] p,
  [data-testid="stFileUploadDropzone"] span,
  [data-testid="stFileUploadDropzone"] small,
  [data-testid="stFileUploader"] span,
  [data-testid="stFileUploader"] small,
  [data-testid="stFileUploader"] p {{
    color: {T['text2']} !important;
    opacity: 1 !important;
  }}
  [data-testid="stFileUploadDropzone"] button,
  [data-testid="stFileUploader"] button {{
    color: {T['accent']} !important;
    font-weight: 600 !important;
  }}
  [data-testid="stFileUploadDropzone"] svg {{
    fill: {T['muted']} !important;
    color: {T['muted']} !important;
  }}

  .chip {{
    display: inline-flex;
    align-items: center;
    gap: 4px;
    background: {T['accent_light']};
    color: {T['accent']};
    border: 1px solid {T['accent']};
    border-radius: 100px;
    padding: 2px 10px;
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 0.04em;
    text-transform: uppercase;
  }}

  .expander-heading {{
    font-size: 0.75rem !important;
    font-weight: 700 !important;
    letter-spacing: 0.04em !important;
    text-transform: uppercase !important;
    color: {T['text2']} !important;
    margin: 1rem 0 0.4rem 0 !important;
    padding: 4px 10px !important;
    background: {T['card2']} !important;
    border-left: 3px solid {T['accent']} !important;
    border-radius: 0 6px 6px 0 !important;
    display: block !important;
  }}

  .step-label {{
    display: flex;
    align-items: center;
    gap: 12px;
    margin: 2.2rem 0 0.75rem 0;
    font-family: 'DM Sans', sans-serif;
  }}
  .step-num {{
    width: 26px;
    height: 26px;
    border-radius: 50%;
    background: {T['accent']};
    color: #ffffff;
    font-size: 0.68rem;
    font-weight: 900;
    display: inline-flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
    letter-spacing: 0;
    box-shadow: 0 2px 8px {T['accent']}55;
  }}
  .step-title {{
    font-size: 0.82rem;
    font-weight: 800;
    letter-spacing: 0.07em;
    text-transform: uppercase;
    color: {T['text']};
  }}
  .step-line {{
    flex: 1;
    height: 1px;
    background: linear-gradient(90deg, {T['border2']} 0%, transparent 100%);
  }}

  .ai-hint {{
    display: flex;
    align-items: center;
    gap: 9px;
    background: {T['accent_light']};
    border: 1px solid {T['accent']}55;
    border-radius: 10px;
    padding: 9px 14px;
    font-size: 0.78rem;
    color: {T['text2']};
    margin: 1.8rem 0 0.6rem 0;
    font-family: 'DM Sans', sans-serif;
    line-height: 1.4;
  }}
  .ai-hint-icon {{ font-size: 1rem; flex-shrink: 0; }}
  .ai-hint strong {{ color: {T['accent']}; font-weight: 700; }}

  .personalizza-wrap [data-testid="stExpander"] {{
    border: 1.5px solid {T['accent']}44 !important;
    border-radius: 14px !important;
    background: {T['card']} !important;
    box-shadow: 0 2px 16px {T['accent']}0f !important;
  }}
  .personalizza-wrap [data-testid="stExpander"] summary {{
    background: {T['accent_light']} !important;
    color: {T['text']} !important;
    font-weight: 700 !important;
    font-size: 0.92rem !important;
    padding: 1rem 1.2rem !important;
  }}
  .personalizza-wrap [data-testid="stExpander"] summary:hover {{
    background: {T['accent_light']} !important;
    filter: brightness(0.96);
  }}
  .personalizza-wrap [data-testid="stExpander"] summary svg {{
    color: {T['accent']} !important;
    fill: {T['accent']} !important;
  }}
  .personalizza-wrap [data-testid="stExpander"] > div > div {{
    padding: 1rem 1.2rem 1.2rem !important;
  }}

  .genera-section {{ margin-top: 2.2rem; margin-bottom: 0.5rem; }}

  /* ── STICKY CTA su mobile ── */
  @media (max-width: 640px) {{
    .genera-section {{
      position: sticky;
      bottom: 0;
      z-index: 999;
      background: {T['bg']};
      padding: 0.8rem 0 0.5rem 0;
      margin: 0;
      border-top: 1px solid {T['border']};
      box-shadow: 0 -4px 20px rgba(0,0,0,0.15);
    }}
  }}

  /* ── ONBOARDING BANNER ── */
  .onboarding-banner {{
    display: flex;
    align-items: flex-start;
    gap: 14px;
    background: linear-gradient(135deg, {T['accent_light']} 0%, {T['card']} 100%);
    border: 1.5px solid {T['accent']};
    border-radius: 14px;
    padding: 1rem 1.3rem;
    margin-bottom: 1.5rem;
    font-family: 'DM Sans', sans-serif;
  }}
  .onboarding-steps {{
    display: flex;
    align-items: center;
    gap: 0.5rem;
    flex-wrap: wrap;
    margin-top: 0.4rem;
  }}
  .onboarding-step {{
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: {T['bg2']};
    border: 1px solid {T['border']};
    border-radius: 20px;
    padding: 5px 12px;
    font-size: 0.76rem;
    font-weight: 600;
    color: {T['text2']};
    white-space: nowrap;
  }}
  .onboarding-step-ico {{
    font-size: 0.85rem;
  }}
  .onboarding-step-arrow {{
    color: {T['accent']};
    font-size: 0.9rem;
    font-weight: 700;
    flex-shrink: 0;
  }}
  /* X close button per onboarding */
  div[data-testid="column"]:has(button[kind="secondary"][title="Chiudi"]) button,
  div[data-testid="column"] button[data-testid*="_dismiss_onboarding"] {{
    background: transparent !important;
    border: 1px solid {T['border']} !important;
    border-radius: 50% !important;
    color: {T['muted']} !important;
    font-size: 0.75rem !important;
    font-weight: 700 !important;
    width: 28px !important;
    height: 28px !important;
    min-height: unset !important;
    padding: 0 !important;
    box-shadow: none !important;
    line-height: 1 !important;
  }}
  div[data-testid="column"]:has(button[kind="secondary"][title="Chiudi"]) button:hover {{
    background: {T['hover']} !important;
    border-color: {T['muted']} !important;
    color: {T['text']} !important;
  }}
  .genera-hint {{
    text-align: center;
    font-size: 0.73rem;
    color: {T['muted']};
    margin-top: 0.6rem;
    font-family: 'DM Sans', sans-serif;
    letter-spacing: 0.02em;
    display: flex;
    align-items: center;
    justify-content: center;
    gap: 5px;
  }}

  .section-label {{
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    color: {T['muted']};
    margin: 1.2rem 0 0.5rem 0;
    padding-bottom: 0.3rem;
    border-bottom: 1px solid {T['border']};
  }}

  .hint {{
    font-size: 0.78rem;
    color: {T['muted']};
    margin-top: 4px;
    line-height: 1.4;
  }}

  .stSelectSlider [data-testid="stMarkdownContainer"] p {{
    color: {T['text']} !important;
    font-size: 0.88rem !important;
  }}
  .stSlider [data-baseweb="slider"] [role="slider"] {{
    background: {T['accent']} !important;
    border-color: {T['accent']} !important;
  }}
  .stSlider [data-baseweb="slider"] div[data-testid="stTickBar"] {{
    color: {T['muted']} !important;
  }}
  .stSlider label p {{
    color: {T['text2']} !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    letter-spacing: 0.01em;
    text-transform: uppercase;
  }}
  .stSlider [data-baseweb="slider"] [data-testid="stSliderThumbValue"] {{
    color: {T['text']} !important;
    font-weight: 700 !important;
    font-size: 0.95rem !important;
  }}

  .stAlert {{
    border-radius: 10px !important;
    border: none !important;
  }}
  .stAlert p,
  .stAlert span,
  .stAlert div,
  [data-testid="stAlert"] p,
  [data-testid="stAlert"] span,
  [data-testid="stAlert"] div {{
      color: #1a1915 !important;
      opacity: 1 !important;
  }}

  .stProgress > div > div {{
    background: {T['accent']} !important;
    border-radius: 100px !important;
  }}
  .stProgress > div {{
    background: {T['border']} !important;
    border-radius: 100px !important;
    height: 6px !important;
  }}

  .stSubheader {{
    font-family: 'DM Sans', sans-serif !important;
    font-size: 1.15rem !important;
    font-weight: 700 !important;
    color: {T['text']} !important;
  }}

  .status-bar {{
    display: flex;
    align-items: center;
    gap: 8px;
    padding: 6px 14px;
    background: {T['card2']};
    border: 1px solid {T['border']};
    border-radius: 8px;
    font-size: 0.78rem;
    color: {T['muted']};
    margin-top: 0.5rem;
    font-family: 'DM Sans', sans-serif;
  }}
  .status-bar .dot {{
    width: 6px;
    height: 6px;
    border-radius: 50%;
    background: {T['success']};
    flex-shrink: 0;
  }}
  .status-bar strong {{
    color: {T['text2']};
    font-weight: 600;
  }}

  .pdf-preview-wrap {{
    margin-top: 1rem;
    border: 1px solid {T['border']};
    border-radius: 14px;
    overflow: hidden;
    box-shadow: {T['shadow_md']};
    background: {T['card2']};
  }}

  .top-bar {{
    display: flex;
    align-items: center;
    justify-content: flex-start;
    gap: 0.75rem;
    margin-bottom: 1.2rem;
  }}
  .top-bar-hint {{
    display: none;
  }}
  @media (max-width: 640px) {{
    .top-bar-hint {{
      display: inline-flex;
      align-items: center;
      gap: 5px;
      background: {T['accent_light']};
      border: 1px solid {T['accent']};
      border-radius: 20px;
      padding: 5px 12px;
      font-size: 0.72rem;
      color: {T['accent']};
      font-weight: 600;
      white-space: nowrap;
    }}
  }}

  .fab-link {{
    position: fixed;
    top: 4.5rem;
    right: 1.5rem;
    z-index: 9999;
    display: inline-flex;
    align-items: center;
    gap: 8px;
    background: {T['accent']};
    color: #ffffff !important;
    text-decoration: none !important;
    border-radius: 50px;
    padding: 10px 18px;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.84rem;
    font-weight: 700;
    box-shadow: 0 4px 18px rgba(217,119,6,0.40);
    transition: transform 0.15s ease, filter 0.15s ease;
    white-space: nowrap;
  }}
  .fab-link:hover {{
    transform: translateY(-2px);
    filter: brightness(1.1);
    color: #ffffff !important;
  }}
  @media (max-width: 640px) {{
    .fab-link {{
      top: 4rem;
      right: 0.8rem;
      padding: 8px 14px;
      font-size: 0.78rem;
    }}
  }}

  .disclaimer {{
    display: flex;
    align-items: flex-start;
    gap: 8px;
    padding: 8px 12px;
    background: {T['card2']};
    border: 1px solid {T['border']};
    border-left: 3px solid {T['muted']};
    border-radius: 8px;
    font-size: 0.74rem;
    color: {T['muted']};
    font-family: 'DM Sans', sans-serif;
    line-height: 1.45;
    margin-bottom: 1rem;
  }}
  .disclaimer-icon {{ flex-shrink: 0; font-size: 0.9rem; margin-top: 1px; }}

  .app-footer {{
    text-align: center;
    font-size: 0.72rem;
    color: {T['muted']};
    font-family: 'DM Sans', sans-serif;
    margin-top: 3rem;
    padding-top: 1.2rem;
    border-top: 1px solid {T['border']};
    line-height: 1.6;
  }}

  .user-pill {{
    display: flex;
    align-items: center;
    gap: 10px;
    background: #1e1d1b;
    border: 1px solid #2e2d28;
    border-radius: 12px;
    padding: 10px 14px;
    margin-top: 0.5rem;
  }}
  .user-avatar {{
    width: 32px;
    height: 32px;
    border-radius: 50%;
    background: linear-gradient(135deg, #D97706, #FF8C00);
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 0.85rem;
    font-weight: 800;
    color: white;
    flex-shrink: 0;
    font-family: 'DM Sans', sans-serif;
  }}
  .user-info {{
    flex: 1;
    min-width: 0;
  }}
  .user-email {{
    font-size: 0.78rem;
    font-weight: 600;
    color: #e8e6e0 !important;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
    font-family: 'DM Sans', sans-serif;
  }}
  .user-role {{
    font-size: 0.65rem;
    color: #6b6960 !important;
    font-family: 'DM Sans', sans-serif;
    margin-top: 1px;
  }}

  /* ── MONTHLY COUNTER BAR ── */
  .monthly-bar {{
    background: #1e1d1b;
    border: 1px solid #2e2d28;
    border-radius: 10px;
    padding: 10px 14px;
    margin-bottom: 0.5rem;
    font-family: 'DM Sans', sans-serif;
  }}
  .monthly-bar-header {{
    display: flex;
    justify-content: space-between;
    align-items: center;
    margin-bottom: 6px;
  }}
  .monthly-bar-label {{
    font-size: 0.72rem;
    font-weight: 700;
    color: #8a8880 !important;
    text-transform: uppercase;
    letter-spacing: 0.07em;
  }}
  .monthly-bar-count {{
    font-size: 0.78rem;
    font-weight: 800;
    color: #e8e6e0 !important;
  }}
  .monthly-bar-count.limit-near {{ color: #F59E0B !important; }}
  .monthly-bar-count.limit-reached {{ color: #EF4444 !important; }}
  .monthly-progress {{
    height: 5px;
    background: #2e2d28;
    border-radius: 100px;
    overflow: hidden;
  }}
  .monthly-progress-fill {{
    height: 100%;
    border-radius: 100px;
    transition: width 0.4s ease;
  }}

  @media (max-width: 640px) {{
    .block-container {{
      padding: 4.5rem 1rem 3rem !important;
    }}
    .hero-title {{ font-size: 56px !important; }}
    .hero-sub {{ font-size: 0.95rem !important; }}
    .hero-wrap {{ margin-bottom: 1.5rem; padding-bottom: 1.2rem; }}
    .top-bar {{
      justify-content: center;
      gap: 0.5rem;
    }}
    .stTextInput input,
    .stNumberInput input {{
      font-size: 1rem !important;
      padding: 14px 16px !important;
      min-height: 52px !important;
      height: 52px !important;
      line-height: 1.4 !important;
      box-sizing: border-box !important;
    }}
    .stTextInput > div > div {{
      min-height: 52px !important;
    }}
    .stTextInput input::placeholder,
    .stNumberInput input::placeholder {{
      font-size: 1rem !important;
      opacity: 1 !important;
    }}
    .stSelectbox [data-baseweb="select"] > div:first-child {{
      padding: 12px 14px !important;
      min-height: 50px !important;
      height: auto !important;
    }}
    .stSelectbox [data-baseweb="select"] span,
    .stSelectbox [data-baseweb="select"] div {{
      font-size: 0.95rem !important;
      white-space: nowrap !important;
      overflow: hidden !important;
      text-overflow: ellipsis !important;
    }}
    [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] > div:first-child {{
      min-height: 48px !important;
      padding: 10px 12px !important;
    }}
    [data-testid="stSidebar"] .stSelectbox [data-baseweb="select"] span {{
      font-size: 0.85rem !important;
      color: #f5f3ed !important;
    }}
    .stTextArea textarea {{
      font-size: 0.95rem !important;
    }}
    .stButton button {{
      min-height: 48px !important;
      font-size: 1rem !important;
    }}
    [data-testid="stSidebar"] .block-container {{
      padding: 1rem !important;
    }}
    .stDownloadButton button {{
      width: 100% !important;
      min-height: 48px !important;
    }}
  }}

  @media (min-width: 641px) and (max-width: 1024px) {{
    .block-container {{
      padding: 1.5rem 1.5rem 3rem !important;
      max-width: 900px !important;
    }}
  }}

  /* ── SIDEBAR DARK OVERRIDE ── */
  [data-testid="stSidebar"] .stButton > button,
  [data-testid="stSidebar"] .stButton > button[kind="secondary"],
  [data-testid="stSidebar"] button[data-testid="baseButton-secondary"],
  [data-testid="stSidebar"] .stDownloadButton button,
  [data-testid="stSidebar"] [data-testid="stDownloadButton"] button {{
    background: #1e1d1b !important;
    color: #e8e6e0 !important;
    border: 1.5px solid #3d3c36 !important;
    border-radius: 8px !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    box-shadow: none !important;
    transform: none !important;
    padding: 8px 14px !important;
    min-height: 36px !important;
    width: 100% !important;
  }}
  [data-testid="stSidebar"] .stButton > button:hover,
  [data-testid="stSidebar"] .stButton > button[kind="secondary"]:hover,
  [data-testid="stSidebar"] button[data-testid="baseButton-secondary"]:hover,
  [data-testid="stSidebar"] .stDownloadButton button:hover,
  [data-testid="stSidebar"] [data-testid="stDownloadButton"] button:hover {{
    background: #2a2926 !important;
    border-color: #D97706 !important;
    color: #f5f3ed !important;
    transform: none !important;
    box-shadow: none !important;
  }}

  [data-testid="stSidebar"] [data-testid="stExpander"] {{
    background: #1e1d1b !important;
    border: 1px solid #2e2d28 !important;
    border-radius: 10px !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] summary {{
    background: #1e1d1b !important;
    color: #c8c6bc !important;
    font-size: 0.82rem !important;
    padding: 0.7rem 1rem !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] summary:hover {{
    background: #2a2926 !important;
    color: #f5f3ed !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] > div > div {{
    background: #1e1d1b !important;
    padding: 0.4rem 1rem 0.8rem !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] p,
  [data-testid="stSidebar"] [data-testid="stExpander"] span {{
    color: #c8c6bc !important;
  }}

  /* Bottoni dentro gli expander della sidebar — stile dark coerente */
  [data-testid="stSidebar"] [data-testid="stExpander"] .stButton > button,
  [data-testid="stSidebar"] [data-testid="stExpander"] div.stButton > button,
  [data-testid="stSidebar"] [data-testid="stExpander"] button[kind="secondary"],
  [data-testid="stSidebar"] [data-testid="stExpander"] button {{
    background: #2a2926 !important;
    background-color: #2a2926 !important;
    color: #c8c6bc !important;
    border: 1px solid #3a3834 !important;
    border-radius: 8px !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    padding: 5px 12px !important;
    min-height: unset !important;
    box-shadow: none !important;
    transition: background 0.15s, border-color 0.15s, color 0.15s !important;
    width: 100% !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] .stButton > button:hover,
  [data-testid="stSidebar"] [data-testid="stExpander"] div.stButton > button:hover {{
    background: #353330 !important;
    background-color: #353330 !important;
    border-color: {T['accent']} !important;
    color: {T['accent']} !important;
    box-shadow: none !important;
    transform: none !important;
  }}
  /* Bottone elimina — bordo rosso sottile */
  [data-testid="stSidebar"] [data-testid="stExpander"] .elimina-btn .stButton > button,
  [data-testid="stSidebar"] [data-testid="stExpander"] .elimina-btn button {{
    background: #1e1008 !important;
    background-color: #1e1008 !important;
    border-color: #5c2222 !important;
    color: #f87171 !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] .elimina-btn .stButton > button:hover,
  [data-testid="stSidebar"] [data-testid="stExpander"] .elimina-btn button:hover {{
    background: #2a0f0f !important;
    background-color: #2a0f0f !important;
    border-color: #f87171 !important;
    color: #fca5a5 !important;
  }}
  /* Bottone stella preferito */
  [data-testid="stSidebar"] [data-testid="stExpander"] .stella-btn .stButton > button,
  [data-testid="stSidebar"] [data-testid="stExpander"] .stella-btn button {{
    background: #1e1d1b !important;
    background-color: #1e1d1b !important;
    border-color: #4a4020 !important;
    color: #9a8a50 !important;
    font-size: 1rem !important;
    padding: 3px 8px !important;
    width: auto !important;
  }}
  [data-testid="stSidebar"] [data-testid="stExpander"] .stella-btn-on .stButton > button,
  [data-testid="stSidebar"] [data-testid="stExpander"] .stella-btn-on button {{
    background: #2a2010 !important;
    background-color: #2a2010 !important;
    border-color: {T['accent']} !important;
    color: #F59E0B !important;
    width: auto !important;
  }}

  [data-testid="stSidebar"] .logout-btn-wrap .stButton > button,
  [data-testid="stSidebar"] .logout-btn-wrap button {{
    background: transparent !important;
    color: #f87171 !important;
    border: 1px solid #5c2222 !important;
    border-radius: 8px !important;
    font-size: 0.78rem !important;
    font-weight: 600 !important;
    padding: 6px 14px !important;
    width: auto !important;
    min-height: unset !important;
    box-shadow: none !important;
  }}
  [data-testid="stSidebar"] .logout-btn-wrap .stButton > button:hover,
  [data-testid="stSidebar"] .logout-btn-wrap button:hover {{
    background: #2a0f0f !important;
    border-color: #f87171 !important;
    color: #fca5a5 !important;
    box-shadow: none !important;
    transform: none !important;
  }}

  @keyframes tooltipSlideIn {{
    0%   {{ opacity: 0; transform: translateX(-12px); }}
    15%  {{ opacity: 1; transform: translateX(4px); }}
    25%  {{ opacity: 1; transform: translateX(0); }}
    75%  {{ opacity: 1; transform: translateX(0); }}
    90%  {{ opacity: 0; transform: translateX(-8px); }}
    100% {{ opacity: 0; transform: translateX(-8px); }}
  }}
  .sidebar-tooltip {{
    position: fixed;
    top: 0.9rem;
    left: 3.8rem;
    z-index: 9998;
    display: flex;
    align-items: center;
    gap: 7px;
    background: {T['accent']};
    color: #ffffff;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.75rem;
    font-weight: 700;
    letter-spacing: 0.02em;
    padding: 7px 13px 7px 10px;
    border-radius: 20px;
    box-shadow: 0 4px 18px {T['accent']}55;
    pointer-events: none;
    white-space: nowrap;
    animation: tooltipSlideIn 4.5s ease forwards;
  }}
  .sidebar-tooltip::before {{
    content: '';
    position: absolute;
    left: -6px;
    top: 50%;
    transform: translateY(-50%);
    border-width: 5px 6px 5px 0;
    border-style: solid;
    border-color: transparent {T['accent']} transparent transparent;
  }}
</style>
""", unsafe_allow_html=True)

# ── FEEDBACK BUTTON ──────────────────────────────────────────────────────────────
st.markdown(f"""
<a class="fab-link" href="{FEEDBACK_FORM_URL}" target="_blank" rel="noopener noreferrer"
   onclick="window.open(this.href,'_blank','noopener,noreferrer'); return false;">
  💬 &nbsp; Feedback & Bug
</a>
""", unsafe_allow_html=True)

# ── SIDEBAR ──────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown('<div class="sidebar-title">Impostazioni</div>', unsafe_allow_html=True)

    st.markdown('<div class="sidebar-label">Classe</div>', unsafe_allow_html=True)
    st.caption("Questa scelta calibra lessico, complessità e riferimenti teorici degli esercizi.")
    difficolta = st.selectbox("livello", SCUOLE, index=0, label_visibility="collapsed")

    st.markdown('<div class="sidebar-label" style="margin-top:1rem;">Opzioni</div>', unsafe_allow_html=True)
    bes_dsa = st.checkbox(
        "Genera versione ridotta (sostegno/certificazioni)",
        value=False,
        help="Verrà generato un secondo file identico ma con una percentuale di esercizi in meno."
    )
    perc_ridotta = None
    if bes_dsa:
        perc_ridotta = st.select_slider(
            "Esercizi da rimuovere",
            help="Es. 20% = verranno eliminati circa 1 esercizio ogni 5, partendo dai più complessi",
            options=[10, 20, 30],
            value=20,
            format_func=lambda x: f"-{x}%",
        )
    doppia_fila = st.checkbox("Genera Versione A e B (due varianti)", value=False)
    genera_soluzioni = st.checkbox(
        "📋 Genera soluzioni della verifica",
        value=False,
        help="Verrà generato un documento separato con le soluzioni complete. Per le domande aperte le risposte saranno sintetiche (max 5 righe)."
    )
    bes_dsa_b = False
    if bes_dsa and doppia_fila:
        bes_dsa_b = st.checkbox(
            "Genera versione ridotta anche per Fila B",
            value=False,
            help="Genera la versione ridotta (BES/DSA) anche per la Fila B"
        )

    esercizio_multidisciplinare = False
    materia2_scelta  = None
    difficolta_multi = None

    st.markdown('<div class="sidebar-label" style="margin-top:1rem;">Punteggi</div>', unsafe_allow_html=True)
    mostra_punteggi = st.checkbox("Mostra punteggio per esercizio", value=False)
    con_griglia     = st.checkbox("Includi griglia dei punteggi", value=False)
    punti_totali    = st.number_input("Punti totali", min_value=10, max_value=200, value=100, step=5,
                                      disabled=not mostra_punteggi)

    st.markdown('<div class="sidebar-label" style="margin-top:1rem;">Modello AI</div>', unsafe_allow_html=True)
    if _is_admin:
        modello_id = MODELLI_DISPONIBILI[
            st.selectbox("modello", list(MODELLI_DISPONIBILI.keys()), label_visibility="collapsed")
        ]
    else:
        modello_id = "gemini-2.5-flash-lite"
        st.markdown(f"""
        <div style="background:{T['card2']};border:1px solid {T['border']};border-radius:10px;
                    padding:8px 12px;font-size:0.8rem;color:{T['text2']};
                    font-family:'DM Sans',sans-serif;">
          Flash 2.5 Lite
          <span style="font-size:0.7rem;color:{T['muted']};margin-left:6px;">modello standard</span>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div class="sidebar-label" style="margin-top:1rem;">Aspetto</div>', unsafe_allow_html=True)
    tema_sel = st.radio(
        "tema",
        ["☀️ Chiaro", "🌙 Scuro"],
        index=0 if st.session_state.theme == "light" else 1,
        horizontal=True,
        label_visibility="collapsed"
    )
    new_theme = "light" if "Chiaro" in tema_sel else "dark"
    if new_theme != st.session_state.theme:
        st.session_state.theme = new_theme
        st.rerun()

    # ── NUOVO: CONTATORE MENSILE ──────────────────────────────────────────────
    st.markdown('<div class="sidebar-label" style="margin-top:1.5rem;">Utilizzo mensile</div>', unsafe_allow_html=True)
    _perc_uso = min(100, int(_verifiche_mese_count / LIMITE_MENSILE * 100))
    _color_bar = "#EF4444" if _limite_raggiunto else ("#F59E0B" if _perc_uso >= 70 else "#10B981")
    _count_class = "limit-reached" if _limite_raggiunto else ("limit-near" if _perc_uso >= 70 else "")
    _gg_reset, _hh_reset = _giorni_al_reset()
    if _gg_reset == 0:
        _reset_str = f"Reset tra {_hh_reset}h"
    elif _gg_reset == 1:
        _reset_str = f"Reset domani"
    else:
        _reset_str = f"Reset tra {_gg_reset}gg"
    st.markdown(f"""
    <div class="monthly-bar">
      <div class="monthly-bar-header">
        <span class="monthly-bar-label">Verifiche questo mese</span>
        <span class="monthly-bar-count {_count_class}">{_verifiche_mese_count} / {LIMITE_MENSILE}</span>
      </div>
      <div class="monthly-progress">
        <div class="monthly-progress-fill" style="width:{_perc_uso}%;background:{_color_bar};"></div>
      </div>
      <div style="text-align:right;font-size:0.68rem;color:#6b6960;margin-top:4px;font-family:'DM Sans',sans-serif;">
        🔄 {_reset_str}
      </div>
    </div>
    """, unsafe_allow_html=True)
    if _limite_raggiunto:
        st.warning(f"Limite mensile raggiunto ({LIMITE_MENSILE} verifiche). {_reset_str}.")

    # ── STORICO VERIFICHE ─────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-label" style="margin-top:1rem;">Le mie verifiche</div>', unsafe_allow_html=True)

    _refresh_key = st.session_state._storico_refresh
    _page_size   = 5
    _storico_limit = st.session_state._storico_page * _page_size
    try:
        storico = supabase_admin.table("verifiche_storico")\
            .select("id, materia, argomento, created_at, latex_a, latex_b, latex_r, scuola")\
            .eq("user_id", st.session_state.utente.id)\
            .is_("deleted_at", "null")\
            .order("created_at", desc=True)\
            .limit(_storico_limit + 1)\
            .execute()  # fetch +1 per sapere se ci sono altri

        if storico.data:
            _ha_altri = len(storico.data) > _storico_limit
            dati_pagina = storico.data[:_storico_limit]

            # Metti in cima i preferiti
            _pref = st.session_state._preferiti
            def _sort_key(v):
                return (0 if v['id'] in _pref else 1, v['created_at'])
            dati_ordinati = sorted(dati_pagina, key=_sort_key)

            for v in dati_ordinati:
                data_str = v['created_at'][:10]
                is_pref  = v['id'] in _pref
                star_ico = "★" if is_pref else "☆"
                star_label_prefix = "⭐ " if is_pref else ""
                label = f"{star_label_prefix}{v['materia']} — {v['argomento'][:20]}{'...' if len(v['argomento'])>20 else ''}"
                with st.expander(f"{label} ({data_str})"):
                    if v.get('scuola'):
                        st.caption(f"{v['scuola'][:35]}")

                    # Riga: stella preferiti
                    _col_star, _col_spacer = st.columns([1, 3])
                    with _col_star:
                        st.markdown(f'<div class="{"stella-btn-on" if is_pref else "stella-btn"}">', unsafe_allow_html=True)
                        if st.button(star_ico, key=f"star_{v['id']}_{_refresh_key}",
                                     help="Aggiungi/rimuovi dai preferiti"):
                            if v['id'] in st.session_state._preferiti:
                                st.session_state._preferiti.discard(v['id'])
                            else:
                                st.session_state._preferiti.add(v['id'])
                            st.rerun()
                        st.markdown('</div>', unsafe_allow_html=True)

                    if v.get('latex_a'):
                        if st.button("♻ Ricarica Fila A", key=f"reload_a_{v['id']}_{_refresh_key}", use_container_width=True):
                            st.session_state.verifiche['A']['latex'] = v['latex_a']
                            pdf, _ = compila_pdf(v['latex_a'])
                            if pdf:
                                st.session_state.verifiche['A']['pdf'] = pdf
                                st.session_state.verifiche['A']['preview'] = True
                            st.rerun()
                    if v.get('latex_b'):
                        if st.button("♻ Ricarica Fila B", key=f"reload_b_{v['id']}_{_refresh_key}", use_container_width=True):
                            st.session_state.verifiche['B']['latex'] = v['latex_b']
                            pdf, _ = compila_pdf(v['latex_b'])
                            if pdf:
                                st.session_state.verifiche['B']['pdf'] = pdf
                                st.session_state.verifiche['B']['preview'] = True
                            st.rerun()
                    # Bottone elimina
                    st.markdown('<div class="elimina-btn">', unsafe_allow_html=True)
                    if st.button("Elimina", key=f"del_{v['id']}_{_refresh_key}",
                                 use_container_width=True,
                                 help="Rimuovi questa verifica dallo storico"):
                        try:
                            from datetime import datetime, timezone
                            supabase_admin.table("verifiche_storico")\
                                .update({"deleted_at": datetime.now(timezone.utc).isoformat()})\
                                .eq("id", v['id'])\
                                .execute()
                            st.session_state._preferiti.discard(v['id'])
                            st.session_state._storico_refresh += 1
                            st.toast("Verifica rimossa dallo storico.", icon="🗑️")
                            st.rerun()
                        except Exception as del_err:
                            st.error(f"Errore: {del_err}")
                    st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.caption("Nessuna verifica salvata ancora.")

        # ── Paginazione: carica altri ─────────────────────────────────────────
        if storico.data and _ha_altri:
            st.markdown('<div style="margin-top:0.5rem;">', unsafe_allow_html=True)
            if st.button("Carica altre verifiche", key="storico_load_more", use_container_width=True):
                st.session_state._storico_page += 1
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        elif storico.data and st.session_state._storico_page > 1:
            st.caption(f"Tutte le {len(dati_pagina)} verifiche caricate.")

    except Exception as e:
        st.caption("Storico non disponibile.")

    # ── USER PILL + LOGOUT ────────────────────────────────────────────────────
    st.markdown("---")
    email_utente = st.session_state.utente.email or ""
    iniziale = email_utente[0].upper() if email_utente else "?"
    st.markdown(f"""
    <div class="user-pill">
      <div class="user-avatar">{iniziale}</div>
      <div class="user-info">
        <div class="user-email">{email_utente}</div>
        <div class="user-role">Docente · Piano gratuito</div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div style="height:0.5rem;"></div>', unsafe_allow_html=True)
    st.markdown('<div class="logout-btn-wrap">', unsafe_allow_html=True)
    if st.button("↩ Esci dall'account", key="logout_btn", use_container_width=False):
        supabase.auth.sign_out()
        st.session_state.utente = None
        st.session_state.pop("_sb_access_token", None)
        st.session_state.pop("_sb_refresh_token", None)
        st.query_params.pop("_at", None)
        st.query_params.pop("_rt", None)
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ── TOPBAR ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="top-bar">
  <div class="top-bar-hint">
    ← Apri le impostazioni per configurare classe, opzioni e modello AI
  </div>
</div>
""", unsafe_allow_html=True)



# ── HEADER ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-wrap">
  <div class="hero-left">
    <h1 class="hero-title"><span class="hero-icon">{APP_ICON}</span> Verific<span class="hero-ai">AI</span></h1>
    <p class="hero-sub">{APP_TAGLINE}</p>
    <span class="hero-beta">Versione Beta</span>
  </div>
</div>
""", unsafe_allow_html=True)

# ── FORM PRINCIPALE ───────────────────────────────────────────────────────────────

# ── ONBOARDING: guida al primo accesso ───────────────────────────────────────────
if not st.session_state._onboarding_done:
    _c_accent       = T['accent']
    _c_text         = T['text']
    _c_text2        = T['text2']
    _c_bg2          = T['bg2']
    _c_border       = T['border']
    _c_card         = T['card']
    _c_accent_light = T['accent_light']
    _c_muted        = T['muted']

    # Banner
    st.markdown(
        f'<div style="background:linear-gradient(135deg,{_c_accent_light} 0%,{_c_card} 100%);'
        f'border:1.5px solid {_c_accent};border-radius:14px;'
        f'padding:1.1rem 1.4rem 1rem 1.4rem;margin-bottom:0.3rem;font-family:DM Sans,sans-serif;">'
        f'<div style="display:flex;align-items:flex-start;gap:12px;">'
        f'<div style="font-size:1.4rem;flex-shrink:0;margin-top:1px;">👋</div>'
        f'<div style="flex:1;">'
        f'<div style="font-size:0.9rem;font-weight:800;color:{_c_text};margin-bottom:0.7rem;">Come iniziare</div>'
        f'<div style="display:flex;align-items:center;gap:8px;padding:0.45rem 0.75rem;margin-bottom:0.5rem;'
        f'background:{_c_bg2};border-radius:8px;border-left:3px solid {_c_accent};">'
        f'<span>⚙️</span>'
        f'<div style="font-size:0.8rem;color:{_c_text2};">Prima di tutto: apri '
        f'<strong style="color:{_c_text};">☰ Impostazioni</strong> in alto a sinistra '
        f'per scegliere classe e modello AI</div>'
        f'</div>'
        f'<div style="display:flex;background:{_c_bg2};border:1px solid {_c_border};border-radius:10px;overflow:hidden;">'
        f'<div style="flex:1;padding:0.6rem 0.85rem;border-right:1px solid {_c_border};">'
        f'<div style="font-size:0.65rem;font-weight:800;color:{_c_accent};text-transform:uppercase;letter-spacing:0.06em;margin-bottom:3px;">01 · Materia</div>'
        f'<div style="font-size:0.76rem;color:{_c_text2};">Scegli la materia</div>'
        f'</div>'
        f'<div style="flex:1;padding:0.6rem 0.85rem;border-right:1px solid {_c_border};">'
        f'<div style="font-size:0.65rem;font-weight:800;color:{_c_accent};text-transform:uppercase;letter-spacing:0.06em;margin-bottom:3px;">02 · Argomento</div>'
        f'<div style="font-size:0.76rem;color:{_c_text2};">Scrivi l\'argomento</div>'
        f'</div>'
        f'<div style="flex:1;padding:0.6rem 0.85rem;">'
        f'<div style="font-size:0.65rem;font-weight:800;color:{_c_accent};text-transform:uppercase;letter-spacing:0.06em;margin-bottom:3px;">03 · Personalizza</div>'
        f'<div style="font-size:0.76rem;color:{_c_text2};">Opzioni avanzate (facoltativo)</div>'
        f'</div>'
        f'</div>'
        f'</div>'
        f'</div>'
        f'</div>',
        unsafe_allow_html=True
    )

    # Riga "Ho capito" — semplice, allineata a destra con CSS
    st.markdown(f"""
    <style>
    .ob-row {{
        display: flex;
        justify-content: flex-end;
        margin-bottom: 0.8rem;
        margin-top: 0;
    }}
    .ob-row div.stButton > button {{
        background: transparent !important;
        border: none !important;
        color: {_c_muted} !important;
        font-size: 0.78rem !important;
        font-weight: 500 !important;
        text-decoration: underline !important;
        text-underline-offset: 2px !important;
        padding: 2px 0 !important;
        min-height: unset !important;
        height: auto !important;
        box-shadow: none !important;
        transform: none !important;
        width: auto !important;
    }}
    .ob-row div.stButton > button:hover {{
        color: {_c_text} !important;
        background: transparent !important;
        box-shadow: none !important;
        transform: none !important;
    }}
    </style>
    <div class="ob-row">
    """, unsafe_allow_html=True)

    if st.button("Ho capito, non mostrare più →", key="_dismiss_onboarding"):
        st.session_state._onboarding_done = True
        st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    # Disclaimer
    st.markdown(
        f'<div style="display:flex;align-items:flex-start;gap:8px;padding:0.6rem 0.9rem;'
        f'margin-bottom:1.2rem;background:{_c_bg2};border-radius:8px;border:1px solid {_c_border};">'
        f'<span style="font-size:0.85rem;flex-shrink:0;margin-top:1px;">⚠️</span>'
        f'<span style="font-size:0.75rem;color:{_c_muted};line-height:1.45;">'
        f'Le verifiche generate dall\'AI sono <strong style="color:{_c_text2};">suggerimenti didattici</strong>. '
        f'Controlla sempre contenuti e punteggi prima di distribuirle agli studenti. '
        f'Il docente rimane responsabile del materiale finale.'
        f'</span>'
        f'</div>',
        unsafe_allow_html=True
    )



        
# STEP 1 — MATERIA
st.markdown(f"""
<div class="step-label">
  <span class="step-num">01</span>
  <span class="step-title">Materia</span>
  <span class="step-line"></span>
</div>
""", unsafe_allow_html=True)
_materie_select = MATERIE + ["✏️ Altra materia..."]
_materia_sel = st.selectbox("Materia", _materie_select, index=0, label_visibility="collapsed")
if _materia_sel == "✏️ Altra materia...":
    materia_scelta = st.text_input("Scrivi materia:", placeholder="es. Economia Aziendale, Scienze Naturali...",
                                   key="_materia_custom_input", label_visibility="collapsed").strip() or "Matematica"
else:
    materia_scelta = _materia_sel or "Matematica"

# STEP 2 — ARGOMENTO
st.markdown(f"""
<div class="step-label">
  <span class="step-num">02</span>
  <span class="step-title">Argomento della verifica</span>
  <span class="step-line"></span>
</div>
""", unsafe_allow_html=True)
argomento_area = st.text_area(
    "argomento",
    placeholder="es. Le equazioni di secondo grado\nes. La Rivoluzione Francese",
    height=100,
    label_visibility="collapsed",
    key="argomento_area"
)
argomento = argomento_area.strip()

# STEP 3 — PERSONALIZZA
st.markdown(f"""
<div class="ai-hint">
  <span class="ai-hint-icon">✨</span>
  <span><strong>Suggerimento:</strong> più dettagli fornisci nell'argomento e nelle opzioni avanzate, più la verifica sarà precisa e su misura.</span>
</div>
""", unsafe_allow_html=True)

st.markdown(f"""
<div class="step-label">
  <span class="step-num">03</span>
  <span class="step-title">Personalizza</span>
  <span class="step-line"></span>
</div>
""", unsafe_allow_html=True)

st.markdown('<div class="personalizza-wrap">', unsafe_allow_html=True)
with st.expander("Personalizza la verifica *(opzionale)*"):

    st.markdown(f'<div class="expander-heading">⏱️ Tempistiche e Struttura</div>', unsafe_allow_html=True)
    _c_dur, _c_num = st.columns(2)
    with _c_dur:
        durata_scelta = st.selectbox(
            "Durata della verifica",
            ["30 min", "1 ora", "1 ora e 30 min", "2 ore"],
            index=1,
        )
    with _c_num:
        num_esercizi_totali = st.slider(
            "Numero di esercizi in verifica",
            min_value=1, max_value=15, value=4,
        )

    with st.expander("🎯 Specifica il tipo di ogni esercizio (opzionale)"):
        n_custom = len(st.session_state.esercizi_custom)
        n_liberi = max(0, num_esercizi_totali - n_custom)

        if n_custom == 0:
            st.info(f"Tutti i {num_esercizi_totali} esercizi generati dall'AI.")
        elif n_custom >= num_esercizi_totali:
            st.warning(f"Totale {n_custom}/{num_esercizi_totali} raggiunto — aumenta il numero.")
        else:
            st.success(f"✅ {n_custom} specifici + {n_liberi} liberi = {num_esercizi_totali} totali")

        if st.session_state.esercizi_custom:
            to_remove = None
            for i, ex in enumerate(st.session_state.esercizi_custom):
                st.markdown(f'<div class="expander-heading">Esercizio {i+1}</div>', unsafe_allow_html=True)
                t = st.selectbox("Tipo esercizio", TIPI_ESERCIZIO,
                                 index=TIPI_ESERCIZIO.index(ex.get('tipo', 'Aperto')),
                                 key=f"tipo_{i}", label_visibility="visible")
                st.session_state.esercizi_custom[i]['tipo'] = t

                d = st.text_input("Descrizione dell'esercizio (opzionale)",
                                  value=ex.get('descrizione', ''),
                                  placeholder="es. Risolvi ax²+bx+c=0 mostrando i passaggi",
                                  key=f"desc_{i}", label_visibility="visible")
                st.session_state.esercizi_custom[i]['descrizione'] = d
                _cimg, _cdel = st.columns([3, 1])
                with _cimg:
                    st.markdown('<div class="compact-uploader">', unsafe_allow_html=True)
                    img = st.file_uploader("📎 Immagine allegata",
                                           type=['png','jpg','jpeg'],
                                           key=f"img_{i}", label_visibility="collapsed")
                    st.markdown('</div>', unsafe_allow_html=True)
                    if img: st.session_state.esercizi_custom[i]['immagine'] = img
                    if st.session_state.esercizi_custom[i].get('immagine'):
                        st.image(st.session_state.esercizi_custom[i]['immagine'], width=60)
                with _cdel:
                    st.markdown('<div style="padding-top:4px;">', unsafe_allow_html=True)
                    if st.button("🗑 Rimuovi", key=f"rm_{i}", use_container_width=True):
                        to_remove = i
                    st.markdown('</div>', unsafe_allow_html=True)
                st.markdown('<hr style="margin:0.8rem 0; opacity:0.15;">', unsafe_allow_html=True)
            if to_remove is not None:
                st.session_state.esercizi_custom.pop(to_remove); st.rerun()

        can_add = len(st.session_state.esercizi_custom) < num_esercizi_totali
        if st.button("＋ Aggiungi esercizio specifico", disabled=not can_add):
            st.session_state.esercizi_custom.append({'tipo': 'Aperto', 'descrizione': '', 'immagine': None, 'materia2': '', 'difficolta_multi': 'Media'})
            st.rerun()

    st.markdown('<div class="expander-heading" style="margin-top:1rem;">🎯 Istruzioni per l\'AI</div>', unsafe_allow_html=True)
    note_generali = st.text_area(
        "note", label_visibility="collapsed",
        placeholder=NOTE_PLACEHOLDER.get(materia_scelta, "es. Argomenti da privilegiare, tipo di esercizi..."),
        height=80
    )

    st.markdown(f'<div class="expander-heading">📂 File di riferimento</div>', unsafe_allow_html=True)
    st.markdown('<div class="compact-uploader">', unsafe_allow_html=True)
    file_ispirazione = st.file_uploader(
        "📎 Allega PDF o immagine",
        type=['pdf','png','jpg','jpeg'],
        label_visibility="collapsed"
    )
    st.markdown('</div>', unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# ── BOTTONE GENERA ────────────────────────────────────────────────────────────────
st.markdown('<div class="genera-section">', unsafe_allow_html=True)

# ── NUOVO: bottone disabilitato se limite raggiunto ──────────────────────────────
genera_btn = st.button(
    "🚀  Genera Verifica",
    use_container_width=True,
    type="primary",
    disabled=_limite_raggiunto
)

if _limite_raggiunto:
    st.markdown(f"""
    <div style="text-align:center;font-size:0.82rem;color:#EF4444;margin-top:0.5rem;
                font-family:'DM Sans',sans-serif;font-weight:600;">
      ⛔ Limite di {LIMITE_MENSILE} verifiche mensili raggiunto. Riprova il mese prossimo.
    </div>
    """, unsafe_allow_html=True)
else:
    # ── hint dinamico: riepilogo file che verranno generati ──────────────
    _files_hint = ["📄 Fila A"]
    if doppia_fila:
        _files_hint.append("📄 Fila B")
    if bes_dsa:
        _files_hint.append("♿ Ridotta A")
    if bes_dsa and doppia_fila and bes_dsa_b:
        _files_hint.append("♿ Ridotta B")
    if genera_soluzioni:
        _files_hint.append("✅ Soluzioni")
    _files_str = " · ".join(_files_hint)
    _n_files   = len(_files_hint)
    st.markdown(f"""
    <div class="genera-hint">
      <strong style="color:{T['text2']};">File richiesti ({_n_files}):</strong> {_files_str}
    </div>
    """, unsafe_allow_html=True)

st.markdown('</div>', unsafe_allow_html=True)

# ── LOGICA GENERAZIONE ───────────────────────────────────────────────────────────
if genera_btn and not _limite_raggiunto:
    if not argomento.strip():
        st.warning("⚠️ Inserisci l'argomento della verifica."); st.stop()
    try:
        model        = genai.GenerativeModel(modello_id)
        materia      = materia_scelta.strip() or "Matematica"
        e_mat        = any(k in materia.lower() for k in ["matem","fis","chim","inform","elettr","meccan"])
        nota_bes = "Svolgere tutti gli esercizi mostrando i passaggi."
        calibrazione = CALIBRAZIONE_SCUOLA.get(difficolta, "")
        s_note       = f"\nNOTE DOCENTE: {note_generali.strip()}" if note_generali.strip() else ""
        s_es, imgs_es = costruisci_prompt_esercizi(
            st.session_state.esercizi_custom, num_esercizi_totali,
            punti_totali if mostra_punteggi else 0, mostra_punteggi)
        titolo_a = "Versione A" if doppia_fila else ""

        _n_steps = 4 + (2 if doppia_fila else 0) + (1 if bes_dsa else 0) + (1 if bes_dsa and doppia_fila and bes_dsa_b else 0) + (1 if genera_soluzioni else 0) + (1 if genera_soluzioni and doppia_fila else 0)
        _step    = [0]
        _prog    = st.empty()

        def _avanza(testo):
            _step[0] += 1
            perc = int(min(_step[0] / _n_steps, 0.97) * 100)
            _prog.markdown(f"""
<div style="margin:0.6rem 0 1rem 0;">
  <div style="font-size:0.82rem;font-weight:600;color:{T['text2']};
              font-family:'DM Sans',sans-serif;margin-bottom:6px;">{testo}</div>
  <div style="background:{T['border']};border-radius:100px;height:8px;overflow:hidden;">
    <div style="background:linear-gradient(90deg,{T['accent']},{T['accent']}cc);
                width:{perc}%;height:100%;border-radius:100px;
                transition:width 0.4s ease;"></div>
  </div>
</div>
""", unsafe_allow_html=True)

        _avanza("✍️  Elaborazione titolo…")

        titolo_resp = model.generate_content(
            f"Sei un docente. Crea un titolo professionale e conciso per una verifica scolastica.\n"
            f"Materia: {materia}\n"
            f"Argomento inserito dall'utente (potrebbe avere errori ortografici o essere informale): \"{argomento}\"\n"
            f"Restituisci SOLO il titolo senza virgolette, senza punteggiatura finale, "
            f"senza prefissi come 'Verifica di'. Esempio: 'Le equazioni di secondo grado'"
        )
        titolo_clean = titolo_resp.text.strip().strip('"').strip("'").strip()
        if not titolo_clean:
            titolo_clean = argomento.strip()
        _avanza("🧠  Generazione esercizi in corso…")

        bes_rule = "- NON inserire mai il simbolo (*) accanto a nessun sottopunto."

        if mostra_punteggi:
            punti_rule = (
                f"- PUNTEGGI OBBLIGATORI: ogni \\item DEVE avere \"(X pt)\" SULLA STESSA RIGA, subito dopo il testo.\n"
                f"- Formato ESATTO e UNICO accettato: (X pt) — esempio: \\item[a)] Risolvi l'equazione. (3 pt)\n"
                f"- NON usare formati alternativi come [X pt], X punti, X p., pt X, ecc.\n"
                f"- La somma di TUTTI i (X pt) deve essere ESATTAMENTE {punti_totali} pt. CONTROLLA prima di terminare.\n"
                f"- Distribuisci i punti in modo che sia facile ottenere almeno 60% svolgendo le parti più semplici.\n"
                f"- NON inserire punti nel titolo \\subsection*, SOLO nei \\item.\n"
                f"- Se dimentichi il punteggio anche su UN SOLO \\item, la griglia di valutazione sarà incompleta."
            )
        else:
            punti_rule = "- NON inserire punti (X pt) in nessun esercizio né sottopunto."

        if esercizio_multidisciplinare:
            materia2_str   = f" con {materia2_scelta}" if materia2_scelta else " (scegli tu la disciplina più adatta)"
            diff_multi_str = f" Difficoltà: {difficolta_multi}." if difficolta_multi else ""
            multi_rule = (
                f"- ESERCIZIO MULTIDISCIPLINARE: uno degli esercizi INCLUSI NEL TOTALE deve collegare "
                f"{materia}{materia2_str}.{diff_multi_str}\n"
                "  Usa SOLO strumenti già acquisiti dagli studenti."
            )
        else:
            multi_rule = "- NON includere esercizi multidisciplinari."

        griglia_rule = ("- NON generare la griglia (sarà aggiunta automaticamente dopo)."
                        if con_griglia else "- NON generare nessuna griglia di valutazione.")

        if e_mat:
            grafici_rule = (
                "- GRAFICI pgfplots: quando il grafico è un DATO fornito allo studente "
                "(es. 'osserva il grafico della parabola e determina...', 'dal grafico ricava...'), "
                "DEVI obbligatoriamente generarlo con pgfplots/tikzpicture. "
                "Esempio per parabola: \\begin{tikzpicture}\\begin{axis}[width=7cm,height=5.5cm,"
                "axis lines=center,xlabel=$x$,ylabel=$y$,grid=both,xtick={-4,...,4},ytick={-4,...,4}]"
                "\\addplot[blue,thick,domain=-3:3,samples=100]{x^2-2*x-3}; \\end{axis}\\end{tikzpicture} "
                "MI RACCOMANDO NON lasciare MAI spazio extra per disegnare se lo studente deve disegnare lui il grafico, lo fara sul suo foglio."
            )
            pgfplots_pkg = "\\usepackage{pgfplots}\n\\pgfplotsset{compat=1.18}\n\\usepackage{tikz}"
        else:
            grafici_rule = ""
            pgfplots_pkg = ""

        titolo_header = f"Verifica di {materia}: {titolo_clean}" + (f" — {titolo_a}" if titolo_a else "")
        _hspace6 = "{6cm}"
        _hspace4 = "{4cm}"
        preambolo_fisso = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[italian]{{babel}}
\\usepackage{{amsmath,amsfonts,amssymb,geometry,array,multicol,enumerate,adjustbox,wasysym}}
{pgfplots_pkg}
\\geometry{{margin=1.5cm}}
\\setlength{{\\parskip}}{{3pt plus1pt minus1pt}}
\\pagestyle{{empty}}
\\begin{{document}}
\\begin{{center}}
  \\textbf{{\\large {titolo_header}}} \\\\
  \\vspace{{0.3cm}}
  \\small \\textbf{{Nome:}} \\underline{{\\hspace{_hspace6}}} \\quad \\textbf{{Classe e Data:}} \\underline{{\\hspace{_hspace4}}} \\\\
  \\vspace{{0.3cm}}
  \\textit{{\\small {nota_bes}}}
\\end{{center}}
"""

        prompt_a = f"""Sei un docente esperto di {materia} e LaTeX. Genera SOLO il corpo degli esercizi (senza preambolo, senza \\documentclass, senza \\begin{{document}}) per una verifica su: {argomento}.
{f'Punti totali da distribuire: {punti_totali} pt.' if mostra_punteggi else ''}

CALIBRAZIONE LIVELLO E TEMPO:
{calibrazione}
- DURATA PREVISTA: {durata_scelta}. Regola la lunghezza dei calcoli, il numero di incognite e la complessità testuale in modo che {num_esercizi_totali} esercizi siano agevolmente fattibili nel tempo scelto.
- BILANCIAMENTO CONTESTO E MODELLAZIONE: NON esagerare con i problemi applicati alla realtà o fortemente interdisciplinari. MASSIMO 1 o 2 esercizi possono essere contestualizzati. I restanti DEVONO essere esercizi canonici, diretti e focalizzati sulla procedura pura.
- REGISTRO LINGUISTICO — REGOLA ASSOLUTA: il testo degli esercizi deve essere CONCISO e DIRETTO.
- DATI PULITI — REGOLA ASSOLUTA: prima di scrivere ogni esercizio, risolvilo mentalmente tu stesso. Scegli SOLO dati che portano a risultati interi o frazioni semplici. MAI scegliere dati che rendono un sistema contraddittorio, sovradeterminato o senza soluzione unica (a meno che non sia esplicitamente richiesto). Se un esercizio chiede di trovare un'equazione soddisfacendo N condizioni, verifica che le N condizioni siano compatibili tra loro.
REGOLE TASSATIVE SUI GRAFICI (LOGICA ANTI-SPOILER):
- Se l'esercizio richiede allo studente di "disegnare", "rappresentare graficamente", "tracciare" o "costruire" una figura/grafico, NON generare il codice TikZ.
- Genera un grafico (TikZ) SOLO se esso è un dato di partenza necessario fornito dal docente.

{s_note}
{s_es}

REGOLE LATEX (TASSATIVE):
{griglia_rule}
{punti_rule}
- NUMERO ESERCIZI: genera ESATTAMENTE {num_esercizi_totali} blocchi \\subsection*. CONTA i tuoi blocchi prima di chiudere.
- Titoli: \\subsection*{{Esercizio N: Titolo}}
- SOTTOPUNTI OBBLIGATORI: usa SEMPRE \\item[a)] \\item[b)] \\item[c)] ecc. con label ESPLICITA tra parentesi quadre.
- PROTEZIONE ESERCIZIO 1 (Saperi Essenziali): nell'Esercizio 1 NON inserire MAI il simbolo (*) su nessun sottopunto.
{multi_rule}
- Scelta multipla: le opzioni DEVONO stare in un \\begin{{enumerate}}[a)] SEPARATO dopo la domanda.
- Vero/Falso: $\\square$ \\textbf{{V}} $\\quad\\square$ \\textbf{{F}}
- Completamento: \\underline{{\\hspace{{3cm}}}}
{grafici_rule}

FORMATO OUTPUT: restituisci SOLO i blocchi \\subsection*{{...}} con relativi esercizi.
TERMINA con \\end{{document}}.
NIENTE preambolo, NIENTE \\documentclass, NIENTE \\begin{{document}}.
SOLO CODICE LATEX del corpo."""

        inp = [prompt_a]
        if file_ispirazione:
            inp.append({"mime_type": file_ispirazione.type, "data": file_ispirazione.getvalue()})
            inp[0] += "\nPrendi spunto dal file allegato per stile e livello."
        for im in imgs_es:
            inp.append({"mime_type": im['mime_type'], "data": im['data']})
            inp[0] += f"\nUsa l'immagine come riferimento per l'Esercizio {im['idx']}."

        ra = model.generate_content(inp)
        _avanza("⚙️  Elaborazione LaTeX…")

        corpo_latex = ra.text.replace("```latex","").replace("```","").strip()
        corpo_latex = pulisci_corpo_latex(corpo_latex)

        # ── SELF-CHECK: VERIFICA E CORREZIONE AUTOMATICA ─────────────────────
        _avanza("🔎  Controllo qualità e correzione errori…")
        prompt_check = f"""Sei un docente esperto di {materia} e devi fare un CONTROLLO DI QUALITÀ RIGOROSO su questa verifica scolastica prima che venga consegnata agli studenti.

MATERIA: {materia}
LIVELLO: {difficolta}
VERIFICA DA CONTROLLARE:
{corpo_latex}

COMPITO: analizza OGNI esercizio e OGNI sottopunto. Per ciascuno verifica:

1. CORRETTEZZA MATEMATICA / DISCIPLINARE: i dati sono coerenti? L'esercizio ha UNA soluzione determinata e corretta? Se risolvo l'esercizio io stesso, ottengo una risposta pulita e sensata?
   - Esempi di ERRORI GRAVI: sistema sovradeterminato o contraddittorio, dati incoerenti (es. due condizioni incompatibili), risposta che richiede conoscenze non adatte al livello, calcoli che portano a risultati assurdi.
   
2. ADEGUATEZZA AL LIVELLO ({difficolta}): la complessità è appropriata? Un esercizio non deve richiedere 30 passaggi se il livello è basso.

3. UNIVOCITÀ: la domanda ha una sola risposta corretta e non è ambigua?

SE trovi problemi: CORREGGILI DIRETTAMENTE modificando i dati dell'esercizio (cambia numeri, coordinate, coefficienti, punti) finché l'esercizio sia corretto, sensato e risolvibile. NON eliminare esercizi, correggili.

SE tutto è corretto: restituisci il testo IDENTICO senza modifiche.

REGOLE OUTPUT:
- Restituisci SOLO il corpo LaTeX corretto (\\subsection* ecc.), senza preambolo.
- Mantieni ESATTAMENTE la stessa struttura LaTeX (\\item[a)], \\item[b)], ecc.).
- NON aggiungere commenti, spiegazioni o note al di fuori del LaTeX.
- TERMINA con \\end{{document}}.
- Se hai modificato dati, mantieni la stessa difficoltà complessiva e lo stesso tipo di esercizio."""

        rc = model.generate_content(prompt_check)
        corpo_latex_corretto = rc.text.replace("```latex","").replace("```","").strip()
        corpo_latex_corretto = pulisci_corpo_latex(corpo_latex_corretto)

        # Usa il corpo corretto solo se il modello ha restituito qualcosa di sensato
        # (stessa struttura, stesso numero di \subsection*)
        _n_orig = len(re.findall(r'\\subsection\*', corpo_latex))
        _n_corr = len(re.findall(r'\\subsection\*', corpo_latex_corretto))
        if corpo_latex_corretto and _n_corr == _n_orig:
            corpo_latex = corpo_latex_corretto
        # Se il modello ha restituito struttura diversa, manteniamo l'originale

        splits = re.split(r'(\\subsection\*\{)', corpo_latex)
        n_blocchi = (len(splits) - 1) // 2
        if n_blocchi > num_esercizi_totali:
            testa = splits[0]
            blocchi_da_tenere = []
            for b in range(num_esercizi_totali):
                blocchi_da_tenere.append(splits[1 + b*2])
                blocchi_da_tenere.append(splits[2 + b*2])
            corpo_troncato = testa + "".join(blocchi_da_tenere)
            corpo_troncato = re.sub(r'\\end\{document\}.*$', '', corpo_troncato, flags=re.DOTALL).rstrip()
            corpo_latex = corpo_troncato + "\n\\end{document}"

        latex_a = preambolo_fisso + corpo_latex
        latex_a = fix_items_environment(latex_a)
        latex_a = rimuovi_vspace_corpo(latex_a)
        if mostra_punteggi:
            latex_a = rimuovi_punti_subsection(latex_a)
            latex_a = riscala_punti(latex_a, punti_totali)

        if con_griglia:
            latex_a_final = inietta_griglia(latex_a, punti_totali)
        else:
            latex_a_final = latex_a

        st.session_state.verifiche['A'] = {**_vf(), 'latex': latex_a_final}
        st.session_state.verifiche['A']['latex_originale'] = latex_a_final

        _avanza("🖨️  Compilazione PDF…")
        pdf_auto, err_auto = compila_pdf(latex_a_final)
        if pdf_auto:
            st.session_state.verifiche['A']['pdf']     = pdf_auto
            st.session_state.verifiche['A']['pdf_ts']  = time.time()
            st.session_state.verifiche['A']['preview'] = True
        else:
            if con_griglia:
                pdf_fallback, _ = compila_pdf(latex_a)
                if pdf_fallback:
                    st.session_state.verifiche['A']['pdf']     = pdf_fallback
                    st.session_state.verifiche['A']['pdf_ts']  = time.time()
                    st.session_state.verifiche['A']['preview'] = True
                    st.warning("⚠️ La griglia di valutazione non è stata inclusa nel PDF.")

        # ── VERIFICA RIDOTTA BES/DSA ──────────────────────────────────────────
        if bes_dsa and perc_ridotta:
            _avanza("⛳ Generazione verifica ridotta…")

            prompt_ridotta = f"""Sei un docente esperto. Hai già generato questa verifica:

{corpo_latex}

Devi creare una versione RIDOTTA per studenti con sostegno o certificazione (BES/DSA/NAI).
La struttura deve essere simile all'originale ma con circa il {perc_ridotta}% di sottopunti IN MENO rispetto al totale.
Scegli quali sottopunti eliminare partendo dai più complessi. Mantieni sempre almeno 1 sottopunto per esercizio.
{'Ridistribuisci i punti in modo che la somma sia ESATTAMENTE ' + str(punti_totali) + ' pt. totali.' if mostra_punteggi else 'NON inserire punteggi.'}
NON aggiungere nessun simbolo (*), nessuna nota BES, nessuna indicazione che si tratta di una verifica ridotta.
TERMINA con \\end{{document}}.
SOLO CODICE LATEX del corpo (\\subsection* ecc.), senza preambolo."""

            rb_bes = model.generate_content(prompt_ridotta)
            corpo_latex_ridotta = rb_bes.text.replace("```latex", "").replace("```", "").strip()
            corpo_latex_ridotta = pulisci_corpo_latex(corpo_latex_ridotta)

            latex_ridotta = preambolo_fisso + corpo_latex_ridotta
            latex_ridotta = fix_items_environment(latex_ridotta)
            latex_ridotta = rimuovi_vspace_corpo(latex_ridotta)
            if mostra_punteggi:
                latex_ridotta = rimuovi_punti_subsection(latex_ridotta)
                latex_ridotta = riscala_punti(latex_ridotta, punti_totali)

            if con_griglia:
                latex_ridotta_final = inietta_griglia(latex_ridotta, punti_totali)
            else:
                latex_ridotta_final = latex_ridotta

            st.session_state.verifiche['R'] = {**_vf(), 'latex': latex_ridotta_final, 'latex_originale': latex_ridotta_final}

            pdf_r, err_r = compila_pdf(latex_ridotta_final)
            if pdf_r:
                st.session_state.verifiche['R']['pdf']    = pdf_r
                st.session_state.verifiche['R']['pdf_ts'] = time.time()
                st.session_state.verifiche['R']['preview'] = True
            else:
                if con_griglia:
                    pdf_r_fallback, _ = compila_pdf(latex_ridotta)
                    if pdf_r_fallback:
                        st.session_state.verifiche['R']['pdf']     = pdf_r_fallback
                        st.session_state.verifiche['R']['pdf_ts']  = time.time()
                        st.session_state.verifiche['R']['preview'] = True

        if doppia_fila:
            _avanza("📄  Generazione Versione B…")
            rb = model.generate_content(
                f"Versione B: stessa struttura, cambia dati e quesiti. "
                f"SOLO corpo esercizi (\\subsection* ecc.), SENZA preambolo/\\documentclass/\\begin{{document}}. "
                f"Sostituisci 'Versione A' con 'Versione B'. TERMINA con \\end{{document}}. SOLO LATEX.\n\n{corpo_latex}")
            corpo_latex_b = rb.text.replace("```latex","").replace("```","").strip()
            corpo_latex_b = pulisci_corpo_latex(corpo_latex_b)

            preambolo_b = preambolo_fisso.replace(
                titolo_header,
                titolo_header.replace("Versione A","Versione B") if "Versione A" in titolo_header
                else titolo_header + " — Versione B"
            )
            latex_b = preambolo_b + corpo_latex_b
            latex_b = fix_items_environment(latex_b)
            latex_b = rimuovi_vspace_corpo(latex_b)
            if mostra_punteggi:
                latex_b = rimuovi_punti_subsection(latex_b)
                latex_b = riscala_punti(latex_b, punti_totali)

            if con_griglia:
                latex_b_final = inietta_griglia(latex_b, punti_totali)
            else:
                latex_b_final = latex_b

            st.session_state.verifiche['B'] = {**_vf(), 'latex': latex_b_final, 'latex_originale': latex_b_final}

            _avanza("🖨️  PDF Versione B…")
            pdf_b_auto, _ = compila_pdf(latex_b_final)
            if pdf_b_auto:
                st.session_state.verifiche['B']['pdf']     = pdf_b_auto
                st.session_state.verifiche['B']['pdf_ts']  = time.time()
                st.session_state.verifiche['B']['preview'] = True
            else:
                if con_griglia:
                    pdf_b_fallback, _ = compila_pdf(latex_b)
                    if pdf_b_fallback:
                        st.session_state.verifiche['B']['pdf'] = pdf_b_fallback
                        st.session_state.verifiche['B']['pdf_ts'] = time.time()
                        st.session_state.verifiche['B']['preview'] = True

        # ── VERIFICA RIDOTTA FILA B ───────────────────────────────────────────
        if doppia_fila and bes_dsa and bes_dsa_b and perc_ridotta and st.session_state.verifiche['B']['latex']:
            _avanza("⛳ Generazione verifica ridotta Fila B…")
            prompt_ridotta_b = f"""Sei un docente esperto. Hai già generato questa verifica (Fila B):

{corpo_latex_b}

Crea una versione RIDOTTA per studenti con sostegno o certificazione (BES/DSA/NAI).
Stessa logica della ridotta A: rimuovi circa il {perc_ridotta}% dei sottopunti più complessi, mantieni almeno 1 sottopunto per esercizio.
{'Ridistribuisci i punti con somma ESATTAMENTE ' + str(punti_totali) + ' pt.' if mostra_punteggi else 'NON inserire punteggi.'}
NON aggiungere note BES o indicazioni che si tratta di versione ridotta.
SOLO CODICE LATEX del corpo (\\subsection* ecc.), senza preambolo. TERMINA con \\end{{document}}."""

            rb_bes_b = model.generate_content(prompt_ridotta_b)
            corpo_lr_b = rb_bes_b.text.replace("```latex","").replace("```","").strip()
            corpo_lr_b = pulisci_corpo_latex(corpo_lr_b)
            preambolo_rb = preambolo_b.replace(titolo_header + " — Versione B", titolo_header + " — Versione B Ridotta") if "Versione B" in preambolo_b else preambolo_b
            latex_ridotta_b = preambolo_rb + corpo_lr_b
            latex_ridotta_b = fix_items_environment(latex_ridotta_b)
            latex_ridotta_b = rimuovi_vspace_corpo(latex_ridotta_b)
            if mostra_punteggi:
                latex_ridotta_b = rimuovi_punti_subsection(latex_ridotta_b)
                latex_ridotta_b = riscala_punti(latex_ridotta_b, punti_totali)
            if con_griglia:
                latex_ridotta_b = inietta_griglia(latex_ridotta_b, punti_totali)
            st.session_state.verifiche['RB'] = {**_vf(), 'latex': latex_ridotta_b, 'latex_originale': latex_ridotta_b}
            pdf_rb, _ = compila_pdf(latex_ridotta_b)
            if pdf_rb:
                st.session_state.verifiche['RB']['pdf']    = pdf_rb
                st.session_state.verifiche['RB']['pdf_ts'] = time.time()
                st.session_state.verifiche['RB']['preview'] = True

        # ── SOLUZIONI ─────────────────────────────────────────────────────────
        if genera_soluzioni:
            _avanza("📋 Generazione soluzioni…")

            def _genera_testo_sol(corpo, versione_label=""):
                _v_tag = f" — {versione_label}" if versione_label else ""
                _prompt = f"""Sei un docente di {materia}. Fornisci le soluzioni SINTETICHE della seguente verifica{_v_tag}.

{corpo}

REGOLE FERREE — RISPETTALE ALLA LETTERA:
- Per ogni esercizio scrivi "Esercizio N: [Titolo]" poi le soluzioni in ordine a), b), c)...
- CALCOLI: mostra SOLO i passaggi essenziali. Niente testo narrativo. Solo la catena di calcolo.
- DOMANDE APERTE / TEORICHE: MASSIMO 3-4 RIGHE. Sii telegraficamente conciso. NON scrivere saggi.
- SCELTA MULTIPLA / VERO-FALSO: una riga sola: "Risposta: X — perché [motivazione breve]."
- NON riscrivere mai il testo della domanda originale, vai diretto alla soluzione.
- SE UN ESERCIZIO HA DATI INCOERENTI O ERRATI: scrivi "Dati da rivedere: [problema in una riga]" e passa avanti.
- Usa $...$ per le formule matematiche inline.
- Risposta totale per esercizio: MASSIMO 15-20 righe inclusi tutti i sottopunti.
- Rispondi con testo strutturato, senza preambolo LaTeX."""
                _rs = model.generate_content(_prompt)
                return _rs.text.strip()

            def _testo_to_latex_body(testo):
                body = ""
                for line in testo.split('\n'):
                    ls = line.strip()
                    if not ls:
                        body += "\n\\vspace{0.15cm}\n"
                    elif re.match(r'^#{1,3}\s', ls):
                        heading = re.sub(r'^#+\s*', '', ls)
                        body += f"\n\\subsection*{{{heading}}}\n"
                    elif re.match(r'^Esercizio\s+\d+', ls, re.IGNORECASE):
                        body += f"\n\\subsection*{{{ls}}}\n"
                    elif re.match(r'^[a-z]\)\s', ls):
                        body += f"\\noindent\\textbf{{{ls[:2]}}} {ls[2:].strip()}\n\n"
                    else:
                        body += ls + "\n"
                return body

            # Genera soluzioni fila A
            testo_sol_a = _genera_testo_sol(corpo_latex, "Fila A" if doppia_fila else "")

            # Genera soluzioni fila B (se esiste)
            testo_sol_b = None
            if doppia_fila and 'corpo_latex_b' in dir() and corpo_latex_b:
                testo_sol_b = _genera_testo_sol(corpo_latex_b, "Fila B")

            # Costruisci unico PDF soluzioni
            _titolo_sol = f"Soluzioni — {materia}: {titolo_clean}"
            latex_sol_body = ""

            if testo_sol_b:
                latex_sol_body += "\\section*{Fila A}\n"
            latex_sol_body += _testo_to_latex_body(testo_sol_a)

            if testo_sol_b:
                latex_sol_body += "\n\\newpage\n\\section*{Fila B}\n"
                latex_sol_body += _testo_to_latex_body(testo_sol_b)

            testo_sol_completo = testo_sol_a
            if testo_sol_b:
                testo_sol_completo += "\n\n---\n\n## Fila B\n\n" + testo_sol_b

            latex_sol = f"""\\documentclass[11pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[italian]{{babel}}
\\usepackage{{amsmath,amsfonts,amssymb,geometry}}
\\geometry{{margin=2cm}}
\\setlength{{\\parskip}}{{4pt}}
\\pagestyle{{empty}}
\\begin{{document}}
\\begin{{center}}
  \\textbf{{\\large {_titolo_sol}}} \\\\
  \\vspace{{0.2cm}}
  {{\\small \\textit{{Documento riservato al docente — non distribuire agli studenti}}}}
\\end{{center}}
\\vspace{{0.4cm}}
{latex_sol_body}
\\end{{document}}"""

            st.session_state.verifiche['S']['latex'] = latex_sol
            st.session_state.verifiche['S']['testo'] = testo_sol_completo
            pdf_sol, _ = compila_pdf(latex_sol)
            if pdf_sol:
                st.session_state.verifiche['S']['pdf']     = pdf_sol
                st.session_state.verifiche['S']['pdf_ts']  = time.time()
                st.session_state.verifiche['S']['preview'] = True

        _prog.markdown(f"""
<div style="margin:0.6rem 0 1rem 0;">
  <div style="font-size:0.82rem;font-weight:600;color:{T['success']};
              font-family:'DM Sans',sans-serif;margin-bottom:6px;">✅  Verifica pronta!</div>
  <div style="background:{T['border']};border-radius:100px;height:8px;overflow:hidden;">
    <div style="background:{T['success']};width:100%;height:100%;border-radius:100px;"></div>
  </div>
</div>
""", unsafe_allow_html=True)
        time.sleep(0.7)
        _prog.empty()
        st.session_state.last_materia   = materia
        st.session_state.last_argomento = titolo_clean
        st.session_state.last_gen_ts    = time.time()
        st.session_state._onboarding_done = True  # dismiss onboarding after first use

        # ── SALVA SU SUPABASE (con analytics) ────────────────────────────────
        try:
            if st.session_state.utente is not None:
                insert_data = {
                    "user_id":      st.session_state.utente.id,
                    "materia":      materia,
                    "argomento":    titolo_clean,
                    "scuola":       difficolta,
                    "latex_a":      st.session_state.verifiche['A']['latex'],
                    "latex_b":      st.session_state.verifiche['B']['latex'] if st.session_state.verifiche['B']['latex'] else None,
                    "latex_r":      st.session_state.verifiche['R']['latex'] if st.session_state.verifiche['R']['latex'] else None,
                    # ── NUOVO: campi analytics ────────────────────────────────
                    "modello":      modello_id,
                    "num_esercizi": num_esercizi_totali,
                }
                result = supabase_admin.table("verifiche_storico").insert(insert_data).execute()
                st.session_state._storico_refresh += 1
                st.toast("✅ Verifica salvata!", icon="💾")
            else:
                st.warning("Utente non loggato, verifica non salvata.")
        except Exception as e:
            st.warning(f"⚠️ Salvataggio non riuscito: {e}")

        st.rerun()

    except Exception as e:
        st.error(f"❌ Errore: {e}")

# ── OUTPUT ────────────────────────────────────────────────────────────────────────
if st.session_state.verifiche['A']['latex']:
    st.divider()
    _df  = doppia_fila   if 'doppia_fila'  in dir() else False
    _arg = st.session_state.last_argomento or (argomento if 'argomento' in dir() else 'verifica')

    attive = ['A','B'] if _df and st.session_state.verifiche['B']['latex'] else ['A']
    if st.session_state.verifiche['R']['latex']:
        attive.append('R')
    if st.session_state.verifiche['RB']['latex']:
        attive.append('RB')

    for idx, fid in enumerate(attive):
        v = st.session_state.verifiche[fid]
        if idx > 0:
            st.divider()
        with st.container():
            if fid == 'R':
                label_ver = "Verifica Ridotta (Fila A)"
            elif fid == 'RB':
                label_ver = "Verifica Ridotta (Fila B)"
            elif _df:
                label_ver = f"Versione {fid}"
            else:
                label_ver = "La tua verifica"
            st.markdown(f"""
            <div style="background:linear-gradient(135deg, {T['accent_light']} 0%, {T['card']} 100%);
                        border:2px solid {T['accent']};border-radius:16px;padding:0;
                        margin-bottom:1.8rem;overflow:hidden;
                        box-shadow:0 4px 20px {T['accent']}22;">
              <div style="background:{T['accent']};padding:1rem 1.3rem;
                          border-bottom:1px solid {T['accent']};">
                <div style="display:flex;align-items:center;gap:12px;">
                  <span style="font-size:1.8rem;">{APP_ICON}</span>
                  <div style="flex:1;">
                    <div style="font-family:'DM Sans',sans-serif;font-size:1.3rem;
                                font-weight:900;color:#ffffff;letter-spacing:-0.02em;">
                      {label_ver}
                    </div>
                    <div style="font-size:0.75rem;color:#ffffff;opacity:0.85;
                                font-weight:600;margin-top:2px;">
                      Generata con successo
                    </div>
                  </div>
                  <div style="background:#ffffff22;backdrop-filter:blur(10px);
                              border:1px solid #ffffff33;border-radius:20px;
                              padding:6px 16px;font-size:0.72rem;font-weight:700;
                              color:#ffffff;letter-spacing:0.05em;text-transform:uppercase;">
                    ✓ Pronta
                  </div>
                </div>
              </div>
              <div style="padding:1.2rem 1.3rem;background:{T['card']};">
                <div style="display:flex;align-items:flex-start;gap:10px;
                            background:{T['accent_light']};border-left:3px solid {T['accent']};
                            border-radius:8px;padding:12px 16px;">
                  <span style="font-size:1.1rem;flex-shrink:0;">⚠️</span>
                  <span style="font-size:0.82rem;color:{T['text2']};line-height:1.5;">
                    Le verifiche generate sono <strong style="color:{T['text']};">suggerimenti didattici</strong>. 
                    Rivedi sempre il contenuto prima della distribuzione — il docente è responsabile del materiale finale.
                  </span>
                </div>
              </div>
            </div>
            """, unsafe_allow_html=True)

            if v['pdf']:
                pdf_size = _stima_dimensione(v['pdf'])
                st.download_button(
                    label=f"📄 Scarica PDF — Alta qualità ({pdf_size})",
                    data=v['pdf'],
                    file_name=f"Verifica_{_arg}_{fid}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"dl_{fid}"
                )
            else:
                if st.button("📄 Genera PDF", key=f"dlc_{fid}", use_container_width=True):
                    with st.spinner("Compilazione…"):
                        pdf, err = compila_pdf(v['latex'])
                    if pdf:
                        st.session_state.verifiche[fid]['pdf'] = pdf
                        st.session_state.verifiche[fid]['pdf_ts'] = time.time()
                        st.rerun()
                    else:
                        st.error("Errore PDF")
                        with st.expander("Log"): st.text(err)

            st.write("")

            with st.expander("✏️ Modifica questa verifica", expanded=False):
                st.markdown(f"""
                <div style="font-size:0.85rem;color:{T['text2']};margin-bottom:0.8rem;line-height:1.5;">
                    Descrivi le modifiche che vuoi apportare. Esempi:<br>
                    • "Aggiungi un esercizio sulla proprietà distributiva"<br>
                    • "Rimuovi l'esercizio 3"<br>
                    • "Cambia i numeri dell'esercizio 2 con valori più piccoli"<br>
                    • "Aumenta la difficoltà dell'esercizio 4"
                </div>
                """, unsafe_allow_html=True)

                richiesta_modifica = st.text_area(
                    "Cosa vuoi modificare?",
                    height=100,
                    placeholder="es. Sostituisci l'esercizio 2 con un problema sulla velocità media",
                    key=f"modifica_input_{fid}",
                    label_visibility="collapsed"
                )

                col_mod1, col_mod2 = st.columns([1, 1])
                with col_mod1:
                    if st.button("🔄 Applica Modifiche", key=f"modifica_btn_{fid}",
                               use_container_width=True, disabled=not richiesta_modifica.strip()):
                        try:
                            with st.spinner("⏳ Modifica in corso..."):
                                model = genai.GenerativeModel(modello_id)
                                latex_modificato = modifica_verifica_con_ai(
                                    v['latex'],
                                    richiesta_modifica.strip(),
                                    model
                                )
                                latex_modificato = fix_items_environment(latex_modificato)
                                latex_modificato = rimuovi_vspace_corpo(latex_modificato)
                                if mostra_punteggi:
                                    latex_modificato = rimuovi_punti_subsection(latex_modificato)
                                    latex_modificato = riscala_punti(latex_modificato, punti_totali)
                                if con_griglia:
                                    latex_modificato = inietta_griglia(latex_modificato, punti_totali)
                                st.session_state.verifiche[fid]['latex'] = latex_modificato
                                pdf_mod, err_mod = compila_pdf(latex_modificato)
                                if pdf_mod:
                                    st.session_state.verifiche[fid]['pdf'] = pdf_mod
                                    st.session_state.verifiche[fid]['pdf_ts'] = time.time()
                                    st.session_state.verifiche[fid]['preview'] = True
                                    st.session_state.verifiche[fid]['docx'] = None
                                    st.success("✅ Modifiche applicate!")
                                    time.sleep(0.8)
                                    st.rerun()
                                else:
                                    st.error("❌ Errore nella compilazione del PDF modificato")
                                    if err_mod:
                                        with st.expander("Log errore"):
                                            st.text(err_mod)
                        except Exception as e:
                            st.error(f"❌ Errore durante la modifica: {str(e)}")

                with col_mod2:
                    ha_originale = 'latex_originale' in v and v.get('latex_originale')
                    if st.button("🗑️ Ripristina Originale", key=f"reset_mod_{fid}",
                               use_container_width=True,
                               disabled=not ha_originale,
                               help="Torna alla versione generata inizialmente"):
                        if ha_originale:
                            st.session_state.verifiche[fid]['latex'] = v['latex_originale']
                            pdf_orig, _ = compila_pdf(v['latex_originale'])
                            if pdf_orig:
                                st.session_state.verifiche[fid]['pdf'] = pdf_orig
                                st.session_state.verifiche[fid]['pdf_ts'] = time.time()
                                st.session_state.verifiche[fid]['docx'] = None
                            st.success("✅ Versione originale ripristinata!")
                            time.sleep(0.5)
                            st.rerun()

            st.write("")

            if v['docx']:
                docx_size = _stima_dimensione(v['docx'])
                st.download_button(
                    label=f"📝 Scarica File Word Modificabile ({docx_size})",
                    data=v['docx'],
                    file_name=f"Verifica_{_arg}_{fid}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key=f"dld_{fid}"
                )
                st.markdown(f"""
                    <div class="hint-docx" style="margin-top: -10px; margin-bottom: 15px;">
                        💡 <b>Nota:</b> La versione Word è modificabile ma ha una resa grafica inferiore.
                        I grafici complessi di funzione (TikZ/Plot) non compaiono in Word e vanno aggiunti a mano.
                    </div>
                """, unsafe_allow_html=True)
            else:
                _docx_gen_key = f"_docx_generating_{fid}"
                if st.button("📝 Genera File Word Modificabile", key=f"dldc_{fid}", use_container_width=True):
                    st.session_state[_docx_gen_key] = True
                if st.session_state.get(_docx_gen_key, False):
                    with st.spinner("⏳ Conversione Word…"):
                        db, de = latex_to_docx_via_ai(v['latex'], con_griglia=con_griglia)
                    if db:
                        st.session_state.verifiche[fid]['docx'] = db
                        st.session_state.verifiche[fid]['docx_ts'] = time.time()
                        st.session_state[_docx_gen_key] = False
                        st.rerun()
                    else:
                        st.session_state[_docx_gen_key] = False
                        st.error("Errore Word")
                        with st.expander("Log"): st.text(de)

            st.write("")

            if v['preview'] and v['pdf']:
                with st.expander("👁 Anteprima PDF", expanded=False):
                    b64 = base64.b64encode(v['pdf']).decode()
                    st.markdown(f"""
                    <iframe src="data:application/pdf;base64,{b64}#toolbar=0&navpanes=0&scrollbar=1"
                            style="width:100%;height:500px;border:none;border-radius:8px;display:block;"></iframe>
                    """, unsafe_allow_html=True)

            _spacer, _tex_col = st.columns([3, 1])
            with _tex_col:
                st.markdown('<div class="tex-btn-wrap">', unsafe_allow_html=True)
                st.download_button(
                    "⬇ Sorgente .tex",
                    data=v['latex'].encode('utf-8'),
                    file_name=f"Verifica_{_arg}_{fid}.tex",
                    mime="text/plain",
                    key=f"dl_tex_{fid}",
                    help="Scarica il sorgente LaTeX per modificarlo"
                )
                st.markdown('</div>', unsafe_allow_html=True)

# ── SOLUZIONI ────────────────────────────────────────────────────────────────────
if st.session_state.verifiche['S'].get('testo') or st.session_state.verifiche['S'].get('pdf'):
    st.divider()
    _arg_s = st.session_state.last_argomento or (argomento if 'argomento' in dir() else 'verifica')
    v_s = st.session_state.verifiche['S']
    st.markdown(f"""
    <div style="background:linear-gradient(135deg, {T['accent_light']} 0%, {T['card']} 100%);
                border:2px solid {T['success']};border-radius:16px;padding:0;
                margin-bottom:1.8rem;overflow:hidden;box-shadow:0 4px 20px {T['success']}22;">
      <div style="background:{T['success']};padding:1rem 1.3rem;">
        <div style="display:flex;align-items:center;gap:12px;">
          <span style="font-size:1.8rem;">📋</span>
          <div style="flex:1;">
            <div style="font-family:'DM Sans',sans-serif;font-size:1.3rem;font-weight:900;color:#ffffff;letter-spacing:-0.02em;">
              Soluzioni
            </div>
            <div style="font-size:0.75rem;color:#ffffff;opacity:0.85;font-weight:600;margin-top:2px;">
              Documento riservato al docente
            </div>
          </div>
          <div style="background:#ffffff22;border:1px solid #ffffff33;border-radius:20px;
                      padding:6px 16px;font-size:0.72rem;font-weight:700;color:#ffffff;
                      letter-spacing:0.05em;text-transform:uppercase;">🔒 Solo docente</div>
        </div>
      </div>
    </div>
    """, unsafe_allow_html=True)

    if v_s.get('pdf'):
        st.download_button(
            label="📄 Scarica Soluzioni PDF",
            data=v_s['pdf'],
            file_name=f"Soluzioni_{_arg_s}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="dl_sol_pdf"
        )
    st.write("")
    if v_s.get('testo'):
        with st.expander("👁 Mostra soluzioni", expanded=False):
            st.markdown(v_s['testo'])
    if v_s.get('pdf') and v_s.get('preview'):
        with st.expander("👁 Anteprima PDF Soluzioni", expanded=False):
            b64_s = base64.b64encode(v_s['pdf']).decode()
            st.markdown(f"""
            <iframe src="data:application/pdf;base64,{b64_s}#toolbar=0&navpanes=0&scrollbar=1"
                    style="width:100%;height:500px;border:none;border-radius:8px;display:block;"></iframe>
            """, unsafe_allow_html=True)

# ── FOOTER ───────────────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="app-footer">
  ⚠️ Le verifiche generate dall'AI sono suggerimenti didattici — rivedi sempre il contenuto
  prima di distribuirlo agli studenti. Il docente è responsabile del materiale finale.<br>
  <span style="opacity:0.55;">VerificAI · Versione Beta</span>
</div>
""", unsafe_allow_html=True)

import streamlit.components.v1 as components
components.html(f"""
<style>
  body {{ margin: 0; padding: 0; background: transparent; }}
  #share-btn {{
    background: none;
    border: none;
    cursor: pointer;
    color: {T['accent']};
    font-weight: 600;
    font-size: 0.72rem;
    font-family: 'DM Sans', sans-serif;
    padding: 0;
    display: block;
    margin: 0 auto;
    text-align: center;
    width: 100%;
  }}
  #share-btn:hover {{ text-decoration: underline; }}
</style>
<button id="share-btn" onclick="copyLink()">🔗 Condividi con i colleghi</button>
<script>
function copyLink() {{
  var url = "{SHARE_URL}";
  var btn = document.getElementById("share-btn");
  var ta = document.createElement("textarea");
  ta.value = url;
  ta.style.cssText = "position:fixed;top:0;left:0;opacity:0;";
  document.body.appendChild(ta);
  ta.focus();
  ta.select();
  var ok = false;
  try {{ ok = document.execCommand("copy"); }} catch(e) {{}}
  document.body.removeChild(ta);
  if (ok) {{
    btn.innerText = "✅ Link copiato!";
    setTimeout(function() {{ btn.innerText = "🔗 Condividi con i colleghi"; }}, 2000);
  }} else {{
    btn.innerText = url;
  }}
}}
</script>
""", height=30)










