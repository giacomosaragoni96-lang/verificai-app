# ── docx_export.py ─────────────────────────────────────────────────────────────
# Conversione da LaTeX a DOCX (Word).
# Nessuna dipendenza da Streamlit — importabile e testabile in isolamento.
#
# Import in app.py:
#   from docx_export import latex_to_docx_via_ai
# ───────────────────────────────────────────────────────────────────────────────

import io
import os
import re
import subprocess
import tempfile
import logging

# Configurazione logging base per vedere output nel terminale
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S'
)

from latex_utils import parse_esercizi

# Setup logger per debugging grafici
logger = logging.getLogger(__name__)


# ── GRAFICO: TikZ/pgfplots → PNG ──────────────────────────────────────────────

_GRAPH_MARKER = "__DOCXGRAPH_{n}__"
_GRAPH_RE = re.compile(r'__DOCXGRAPH_(\d+)__')


def _clean_tikz_code(tikz_code: str) -> str:
    """
    Rimuove etichette automatiche e titoli non richiesti dal codice TikZ.
    Mantiene il codice "puro" per la conversione in PNG.
    """
    logger.debug("Cleaning TikZ code - removing automatic labels...")
    
    # Rimuovi comandi \label{} che potrebbero generare etichette automatiche
    cleaned = re.sub(r'\\label\{[^}]*\}', '', tikz_code)
    
    # Rimuovi \caption{} e titoli automatici
    cleaned = re.sub(r'\\caption\{[^}]*\}', '', cleaned)
    
    # Rimuovi \node con testo che assomiglia a etichette automatiche (Figure, etc.)
    cleaned = re.sub(r'\\node\s*[^(]*\([^)]*\)\s*\{[^F]*(Figure|Fig|Grafico|Diagramma)[^}]*\}', '', cleaned, flags=re.IGNORECASE)
    
    # Rimuovi commenti che potrebbero contenere riferimenti a etichette
    cleaned = re.sub(r'%.*[Ff]igure.*', '', cleaned)
    cleaned = re.sub(r'%.*[Gg]rafico.*', '', cleaned)
    
    # Rimuovi spazi multipli e righe vuote eccedenti
    cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned)
    cleaned = cleaned.strip()
    
    if cleaned != tikz_code:
        logger.debug("TikZ code cleaned - automatic labels removed")
    
    return cleaned


def _tikz_to_png_bytes(tikz_code: str) -> bytes | None:
    """
    Compiles a TikZ/pgfplots block to a PNG image.
    Conversion chain: PyMuPDF → pdf2image → pdftoppm CLI.
    Returns None on total failure.
    """
    logger.info(f"TikZ Conversion Attempt - Code length: {len(tikz_code)} chars")
    logger.debug(f"TikZ Code Preview: {tikz_code[:200]}...")
    
    # Pulisci il codice TikZ da etichette automatiche
    tikz_code = _clean_tikz_code(tikz_code)
    logger.debug(f"TikZ Code after cleaning: {tikz_code[:200]}...")
    
    needs_pgfplots = (
        r'\begin{axis}' in tikz_code
        or 'pgfplots' in tikz_code
        or 'addplot' in tikz_code
    )
    logger.info(f"Requires pgfplots: {needs_pgfplots}")
    
    latex = (
        r'\documentclass[border=4pt]{standalone}' '\n'
        r'\usepackage{tikz}' '\n'
    )
    if needs_pgfplots:
        latex += r'\usepackage{pgfplots}' '\n'
        latex += r'\pgfplotsset{compat=1.18}' '\n'
    latex += r'\begin{document}' '\n' + tikz_code + '\n' r'\end{document}'
    
    logger.debug(f"Generated LaTeX document: {len(latex)} chars")

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            logger.info(f"Using temp directory: {tmpdir}")
            tex_path = os.path.join(tmpdir, 'g.tex')
            pdf_path = os.path.join(tmpdir, 'g.pdf')
            
            with open(tex_path, 'w', encoding='utf-8') as f:
                f.write(latex)
            
            logger.info("Running pdflatex compilation...")
            result = subprocess.run(
                ['pdflatex', '-interaction=nonstopmode',
                 '-output-directory', tmpdir, tex_path],
                capture_output=True, timeout=30,
            )
            
            if result.returncode != 0:
                logger.error(f"pdflatex failed - Return code: {result.returncode}")
                logger.error(f"pdflatex stderr: {result.stderr.decode() if result.stderr else 'None'}")
                logger.error(f"pdflatex stdout: {result.stdout.decode() if result.stdout else 'None'}")
                return None
            
            if not os.path.exists(pdf_path):
                logger.error(f"PDF file not created at {pdf_path}")
                return None
            
            logger.info("PDF compilation successful")

            # ── Method 1: PyMuPDF — self-contained, works on Windows ──────────
            try:
                logger.info("Attempting PyMuPDF conversion...")
                import fitz  # PyMuPDF
                fitz_doc = fitz.open(pdf_path)
                page = fitz_doc[0]
                # 2.5× zoom ≈ 180 DPI (base 72 dpi)
                mat  = fitz.Matrix(2.5, 2.5)
                pix  = page.get_pixmap(matrix=mat, alpha=False)
                fitz_doc.close()
                png_bytes = pix.tobytes('png')
                logger.info(f"PyMuPDF conversion successful - PNG size: {len(png_bytes)} bytes")
                return png_bytes
            except Exception as e:
                logger.warning(f"PyMuPDF conversion failed: {e}")
                pass

            # ── Method 2: pdf2image (needs poppler in PATH) ───────────────────
            try:
                logger.info("Attempting pdf2image conversion...")
                from pdf2image import convert_from_path
                pages = convert_from_path(pdf_path, dpi=180)
                if pages:
                    buf = io.BytesIO()
                    pages[0].save(buf, 'PNG')
                    png_bytes = buf.getvalue()
                    logger.info(f"pdf2image conversion successful - PNG size: {len(png_bytes)} bytes")
                    return png_bytes
            except Exception as e:
                logger.warning(f"pdf2image conversion failed: {e}")
                pass

            # ── Method 3: pdftoppm CLI ────────────────────────────────────────
            try:
                logger.info("Attempting pdftoppm CLI conversion...")
                out_base = os.path.join(tmpdir, 'out')
                result = subprocess.run(
                    ['pdftoppm', '-png', '-r', '180', '-singlefile',
                     pdf_path, out_base],
                    capture_output=True, timeout=15,
                )
                
                if result.returncode != 0:
                    logger.error(f"pdftoppm failed - Return code: {result.returncode}")
                    logger.error(f"pdftoppm stderr: {result.stderr.decode() if result.stderr else 'None'}")
                
                out_path = out_base + '.png'
                if os.path.exists(out_path):
                    with open(out_path, 'rb') as f:
                        png_bytes = f.read()
                    logger.info(f"pdftoppm conversion successful - PNG size: {len(png_bytes)} bytes")
                    return png_bytes
                else:
                    logger.error(f"pdftoppm: PNG file not created at {out_path}")
            except Exception as e:
                logger.warning(f"pdftoppm conversion failed: {e}")
                pass
    except Exception as e:
        logger.error(f"Complete TikZ conversion failure: {e}")
        pass
    
    logger.error("All TikZ conversion methods failed")
    return None


def _pre_render_graphs(latex: str) -> tuple[str, dict]:
    """
    Extracts TikZ/pgfplots blocks from latex, renders each to PNG, and
    replaces the block with a __DOCXGRAPH_N__ placeholder.
    Returns (modified_latex, {N: png_bytes_or_None}).
    """
    logger.info("Starting pre-render of graphs...")
    graphs: dict[int, bytes | None] = {}
    counter = [0]

    def _replace(m: re.Match) -> str:
        n = counter[0]
        counter[0] += 1
        tikz_code = m.group(0)
        logger.info(f"Processing graph #{n} - Length: {len(tikz_code)} chars")
        
        png_bytes = _tikz_to_png_bytes(tikz_code)
        graphs[n] = png_bytes
        
        if png_bytes:
            logger.info(f"Graph #{n} conversion SUCCESS - {len(png_bytes)} bytes")
        else:
            logger.warning(f"Graph #{n} conversion FAILED - will show placeholder")
        
        # Aggiungi \n\ prima e dopo per forzare a capo nel LaTeX
        return f'\n\n__DOCXGRAPH_{n}__\n\n'

    modified = re.sub(
        r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
        _replace, latex, flags=re.DOTALL,
    )
    modified = re.sub(
        r'\\begin\{axis\}.*?\\end\{axis\}',
        _replace, modified, flags=re.DOTALL,
    )
    
    total_graphs = counter[0]
    successful = sum(1 for png in graphs.values() if png is not None)
    logger.info(f"Pre-render complete: {successful}/{total_graphs} graphs converted successfully")
    
    return modified, graphs


def _add_content_with_graphs(
    doc,
    text: str,
    graphs: dict,
    *,
    left_indent_cm: float = 0.0,
    space_before_pt: float = 0.0,
    space_after_pt: float = 4.0,
    base_size_pt: int = 11,
    img_width_cm: float = 13.0,
) -> None:
    """
    Splits `text` on __DOCXGRAPH_N__ markers.
    Text parts are added as formatted paragraphs; graph markers become centered images.
    """
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Debug visivo: controlla se i placeholder sono presenti
    logger.info(f"Testo esercizio prima dell'inserimento: {text[:100]}...")
    placeholder_count = len(_GRAPH_RE.findall(text))
    logger.info(f"Placeholder __DOCXGRAPH__ trovati nel testo: {placeholder_count}")
    
    # Fallimento esplicito: se ci sono grafici ma nessun placeholder
    if graphs and placeholder_count == 0:
        logger.error("Placeholder non trovato nel testo, l'immagine non può essere inserita!")
        logger.error(f"Grafici disponibili: {list(graphs.keys())}")
    
    parts = _GRAPH_RE.split(text)
    logger.info(f"Testo diviso in {len(parts)} parti: {parts}")
    first = True
    for i, part in enumerate(parts):
        logger.debug(f"Processing part {i}: '{part}'")
        # Dopo lo split, i numeri dei grafici sono nelle posizioni dispari
        if i % 2 == 1:  # Posizioni dispari contengono i numeri dei grafici
            try:
                n = int(part)
                png = graphs.get(n)
                
                logger.info(f"Processando placeholder per grafico #{n}")
                
                # Aggiungi un paragrafo vuoto per separazione dal testo
                p_separator = doc.add_paragraph()
                p_separator.paragraph_format.space_before = Pt(6)
                p_separator.paragraph_format.space_after = Pt(6)
                
                # Forza il paragrafo: crea nuovo spazio dedicato per l'immagine
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img.paragraph_format.space_before = Pt(12)  # Spazio maggiore per separazione
                p_img.paragraph_format.space_after  = Pt(12)
                p_img.paragraph_format.left_indent = Cm(0.0)
                
                if png:
                    try:
                        run = p_img.add_run()
                        # Usa Inches per migliore controllo dimensioni
                        run.add_picture(io.BytesIO(png), width=Inches(4))
                        logger.info(f"Final Step: Image {n} actually appended to Word document - {len(png)} bytes")
                        logger.debug(f"Graph #{n} successfully added to DOCX - centered layout, Inches(4)")
                    except Exception as e:
                        logger.error(f"Failed to insert graph #{n} into DOCX: {e}")
                        p_img.add_run('⚠️ Grafico non disponibile - errore inserimento').italic = True
                else:
                    logger.warning(f"Graph #{n} placeholder - conversion failed")
                    p_img.add_run('⚠️ Grafico non disponibile - conversione fallita').italic = True
                
                # Aggiungi un altro paragrafo vuoto dopo il grafico
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_before = Pt(6)
                p_after.paragraph_format.space_after = Pt(6)
                
                first = False
            except ValueError:
                logger.warning(f"Parte '{part}' non è un numero di grafico valido")
                continue
        else:
            # Questo è testo normale
            stripped = part.strip()
            if not stripped:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(left_indent_cm)
            p.paragraph_format.space_before = Pt(space_before_pt if first else 0)
            p.paragraph_format.space_after  = Pt(space_after_pt)
            _add_rich_runs(p, stripped, base_size_pt)
            first = False
    
    # Log finale per verifica completa
    processed_graphs = sum(1 for i, part in enumerate(parts) if i % 2 == 1 and part.strip().isdigit())
    logger.info(f"Elaborazione completata: {processed_graphs} placeholder processati su {len(graphs)} grafici disponibili")


# ── HELPER XML/DOCX INTERNI ────────────────────────────────────────────────────

def _make_tc_xml(text, width_dxa, bold=False, gridSpan=1):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE

    tc = _OE('w:tc')
    tcPr = _OE('w:tcPr')
    tc.append(tcPr)

    if gridSpan > 1:
        gs = _OE('w:gridSpan')
        gs.set(_qn('w:val'), str(gridSpan))
        tcPr.append(gs)

    tcW = _OE('w:tcW')
    tcW.set(_qn('w:w'), str(width_dxa))
    tcW.set(_qn('w:type'), 'dxa')
    tcPr.append(tcW)

    p = _OE('w:p')
    tc.append(p)
    pPr = _OE('w:pPr')
    p.append(pPr)
    jc = _OE('w:jc')
    jc.set(_qn('w:val'), 'center')
    pPr.append(jc)

    if text:
        r_el = _OE('w:r')
        p.append(r_el)
        if bold:
            rPr = _OE('w:rPr')
            b_el = _OE('w:b')
            rPr.append(b_el)
            r_el.append(rPr)
        t = _OE('w:t')
        t.text = str(text)
        r_el.append(t)

    return tc


def _fix_tbl_grid(tbl_el, col_widths, _qn, _OE):
    old_grid = tbl_el.find(_qn('w:tblGrid'))
    if old_grid is not None:
        tbl_el.remove(old_grid)

    new_grid = _OE('w:tblGrid')
    for w in col_widths:
        gc = _OE('w:gridCol')
        gc.set(_qn('w:w'), str(max(1, w)))
        new_grid.append(gc)

    tbl_pr = tbl_el.find(_qn('w:tblPr'))
    if tbl_pr is not None:
        tbl_pr.addnext(new_grid)
    else:
        tbl_el.insert(0, new_grid)


def _build_griglia_xml(doc, row_es, row_sotto, row_max, PAGE_W_DXA=9638):
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    n_cols    = len(row_es)
    row_punti = ["Punti"] + [""] * (n_cols - 1)
    row_voto  = ["Voto"] + ["\\hrulefill"] * (n_cols - 1)  # Add Voto row with fill lines

    first_col = 1400
    last_col  = 900
    n_mid     = max(1, n_cols - 2)
    available = PAGE_W_DXA - first_col - last_col
    mid_col   = max(400, available // n_mid)
    col_widths = [first_col] + [mid_col] * n_mid + [last_col]
    diff = PAGE_W_DXA - sum(col_widths)
    col_widths[-1] += diff

    def _setup_tbl(tbl, total_w, widths):
        tbl_el = tbl._tbl
        tbl_pr = tbl_el.find(_qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = _OE('w:tblPr')
            tbl_el.insert(0, tbl_pr)

        tbl_w = _OE('w:tblW')
        tbl_w.set(_qn('w:w'), str(total_w))
        tbl_w.set(_qn('w:type'), 'dxa')
        ex = tbl_pr.find(_qn('w:tblW'))
        if ex is not None:
            tbl_pr.remove(ex)
        tbl_pr.append(tbl_w)

        tbl_lay = _OE('w:tblLayout')
        tbl_lay.set(_qn('w:type'), 'fixed')
        ex2 = tbl_pr.find(_qn('w:tblLayout'))
        if ex2 is not None:
            tbl_pr.remove(ex2)
        tbl_pr.append(tbl_lay)

        tbl_cm = _OE('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            cm_el = _OE(f'w:{side}')
            cm_el.set(_qn('w:w'), '50')
            cm_el.set(_qn('w:type'), 'dxa')
            tbl_cm.append(cm_el)
        ex3 = tbl_pr.find(_qn('w:tblCellMar'))
        if ex3 is not None:
            tbl_pr.remove(ex3)
        tbl_pr.append(tbl_cm)

        _fix_tbl_grid(tbl_el, widths, _qn, _OE)

    def _fill_cell(cell, text, bold=False, w=None, font_pt=9):
        cell.text = str(text)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(font_pt)
            if bold:
                run.bold = True
            # Special handling for Voto row - add underline
            if text == "\\hrulefill":
                run.underline = True
                # Clear the text and just keep the underline
                run.text = "           "  # Space for manual writing
        if w is not None:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = _OE('w:tcW')
            tcW.set(_qn('w:w'), str(w))
            tcW.set(_qn('w:type'), 'dxa')
            ex = tcPr.find(_qn('w:tcW'))
            if ex is not None:
                tcPr.remove(ex)
            tcPr.append(tcW)

    tbl = doc.add_table(rows=4, cols=n_cols)
    tbl.style = 'Table Grid'
    _setup_tbl(tbl, PAGE_W_DXA, col_widths)

    for r_idx, riga in enumerate([row_sotto, row_max, row_punti, row_voto], start=1):
        for c_idx in range(n_cols):
            val = riga[c_idx] if c_idx < len(riga) else ''
            _fill_cell(
                tbl.cell(r_idx, c_idx), val,
                bold=(c_idx == 0),
                w=col_widths[c_idx] if c_idx < len(col_widths) else mid_col
            )

    groups = []
    c = 0
    while c < n_cols:
        val = row_es[c] if c < len(row_es) else ''
        if c == 0 or c == n_cols - 1:
            groups.append((val, 1, col_widths[c]))
            c += 1
        else:
            span = 1
            while (c + span < n_cols - 1 and
                   c + span < len(row_es) and
                   row_es[c + span] == val):
                span += 1
            total_w = sum(col_widths[c:c + span])
            groups.append((val, span, total_w))
            c += span

    tr_el = tbl.rows[0]._tr
    for tc_el in list(tr_el.findall(_qn('w:tc'))):
        tr_el.remove(tc_el)
    for label, span, width in groups:
        tr_el.append(_make_tc_xml(label, width, bold=True, gridSpan=span))

    return tbl


# ── HELPER FUNZIONI DOCX ───────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Helper python-docx: imposta colore sfondo cella."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for _s in tcPr.findall(qn('w:shd')):
        tcPr.remove(_s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=80, right=80) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


# ── CONVERSIONE TESTO LATEX → TESTO PLAIN ──────────────────────────────────────

def _strip_latex_math(text: str) -> str:
    if not text:
        return text

    text = re.sub(r'\$\$(.+?)\$\$', lambda m: m.group(1).strip(), text, flags=re.DOTALL)
    text = re.sub(r'\$(.+?)\$',     lambda m: m.group(1).strip(), text)
    text = re.sub(r'\\mathcal\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathbf\{([^}]+)\}',  r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]+)\}',  r'\1', text)
    text = re.sub(r'\\text\{([^}]+)\}',    r'\1', text)
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', text)
    text = re.sub(r'\\sqrt\{([^}]+)\}',    r'sqrt(\1)', text)
    text = re.sub(r'\^\{([^}]+)\}',         r'^\1', text)
    text = re.sub(r'_\{([^}]+)\}',          r'_\1', text)

    replacements = [
        ("\\\\leq", "\u2264"), ("\\\\geq", "\u2265"), ("\\\\neq", "\u2260"),
        ("\\\\approx", "\u2248"), ("\\\\cdot", "\u00b7"), ("\\\\times", "\u00d7"),
        ("\\\\pm", "\u00b1"), ("\\\\infty", "\u221e"), ("\\\\alpha", "\u03b1"),
        ("\\\\beta", "\u03b2"), ("\\\\gamma", "\u03b3"), ("\\\\delta", "\u03b4"),
        ("\\\\theta", "\u03b8"), ("\\\\pi", "\u03c0"), ("\\\\sin", "sin"),
        ("\\\\cos", "cos"), ("\\\\tan", "tan"), ("\\\\log", "log"),
        ("\\\\ln", "ln"), ("\\\\lim", "lim"), ("\\\\forall", "\u2200"),
        ("\\\\exists", "\u2203"), ("\\\\in", "\u2208"), ("\\\\notin", "\u2209"),
        ("\\\\subset", "\u2282"), ("\\\\cup", "\u222a"), ("\\\\cap", "\u2229"),
        ("\\\\emptyset", "\u2205"), ("^2", "\u00b2"), ("^3", "\u00b3"),
        ("\\\\left(", "("), ("\\\\right)", ")"),
        ("\\\\left[", "["), ("\\\\right]", "]"),
    ]
    for fr, to in replacements:
        text = text.replace(fr, to)

    text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    text = text.replace('{', '').replace('}', '')
    return text.strip()


def _clean_latex_line(text: str) -> str:
    """Strip LaTeX → plain text (formatting lost). Used for headers and parsing."""
    if not text:
        return ''

    text = re.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
                  '[Figura]', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{axis\}.*?\\end\{axis\}',
                  '[Grafico]', text, flags=re.DOTALL)
    text = re.sub(r'\\vspace\*?\{[^}]*\}', '', text)
    text = re.sub(r'\\hspace\*?\{[^}]*\}', '', text)
    text = re.sub(r'\\noindent\b', '', text)
    text = re.sub(r'\\newline\b', '', text)
    text = re.sub(r'\\\\(\s)', r'\1', text)
    text = re.sub(r'\\\\\s*$', '', text, flags=re.MULTILINE)

    for cmd in ('textbf', 'textit', 'emph', 'underline', 'textrm', 'texttt'):
        text = re.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', text)

    text = re.sub(
        r'\\(?:small|large|Large|LARGE|huge|Huge|normalsize|footnotesize)\b', '', text
    )
    return _strip_latex_math(text).strip()


def _add_rich_runs(paragraph, text: str, base_size_pt: int = 11) -> None:
    """
    Aggiunge runs al paragrafo preservando \\textbf{} → bold e \\textit{} → italic.
    Math inline ($...$) viene preservata come testo. TikZ → [Figura].
    """
    from docx.shared import Pt

    # Rimuovi TikZ e grafico prima
    text = re.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
                  ' [Figura] ', text, flags=re.DOTALL)
    text = re.sub(r'\\begin\{axis\}.*?\\end\{axis\}',
                  ' [Grafico] ', text, flags=re.DOTALL)

    # Tokenizza: cerca \textbf{...} e \textit{...}/\emph{...}
    # Token: ('bold', content), ('italic', content), ('plain', content)
    tokens: list[tuple[str, str]] = []
    pattern = re.compile(
        r'\\textbf\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\(?:textit|emph)\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
        r'|\\underline\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}'
    )
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            tokens.append(('plain', text[last_end:m.start()]))
        if m.group(1) is not None:
            tokens.append(('bold', m.group(1)))
        elif m.group(2) is not None:
            tokens.append(('italic', m.group(2)))
        else:
            tokens.append(('underline', m.group(3)))
        last_end = m.end()
    if last_end < len(text):
        tokens.append(('plain', text[last_end:]))

    for kind, content in tokens:
        # Pulisci contenuto
        content = _strip_latex_math(content)
        content = re.sub(r'\\[a-zA-Z]+\s*', '', content)
        content = content.replace('{', '').replace('}', '').strip()
        if not content:
            continue
        run = paragraph.add_run(content)
        run.font.size = Pt(base_size_pt)
        if kind == 'bold':
            run.bold = True
        elif kind == 'italic':
            run.italic = True
        elif kind == 'underline':
            run.underline = True


def _estrai_punti(text: str) -> str:
    m = re.search(r'\((\d+(?:[.,]\d+)?)\s*pt\)', text)
    return m.group(1) if m else ''


# ── PARSING LATEX → STRUTTURA DATI ────────────────────────────────────────────

def _parse_latex_to_data(codice_latex: str) -> dict:
    LETTERE = 'abcdefghijklmnopqrstuvwxyz'
    ROMANI  = ['i','ii','iii','iv','v','vi','vii','viii','ix','x',
               'xi','xii','xiii','xiv','xv','xvi','xvii','xviii','xix','xx']

    data = {'titolo': '', 'intestazione_nota': '', 'esercizi': []}

    m = re.search(r'\\textbf\{\\large ([^}]+)\}', codice_latex)
    if m:
        data['titolo'] = _clean_latex_line(m.group(1))

    m2 = re.search(r'\\textit\{\\small ([^}]+)\}', codice_latex)
    if m2:
        data['intestazione_nota'] = m2.group(1).strip()

    body_start = codice_latex.find('\\end{center}')
    body_start = (body_start + len('\\end{center}')) if body_start != -1 else 0
    corpus = codice_latex[body_start:].replace('\\end{document}', '')

    blocks = re.split(r'\\subsection\*\s*\{', corpus)

    for block in blocks[1:]:
        brace_depth = 0
        header_end  = 0
        for ci, ch in enumerate(block):
            if ch == '{':
                brace_depth += 1
            elif ch == '}':
                if brace_depth == 0:
                    header_end = ci
                    break
                brace_depth -= 1

        titolo_ex = _clean_latex_line(block[:header_end])
        body = block[header_end + 1:]

        body = re.sub(r'\\begin\{tikzpicture\}.*?\\end\{tikzpicture\}',
                      '\n[Grafico]\n', body, flags=re.DOTALL)
        body = re.sub(r'\\begin\{axis\}.*?\\end\{axis\}',
                      '\n[Grafico]\n', body, flags=re.DOTALL)

        first_env = len(body)
        for marker in [r'\begin{enumerate}', r'\begin{itemize}', r'\item']:
            idx = body.find(marker)
            if idx != -1 and idx < first_env:
                first_env = idx

        raw_intro = body[:first_env]
        raw_intro = re.sub(r'\\begin\{tabular\}.*?\\end\{tabular\}',
                           '', raw_intro, flags=re.DOTALL)
        raw_intro = re.sub(r'\\begin\{center\}.*?\\end\{center\}',
                           '', raw_intro, flags=re.DOTALL)
        testo_intro = _clean_latex_line(raw_intro)

        sottopunti = []

        def _parse_items(items_block, label_style='alpha'):
            risultati = []
            pat = re.compile(
                r'\\item(?:\[([^\]]*)\])?\s*(.*?)(?=\\item(?:\[|\s)|$)',
                re.DOTALL
            )
            auto_idx = 0
            for im in pat.finditer(items_block):
                explicit = im.group(1)
                if explicit is not None:
                    label = explicit.strip()
                    auto_idx += 1
                else:
                    if label_style == 'roman':
                        label = ROMANI[min(auto_idx, len(ROMANI)-1)] + ')'
                    else:
                        label = LETTERE[auto_idx % 26] + ')'
                    auto_idx += 1

                raw_text = im.group(2).strip()
                if not label and not raw_text:
                    continue

                opzioni = []

                inner = re.search(
                    r'\\begin\{enumerate\}\s*\[([^\]]*)\]\s*(.*?)\\end\{enumerate\}',
                    raw_text, re.DOTALL
                )
                if not inner:
                    inner = re.search(
                        r'\\begin\{enumerate\}\s*(.*?)\\end\{enumerate\}',
                        raw_text, re.DOTALL
                    )
                if inner:
                    inner_content = inner.group(inner.lastindex)
                    inner_label   = inner.group(1) if inner.lastindex > 1 else 'a)'
                    inner_style   = 'roman' if 'i' in inner_label else 'alpha'
                    for sub_label, sub_testo, _, _ in _parse_items(inner_content, inner_style):
                        opt_c = (f"{sub_label} {sub_testo}").strip()
                        if opt_c:
                            opzioni.append(opt_c)
                    raw_text = raw_text[:inner.start()].strip()

                vf_pairs = re.findall(r'\$\\square\$\s*\\textbf\{([VF])\}', raw_text)
                if vf_pairs:
                    opzioni = [f"☐ {v}" for v in vf_pairs]
                    raw_text = re.sub(
                        r'\$\\square\$\s*\\textbf\{[VF]\}\s*(?:\\quad)?', '', raw_text
                    ).strip()

                testo_clean = _clean_latex_line(raw_text)
                testo_clean = re.sub(r'\n{2,}', '\n', testo_clean).strip()
                punti = _estrai_punti(label + ' ' + testo_clean)
                risultati.append((label, testo_clean, opzioni, punti))

            return risultati

        env_pat = re.compile(
            r'\\begin\{(enumerate|itemize)\}(\s*\[[^\]]*\])?\s*(.*?)\\end\{(?:enumerate|itemize)\}',
            re.DOTALL
        )
        used_ranges = []
        for em in env_pat.finditer(body):
            used_ranges.append((em.start(), em.end()))
            opt_arg     = (em.group(2) or '').strip()
            items_block = em.group(3)
            label_style = 'roman' if opt_arg.startswith('[i') else 'alpha'

            for label, testo_clean, opzioni, punti in _parse_items(items_block, label_style):
                sottopunti.append({
                    'label':   label,
                    'testo':   testo_clean,
                    'opzioni': opzioni,
                    'punti':   punti,
                })

        if not sottopunti:
            # Try to find sub-points in format a) b) c) without \item
            simple_pattern = re.compile(
                r'([a-z])\)\s*([^a-z]*?)(?=\s*[a-z]\)|$)',
                re.DOTALL
            )
            simple_matches = simple_pattern.findall(body)
            if simple_matches:
                logger.info(f"Found {len(simple_matches)} simple sub-points in a) b) c) format")
                for letter, text in simple_matches:
                    testo_clean = _clean_latex_line(text.strip())
                    punti = _estrai_punti(letter + ') ' + testo_clean)
                    sottopunti.append({
                        'label': letter + ')',
                        'testo': testo_clean,
                        'opzioni': [],
                        'punti': punti,
                    })
            else:
                # Try another pattern for different formats
                alt_pattern = re.compile(
                    r'([a-z])\)\s*([^)]+?)(?=\s*[a-z]\)|\s*$)',
                    re.DOTALL
                )
                alt_matches = alt_pattern.findall(body)
                if alt_matches:
                    logger.info(f"Found {len(alt_matches)} alt format sub-points")
                    for letter, text in alt_matches:
                        testo_clean = _clean_latex_line(text.strip())
                        punti = _estrai_punti(letter + ') ' + testo_clean)
                        sottopunti.append({
                            'label': letter + ')',
                            'testo': testo_clean,
                            'opzioni': [],
                            'punti': punti,
                        })
            
        if not sottopunti:
            free_items = re.compile(
                r'\\item(?:\[([^\]]*)\])?\s*(.*?)(?=\\item(?:\[|\s)|\\end\{|$)',
                re.DOTALL
            )
            auto_idx = 0
            for im in free_items.finditer(body):
                skip = any(s <= im.start() < e for s, e in used_ranges)
                if skip:
                    continue
                explicit = im.group(1)
                label    = explicit.strip() if explicit is not None else (LETTERE[auto_idx % 26] + ')')
                auto_idx += 1
                raw_text    = im.group(2).strip()
                testo_clean = _clean_latex_line(raw_text)
                testo_clean = re.sub(r'\n{2,}', '\n', testo_clean).strip()
                sottopunti.append({
                    'label':   label,
                    'testo':   testo_clean,
                    'opzioni': [],
                    'punti':   _estrai_punti(label + ' ' + testo_clean),
                })

        # Last resort: if still no sub-points found, try manual split
        if not sottopunti:
            # Look for patterns like "a) something" followed by "b) something"
            split_pattern = re.compile(r'([a-z]\)[^a-z]*?)(?=[a-z]\)|$)', re.DOTALL)
            manual_matches = split_pattern.findall(body)
            if manual_matches and len(manual_matches) > 1:
                logger.info(f"Manual split found {len(manual_matches)} sub-points")
                for match in manual_matches:
                    # Extract letter and text
                    letter_match = re.match(r'([a-z]\))(.*)', match.strip(), re.DOTALL)
                    if letter_match:
                        letter = letter_match.group(1)
                        text = letter_match.group(2).strip()
                        testo_clean = _clean_latex_line(text)
                        punti = _estrai_punti(letter + ' ' + testo_clean)
                        sottopunti.append({
                            'label': letter,
                            'testo': testo_clean,
                            'opzioni': [],
                            'punti': punti,
                        })

        if titolo_ex or sottopunti:
            logger.info(f"Found exercise: '{titolo_ex}' with {len(sottopunti)} sub-points")
            for i, sp in enumerate(sottopunti):
                logger.info(f"  Sub-point {i+1}: label='{sp.get('label')}', testo='{sp.get('testo', '')[:50]}...'")
            data['esercizi'].append({
                'titolo':      titolo_ex,
                'testo_intro': testo_intro,
                'sottopunti':  sottopunti,
            })

    return data


# ── EXPORT PRINCIPALE ──────────────────────────────────────────────────────────

def latex_to_docx_via_ai(codice_latex: str, con_griglia: bool = True) -> tuple[bytes | None, str | None]:
    """
    Converte un documento LaTeX in un file Word (.docx).
    Restituisce (docx_bytes, None) in caso di successo,
    oppure (None, messaggio_errore) in caso di fallimento.

    Parametri:
        codice_latex  – il LaTeX completo (preambolo + corpo)
        con_griglia   – se True, aggiunge la griglia di valutazione
    """
    logger.info("=== Starting LaTeX to DOCX conversion ===")
    logger.info(f"Input LaTeX length: {len(codice_latex)} chars")
    logger.info(f"Grid enabled: {con_griglia}")
    
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        logger.info("python-docx imports successful")
    except ImportError as e:
        logger.error(f"python-docx import failed: {e}")
        return None, "python-docx non installato. Esegui: pip install python-docx"

    # Pre-render TikZ/pgfplots blocks to PNG images
    logger.info("Step 1: Pre-rendering TikZ/pgfplots graphs...")
    codice_latex, _graphs = _pre_render_graphs(codice_latex)
    logger.info(f"Graphs pre-rendered: {len(_graphs)} total")

    try:
        logger.info("Step 2: Parsing LaTeX structure...")
        data = _parse_latex_to_data(codice_latex)
        logger.info(f"Parsed {len(data.get('esercizi', []))} exercises")
    except Exception as e:
        import traceback
        logger.error(f"LaTeX parsing failed: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None, f"Errore parsing LaTeX: {e}\n{traceback.format_exc()}"

    try:
        logger.info("Step 3: Creating DOCX document...")
        doc        = DocxDocument()
        MARGIN_CM  = 1.5
        PAGE_W_DXA = int((21.0 - 2 * MARGIN_CM) / 2.54 * 1440)
        GRIGLIA_W_DXA = PAGE_W_DXA - 80
        logger.info(f"Document setup: margins={MARGIN_CM}cm, page_width={PAGE_W_DXA} DXA")

        for section in doc.sections:
            section.page_width    = Cm(21.0)
            section.page_height   = Cm(29.7)
            section.left_margin   = Cm(MARGIN_CM)
            section.right_margin  = Cm(MARGIN_CM)
            section.top_margin    = Cm(MARGIN_CM)
            section.bottom_margin = Cm(MARGIN_CM)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        logger.info("Document styling applied")

        # ── INTESTAZIONE PROFESSIONALE ────────────────────────────────────────────
        # Riga istituto (campo vuoto da compilare)
        p_ist = doc.add_paragraph()
        p_ist.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ist = p_ist.add_run("Istituto: _________________________________________")
        r_ist.font.size = Pt(9)
        r_ist.font.color.rgb = None   # grigio di sistema
        p_ist.paragraph_format.space_before = Pt(0)
        p_ist.paragraph_format.space_after  = Pt(2)

        # Titolo principale (grande, grassetto, centrato)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        rt = p.add_run(data.get('titolo', 'Verifica'))
        rt.bold = True
        rt.font.size = Pt(15)

        # Separatore orizzontale (simulato con paragrafo con bordo inferiore)
        p_sep = doc.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(2)
        p_sep.paragraph_format.space_after  = Pt(6)
        pPr = p_sep._p.get_or_add_pPr()
        pBdr = _OE('w:pBdr')
        bdr_el = _OE('w:bottom')
        bdr_el.set(_qn('w:val'), 'single')
        bdr_el.set(_qn('w:sz'), '6')
        bdr_el.set(_qn('w:space'), '1')
        bdr_el.set(_qn('w:color'), '444444')
        pBdr.append(bdr_el)
        pPr.append(pBdr)

        # Tabella intestazione (Nome / Classe / Data — 3 colonne)
        _cw3 = [
            int(PAGE_W_DXA * 0.40),
            int(PAGE_W_DXA * 0.30),
            PAGE_W_DXA - int(PAGE_W_DXA * 0.40) - int(PAGE_W_DXA * 0.30),
        ]
        hdr_tbl = doc.add_table(rows=1, cols=3)
        hdr_el  = hdr_tbl._tbl
        tblPr_h = hdr_el.find(_qn('w:tblPr'))
        if tblPr_h is None:
            tblPr_h = _OE('w:tblPr')
            hdr_el.insert(0, tblPr_h)

        for tag_h, attrs_h in [
            ('w:tblW',      {'w:w': str(PAGE_W_DXA), 'w:type': 'dxa'}),
            ('w:tblLayout', {'w:type': 'fixed'}),
        ]:
            el_h = _OE(tag_h)
            for k_h, v_h in attrs_h.items():
                el_h.set(_qn(k_h), v_h)
            ex_h = tblPr_h.find(_qn(tag_h))
            if ex_h is not None:
                tblPr_h.remove(ex_h)
            tblPr_h.append(el_h)

        tblB_h = _OE('w:tblBorders')
        for side_h in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b_h = _OE(f'w:{side_h}')
            b_h.set(_qn('w:val'), 'nil')
            tblB_h.append(b_h)
        ex_b_h = tblPr_h.find(_qn('w:tblBorders'))
        if ex_b_h is not None:
            tblPr_h.remove(ex_b_h)
        tblPr_h.append(tblB_h)
        _fix_tbl_grid(hdr_el, _cw3, _qn, _OE)

        hdr_cells_data = [
            [("Nome: ", "_____________________________")],
            [("Classe: ", "_______________")],
            [("Data: ", "_______________")],
        ]
        for ci_h, (runs, cw_val) in enumerate(zip(hdr_cells_data, _cw3)):
            cell_h = hdr_tbl.cell(0, ci_h)
            tc_h   = cell_h._tc
            tcPr_h = tc_h.get_or_add_tcPr()
            tcW_h  = _OE('w:tcW')
            tcW_h.set(_qn('w:w'), str(cw_val))
            tcW_h.set(_qn('w:type'), 'dxa')
            ex_w_h = tcPr_h.find(_qn('w:tcW'))
            if ex_w_h is not None:
                tcPr_h.remove(ex_w_h)
            tcPr_h.append(tcW_h)

            tcBdr_h = _OE('w:tcBorders')
            for s_h in ('top', 'left', 'bottom', 'right'):
                b2_h = _OE(f'w:{s_h}')
                b2_h.set(_qn('w:val'), 'nil')
                tcBdr_h.append(b2_h)
            ex_bdr_h = tcPr_h.find(_qn('w:tcBorders'))
            if ex_bdr_h is not None:
                tcPr_h.remove(ex_bdr_h)
            tcPr_h.append(tcBdr_h)

            p_h = cell_h.paragraphs[0]
            for lbl_h, line_h in runs:
                r1_h = p_h.add_run(lbl_h)
                r1_h.bold = True
                r1_h.font.size = Pt(10)
                r2_h = p_h.add_run(line_h)
                r2_h.font.size = Pt(10)

        # Nota intestazione (durata, istruzioni)
        nota = data.get('intestazione_nota', '')
        if nota:
            p3 = doc.add_paragraph()
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_before = Pt(6)
            p3.paragraph_format.space_after  = Pt(2)
            r3 = p3.add_run(nota)
            r3.italic = True
            r3.font.size = Pt(10)

        logger.info("Step 4: Building document content...")
        logger.info(f"Processing {len(data.get('esercizi', []))} exercises with graphs")

        # ── ESERCIZI ──────────────────────────────────────────────────────────────
        for _ex_idx, ex in enumerate(data.get('esercizi', [])):
            logger.debug(f"Processing exercise {_ex_idx + 1}: {ex.get('titolo', 'Untitled')}")
            pe = doc.add_paragraph()
            pe.paragraph_format.space_before = Pt(3) if _ex_idx == 0 else Pt(6)
            pe.paragraph_format.space_after  = Pt(2)
            rt = pe.add_run(ex.get('titolo', ''))
            rt.bold = True
            rt.font.size = Pt(12)

            intro = ex.get('testo_intro', '').strip()
            if intro:
                if _GRAPH_RE.search(intro):
                    _add_content_with_graphs(
                        doc, intro, _graphs,
                        space_before_pt=0, space_after_pt=4, base_size_pt=11,
                    )
                else:
                    pi = doc.add_paragraph()
                    pi.paragraph_format.space_before = Pt(0)
                    pi.paragraph_format.space_after  = Pt(4)
                    pi.paragraph_format.left_indent  = Cm(0.0)
                    _add_rich_runs(pi, intro, base_size_pt=11)

            for sp in ex.get('sottopunti', []):
                label    = sp.get('label', '').strip()
                testo    = sp.get('testo', '').strip()
                opzioni  = sp.get('opzioni', [])
                punti_sp = sp.get('punti', '')
                
                # DEBUG: Log what we're processing
                logger.info(f"Processing sub-point: label='{label}', testo='{testo[:50]}...', punti='{punti_sp}'")

                # Check if this sub-point text contains a graph placeholder
                if testo and _GRAPH_RE.search(testo):
                    # Add label paragraph first if needed
                    if label:
                        pl = doc.add_paragraph()
                        pl.paragraph_format.left_indent  = Cm(0.5)
                        pl.paragraph_format.space_before = Pt(3)
                        pl.paragraph_format.space_after  = Pt(0)
                        rl = pl.add_run(label + "  ")
                        rl.bold = True
                        rl.font.size = Pt(11)
                    _add_content_with_graphs(
                        doc, testo, _graphs,
                        left_indent_cm=0.5, space_before_pt=2, space_after_pt=8,  # Increased spacing
                        base_size_pt=11,
                    )
                else:
                    ps = doc.add_paragraph()
                    ps.paragraph_format.left_indent  = Cm(0.5)
                    ps.paragraph_format.space_before = Pt(3)
                    ps.paragraph_format.space_after  = Pt(8)  # Increased spacing for better separation
                    if label:
                        rl = ps.add_run(label + "  ")
                        rl.bold = True
                        rl.font.size = Pt(11)
                    if testo:
                        _add_rich_runs(ps, testo, base_size_pt=11)

                if opzioni:
                    for opt in opzioni:
                        po = doc.add_paragraph()
                        po.paragraph_format.left_indent  = Cm(1.3)
                        po.paragraph_format.space_before = Pt(0)
                        po.paragraph_format.space_after  = Pt(2)
                        _add_rich_runs(po, str(opt), base_size_pt=11)
                elif not (testo and _GRAPH_RE.search(testo)):
                    # Spazio risposta proporzionale ai punti: min 10pt, max 36pt
                    try:
                        _pts_val = float(str(punti_sp).replace(',', '.'))
                    except (ValueError, TypeError):
                        _pts_val = 0
                    # Regola: ~3pt per punto, minimo 10, massimo 36
                    _space_after = min(36, max(10, int(_pts_val * 3)))
                    pr = doc.add_paragraph()
                    pr.paragraph_format.left_indent  = Cm(0.5)
                    pr.paragraph_format.space_before = Pt(1)
                    pr.paragraph_format.space_after  = Pt(_space_after)

        # Griglia di valutazione
        esercizi_parsed = parse_esercizi(codice_latex) if con_griglia else []
        if con_griglia and esercizi_parsed:
            pg = doc.add_paragraph()
            rg = pg.add_run("Griglia Punteggi")
            rg.bold = True
            rg.font.size = Pt(12)
            pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pg.paragraph_format.space_before = Pt(18)  # Increased spacing before table
            pg.paragraph_format.space_after = Pt(6)   # Add spacing after title

            n_sotto_totali = sum(len(ex['items']) for ex in esercizi_parsed)

            if n_sotto_totali > 12:
                # Griglia compatta (un colonne per esercizio, non per sottopunto)
                row_es  = ["Es."]
                row_max = ["Max"]
                for ex in esercizi_parsed:
                    row_es.append(f"Es. {ex['num']}")
                    tot_ex = 0
                    for _, pts in ex['items']:
                        try:
                            tot_ex += float(pts.replace(',', '.'))
                        except Exception:
                            pass
                    row_max.append(
                        str(int(tot_ex)) if tot_ex == int(tot_ex) else str(round(tot_ex, 1))
                    )
                row_es.append("Tot")
                tot_pts = sum(
                    float(x.replace(',', '.'))
                    for x in row_max[1:]
                    if x.replace(',', '.').replace('.', '').isdigit()
                )
                row_max.append(
                    str(int(tot_pts)) if tot_pts == int(tot_pts) else str(round(tot_pts, 1))
                )
                row_sotto_c = [""] * len(row_es)
                _build_griglia_xml(doc, row_es, row_sotto_c, row_max, GRIGLIA_W_DXA)
            else:
                # Griglia dettagliata (una colonna per sottopunto)
                row_es    = ["Es."]
                row_sotto = ["Sotto."]
                row_max   = ["Max"]
                for ex in esercizi_parsed:
                    for label_clean, pts in ex['items']:
                        lbl = label_clean.replace('*', '').replace(')', '').strip()[:3]
                        row_es.append(str(ex['num']))
                        row_sotto.append(lbl)
                        row_max.append(str(pts))
                row_es.append("Tot")
                row_sotto.append("")
                tot_pts = 0
                for x in row_max[1:]:
                    try:
                        tot_pts += float(x.replace(',', '.'))
                    except Exception:
                        pass
                row_max.append(
                    str(int(tot_pts)) if tot_pts == int(tot_pts) else str(round(tot_pts, 1))
                )
                if len(row_es) >= 3:
                    _build_griglia_xml(doc, row_es, row_sotto, row_max, GRIGLIA_W_DXA)

        logger.info("Step 5: Saving DOCX document...")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        logger.info(f"DOCX conversion SUCCESS - Final file size: {len(docx_bytes)} bytes")
        logger.info("=== LaTeX to DOCX conversion completed successfully ===")
        return docx_bytes, None
    except Exception as e:
        import traceback
        logger.error(f"DOCX conversion FAILED: {e}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return None, f"Errore durante creazione DOCX: {e}"


# ── RUBRICA VALUTAZIONE → DOCX ───────────────────────────────────────────────────

def _rubrica_to_docx(
    rubrica_testo: str,
    materia: str = "",
    livello: str = "",
    argomento: str = "",
) -> bytes | None:
    """
    Genera un DOCX professionale dalla rubrica di valutazione.
    Input: testo pipe-delimited dall'AI (FASCE: + ESERCIZI:).
    Output: bytes del file .docx.
    """
    try:
        import io
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # ── Parse pipe-delimited content ─────────────────────────────────────
        fasce: list[list[str]] = []
        esercizi: list[list[str]] = []
        section = None
        for raw_line in rubrica_testo.strip().split('\n'):
            line = raw_line.strip()
            if line == 'FASCE:':
                section = 'fasce'
                continue
            if line == 'ESERCIZI:':
                section = 'esercizi'
                continue
            if '|' not in line or line.startswith('#') or line.startswith('('):
                continue
            parts = [p.strip() for p in line.split('|')]
            if section == 'fasce' and len(parts) >= 6:
                fasce.append(parts[:6])
            elif section == 'esercizi' and len(parts) >= 6:
                esercizi.append(parts[:6])

        # ── Colors ───────────────────────────────────────────────────────────
        C_TEAL      = RGBColor(0x0A, 0x8F, 0x72)
        C_TEAL_DARK = RGBColor(0x07, 0x6B, 0x56)
        C_GRAY      = RGBColor(0x6B, 0x72, 0x80)
        C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
        C_DARK      = RGBColor(0x1A, 0x1A, 0x2E)

        FASCIA_BG = ['D4EDDA', 'D1ECF1', 'FFF3CD', 'F8D7DA']  # verde, azzurro, giallo, rosso

        doc = Document()

        # ── Page layout ───────────────────────────────────────────────────────
        sec = doc.sections[0]
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

        # ── Title block ───────────────────────────────────────────────────────
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_title.add_run('GRIGLIA DI VALUTAZIONE PER COMPETENZE')
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = C_TEAL_DARK

        meta_parts = [x for x in [materia, livello] if x]
        if meta_parts:
            p_meta = doc.add_paragraph()
            p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rm = p_meta.add_run(' · '.join(meta_parts))
            rm.font.size = Pt(10)
            rm.font.color.rgb = C_GRAY

        if argomento:
            p_arg = doc.add_paragraph()
            p_arg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ra = p_arg.add_run(f'Argomento: {argomento}')
            ra.font.size = Pt(9.5)
            ra.font.color.rgb = C_GRAY
            ra.italic = True

        # Horizontal rule (via paragraph border)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        p_rule = doc.add_paragraph()
        p_rule.paragraph_format.space_before = Pt(2)
        p_rule.paragraph_format.space_after  = Pt(2)
        pPr = p_rule._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom_bdr = OxmlElement('w:bottom')
        bottom_bdr.set(qn('w:val'), 'single')
        bottom_bdr.set(qn('w:sz'), '12')
        bottom_bdr.set(qn('w:space'), '1')
        bottom_bdr.set(qn('w:color'), '0A8F72')
        pBdr.append(bottom_bdr)
        pPr.append(pBdr)

        # ── Section 1: Fasce di Voto ──────────────────────────────────────────
        p_s1 = doc.add_paragraph()
        p_s1.paragraph_format.space_before = Pt(8)
        p_s1.paragraph_format.space_after  = Pt(4)
        rs1 = p_s1.add_run('Fasce di Voto')
        rs1.bold = True
        rs1.font.size = Pt(11)
        rs1.font.color.rgb = C_TEAL

        HDR_FASCE = ['Fascia', 'Punti', 'Voto', 'Comprensione dei contenuti',
                     'Applicazione delle conoscenze', 'Esposizione e precisione']
        W_FASCE   = [Cm(2.2), Cm(2.0), Cm(1.4), Cm(4.2), Cm(4.2), Cm(3.0)]

        tbl1 = doc.add_table(rows=1 + len(fasce), cols=6)
        tbl1.style = 'Table Grid'

        # Header row
        hdr_row = tbl1.rows[0]
        for ci, (hdr_txt, w) in enumerate(zip(HDR_FASCE, W_FASCE)):
            cell = hdr_row.cells[ci]
            cell.width = w
            _set_cell_bg(cell, '0A7A62')
            _set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = p.add_run(hdr_txt)
            rr.bold = True
            rr.font.size = Pt(8.5)
            rr.font.color.rgb = C_WHITE

        # Data rows
        for ri, row_data in enumerate(fasce):
            row = tbl1.rows[ri + 1]
            bg = FASCIA_BG[ri] if ri < len(FASCIA_BG) else 'FFFFFF'
            for ci, val in enumerate(row_data[:6]):
                cell = row.cells[ci]
                cell.width = W_FASCE[ci]
                _set_cell_bg(cell, bg)
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if ci < 3
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                rr = p.add_run(val)
                rr.font.size = Pt(8.5)
                if ci == 0:
                    rr.bold = True
                    rr.font.color.rgb = C_DARK

        # ── Section 2: Griglia per Esercizio ─────────────────────────────────
        if esercizi:
            p_s2 = doc.add_paragraph()
            p_s2.paragraph_format.space_before = Pt(10)
            p_s2.paragraph_format.space_after  = Pt(4)
            rs2 = p_s2.add_run('Griglia per Esercizio')
            rs2.bold = True
            rs2.font.size = Pt(11)
            rs2.font.color.rgb = C_TEAL

            HDR_ES = ['Es.', 'Argomento', 'Punti', 'Criterio valutato',
                      'Risposta eccellente (9-10)', 'Risposta sufficiente (6)']
            W_ES   = [Cm(1.0), Cm(3.2), Cm(1.4), Cm(3.5), Cm(4.4), Cm(3.5)]

            tbl2 = doc.add_table(rows=1 + len(esercizi), cols=6)
            tbl2.style = 'Table Grid'

            hdr_row2 = tbl2.rows[0]
            for ci, (hdr_txt, w) in enumerate(zip(HDR_ES, W_ES)):
                cell = hdr_row2.cells[ci]
                cell.width = w
                _set_cell_bg(cell, '0A7A62')
                _set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = p.add_run(hdr_txt)
                rr.bold = True
                rr.font.size = Pt(8.5)
                rr.font.color.rgb = C_WHITE

            for ri, row_data in enumerate(esercizi):
                row = tbl2.rows[ri + 1]
                bg = 'F0FAF8' if ri % 2 == 0 else 'FFFFFF'
                for ci, val in enumerate(row_data[:6]):
                    cell = row.cells[ci]
                    cell.width = W_ES[ci]
                    _set_cell_bg(cell, bg)
                    _set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if ci < 3
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    rr = p.add_run(val)
                    rr.font.size = Pt(8.5)
                    if ci == 0:
                        rr.bold = True

        # ── Footer ────────────────────────────────────────────────────────────
        p_foot = doc.add_paragraph()
        p_foot.paragraph_format.space_before = Pt(10)
        p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = p_foot.add_run(
            'Generato con VerificAI  ·  Allineato alle Linee Guida MIM sulla valutazione per competenze'
        )
        rf.font.size = Pt(7.5)
        rf.font.color.rgb = C_GRAY
        rf.italic = True

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except Exception:
        return None
